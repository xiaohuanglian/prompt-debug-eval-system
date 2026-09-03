"""All step panel widgets for the Prompt GUI application."""

import os
import re
import json
import traceback
from datetime import datetime

from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QApplication, QTextEdit, QLineEdit, QComboBox, QSpinBox, QGroupBox,
    QTableWidget, QTableWidgetItem, QHeaderView, QSplitter,
    QProgressBar, QMessageBox, QFileDialog, QCheckBox,
    QTabWidget, QListWidget, QListWidgetItem, QTreeWidget, QTreeWidgetItem, QAbstractItemView,
    QGridLayout, QFrame, QScrollArea, QPlainTextEdit, QSizePolicy,
    QShortcut, QDialog, QDialogButtonBox, QInputDialog,
)
from PyQt5.QtCore import Qt, QTimer, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor, QPalette, QTextCharFormat, QSyntaxHighlighter
import docx

from ._project_paths import PROJECT_ROOT

from code.pipeline.model_resolver import load_model
from code.pipeline.prompt_generator import (
    extract_placeholders,
    normalize_placeholders_to_double_braces,
    strip_auto_output_schema,
)
from code.pipeline.eval_code_generator import generate_eval_script
from code.pipeline.version_manager import (
    VersionManager,
    build_scratch_scenario_name,
    is_generic_scratch_scenario_name,
    normalize_scenario_name,
)

from .workers import (
    ListModelsWorker, RawDocsAnalyzeWorker, GeneratePromptWorker,
    ContinuePromptOptimizeWorker,
    GenerateEvalDataAutoWorker, GenerateEvalDataSeedWorker,
    RunEvalWorker, AnalyzeErrorsWorker, SuggestImprovementsWorker,
    ConvertTemplateWorker, GatewayValidateWorker, GatewayDeployWorker,
    convert_python_to_go_template,
)


# ─── Styles ─────────────────────────────────────────────────────────────────

STYLE = """
QGroupBox {
    font-weight: bold;
    border: 1px solid #cccccc;
    border-radius: 6px;
    margin-top: 12px;
    padding-top: 16px;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 10px;
    padding: 0 5px;
}
QPushButton {
    padding: 6px 16px;
    border-radius: 4px;
    border: 1px solid #aaa;
    background: #f0f0f0;
    min-height: 24px;
}
QPushButton:hover { background: #e0e0e0; }
QPushButton:pressed { background: #d0d0d0; }
QPushButton:disabled { background: #f5f5f5; color: #999; }
QPushButton.btn-primary {
    background: #1976D2;
    color: white;
    border: 1px solid #1565C0;
}
QPushButton.btn-primary:hover { background: #1565C0; }
QPushButton.btn-success {
    background: #388E3C;
    color: white;
    border: 1px solid #2E7D32;
}
QPushButton.btn-success:hover { background: #2E7D32; }
QPushButton.btn-danger {
    background: #D32F2F;
    color: white;
    border: 1px solid #C62828;
}
QPushButton.btn-danger:hover { background: #C62828; }
QPushButton.btn-warning {
    background: #F57C00;
    color: white;
    border: 1px solid #E65100;
}
QPushButton.btn-warning:hover { background: #E65100; }
QTextEdit, QPlainTextEdit, QLineEdit {
    border: 1px solid #ccc;
    border-radius: 4px;
    padding: 4px;
}
QTableWidget {
    border: 1px solid #ccc;
    gridline-color: #e0e0e0;
}
QHeaderView::section {
    background: #f5f5f5;
    border: 1px solid #ddd;
    padding: 4px;
}
QProgressBar {
    border: 1px solid #ccc;
    border-radius: 4px;
    text-align: center;
}
QProgressBar::chunk {
    background: #1976D2;
    border-radius: 3px;
}
"""


def _resolve_scratch_scenario_name(parent, scenario_name: str):
    """
    For new scratch work, prevent saving into the shared bare "测试" bucket.
    Returns a concrete scenario name, or None if the user cancels.
    """
    name = normalize_scenario_name(scenario_name)
    if not is_generic_scratch_scenario_name(name):
        return name

    note, ok = QInputDialog.getText(
        parent,
        "填写临时测试用途",
        "场景名“测试”会复用同一个版本目录。\n"
        "请输入这次临时测试的具体用途，系统将保存为“测试_具体用途”：",
    )
    if not ok:
        return None
    try:
        return build_scratch_scenario_name(note)
    except ValueError as exc:
        QMessageBox.warning(parent, "提示", str(exc))
        return None


# ─── Helper widgets ─────────────────────────────────────────────────────────

def _make_labels_selectable(widget: QWidget):
    """递归遍历 widget 树，将所有 QLabel 设为文字可选中复制。"""
    from PyQt5.QtWidgets import QLabel
    for child in widget.findChildren(QLabel):
        child.setTextInteractionFlags(Qt.TextSelectableByMouse)
    # 自身也可能是 QLabel
    if isinstance(widget, QLabel):
        widget.setTextInteractionFlags(Qt.TextSelectableByMouse)


def _get_main_window_context(widget: QWidget) -> dict:
    parent = widget.parent()
    while parent and not hasattr(parent, 'context'):
        parent = parent.parent()
    if parent and hasattr(parent, 'context'):
        return parent.context
    return {}


def _notify_long_task_done(widget: QWidget, title: str, message: str, force: bool = False):
    """Notify the user when a long-running worker finishes."""
    context = _get_main_window_context(widget)
    if not force and not context.get("notify_on_finish", True):
        return
    window = widget.window()
    try:
        QApplication.alert(window, 0)
    except Exception:
        pass
    try:
        import winsound
        winsound.MessageBeep(winsound.MB_ICONASTERISK)
    except Exception:
        pass
    QMessageBox.information(widget, title, message)


class ContinuePromptDialog(QDialog):
    """Collect continue-scenario suggestions with wrapping and template choice."""

    def __init__(self, templates: list, parent=None, initial_suggestions: str = ""):
        super().__init__(parent)
        self.setWindowTitle("优化 Prompt — 输入修改意见")
        self.resize(560, 460)
        layout = QVBoxLayout(self)

        desc = QLabel(
            "该场景已有完整评测结果。\n\n"
            "在下方输入你的修改意见，系统将后台生成候选 Prompt。\n"
            "无需修改可直接点 OK，将使用当前 Prompt。"
        )
        desc.setWordWrap(True)
        layout.addWidget(desc)

        tpl_row = QHBoxLayout()
        tpl_row.addWidget(QLabel("参考模板:"))
        self.template_combo = QComboBox()
        self.template_combo.addItem("(无 - 自由生成)", "")
        self._template_map = {}
        for item in templates:
            display = f"{item['name']}  ({len(item['content'])} chars)"
            self.template_combo.addItem(display, item["name"])
            self._template_map[item["name"]] = item["content"]
        if templates:
            self.template_combo.setCurrentIndex(1)
        tpl_row.addWidget(self.template_combo, 1)
        layout.addLayout(tpl_row)

        self.suggestions_edit = QTextEdit()
        self.suggestions_edit.setMinimumHeight(230)
        if initial_suggestions:
            self.suggestions_edit.setPlainText(initial_suggestions)
        self.suggestions_edit.setPlaceholderText("输入你希望新版 Prompt 修改的地方...")
        _enable_width_wrap(self.suggestions_edit)
        layout.addWidget(self.suggestions_edit, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def suggestions(self) -> str:
        return self.suggestions_edit.toPlainText().strip()

    def template_ref(self) -> str:
        name = self.template_combo.currentData()
        return self._template_map.get(name, "") if name else ""

    def template_name(self) -> str:
        return self.template_combo.currentData() or ""


class PromptRevisionDialog(QDialog):
    """Collect prompt revision directions from Step 4."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("修改 Prompt")
        self.resize(560, 360)
        layout = QVBoxLayout(self)

        desc = QLabel("输入你希望新版 Prompt 改进的方向，点击 OK 后会回到 Step 3 并自动重新生成 Prompt。")
        desc.setWordWrap(True)
        layout.addWidget(desc)

        self.revision_edit = QTextEdit()
        self.revision_edit.setMinimumHeight(220)
        self.revision_edit.setPlaceholderText("例如：不要编造不存在的错误类型；输出要更短；必须根据 target_language 切换语言...")
        _enable_width_wrap(self.revision_edit)
        layout.addWidget(self.revision_edit, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def revision_text(self) -> str:
        return self.revision_edit.toPlainText().strip()


def _enable_width_wrap(editor):
    """Wrap long lines to the widget width so text is readable without horizontal scrolling."""
    if isinstance(editor, QTextEdit):
        editor.setLineWrapMode(QTextEdit.WidgetWidth)
    elif isinstance(editor, QPlainTextEdit):
        editor.setLineWrapMode(QPlainTextEdit.WidgetWidth)
    editor.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)


class LogTextEdit(QTextEdit):
    """Read-only log output widget."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setReadOnly(True)
        self.setFont(QFont("Consolas", 9))
        self.setStyleSheet("background: #1e1e1e; color: #d4d4d4;")
        self.document().setMaximumBlockCount(1000)

    def log(self, message: str):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.append(f"[{timestamp}] {message}")

    def log_error(self, message: str):
        self.log(f"ERROR: {message}")


class StepNavWidget(QListWidget):
    """Sidebar step navigation."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.steps = [
            "1. 模型选择",
            "2. 需求输入",
            "3. 生成 Prompt",
            "4. 评测数据",
            "5. 评测脚本",
            "6. 运行评测",
            "7. 结果分析",
            "8. 迭代优化",
            "9. 上线部署",
        ]
        for s in self.steps:
            self.addItem(s)
        self.setFixedWidth(160)
        self.setCurrentRow(0)
        self.setStyleSheet("""
            QListWidget {
                border: none;
                background: #fafafa;
                font-size: 13px;
            }
            QListWidget::item {
                padding: 10px 12px;
                border-bottom: 1px solid #eee;
            }
            QListWidget::item:selected {
                background: #1976D2;
                color: white;
            }
            QListWidget::item:hover:!selected {
                background: #e3f2fd;
            }
        """)


class JsonHighlighter(QSyntaxHighlighter):
    """Simple JSON syntax highlighter."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.rules = []
        fmt_key = QTextCharFormat()
        fmt_key.setForeground(QColor("#1565C0"))
        fmt_key.setFontWeight(QFont.Bold)
        self.rules.append((r'"[^"]*"\s*:', fmt_key))

        fmt_str = QTextCharFormat()
        fmt_str.setForeground(QColor("#2E7D32"))
        self.rules.append((r'"[^"]*"', fmt_str))

        fmt_num = QTextCharFormat()
        fmt_num.setForeground(QColor("#E65100"))
        self.rules.append((r'\b\d+\.?\d*\b', fmt_num))

        fmt_bool = QTextCharFormat()
        fmt_bool.setForeground(QColor("#6A1B9A"))
        self.rules.append((r'\b(true|false|null)\b', fmt_bool))

    def highlightBlock(self, text):
        for pattern, fmt in self.rules:
            import re
            for match in re.finditer(pattern, text):
                start, end = match.start(), match.end()
                self.setFormat(start, end - start, fmt)


# ─── Step 1: Model Panel ────────────────────────────────────────────────────

class ModelPanel(QWidget):
    """Step 1: Model selection."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.helper_llm = None
        self.target_llm = None
        self.helper_info = None
        self.target_info = None
        self.max_workers = 8
        self.available_models = []
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)

        title = QLabel("Step 1: 模型选择")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        desc = QLabel("分别选择 Helper 模型（生成 prompt/评测数据/优化建议）和 Target 模型（被评测的模型）。")
        desc.setWordWrap(True)
        layout.addWidget(desc)

        # API Keys path info
        api_path = os.path.join(PROJECT_ROOT, "code", "models", "api_keys.py")
        path_label = QLabel(f"API 配置: {api_path}")
        path_label.setStyleSheet("color: #666; font-size: 11px;")
        path_label.setWordWrap(True)
        layout.addWidget(path_label)

        # Model selection
        grid = QGridLayout()
        grid.setSpacing(10)

        grid.addWidget(QLabel("Helper 模型:"), 0, 0)
        self.helper_combo = QComboBox()
        self.helper_combo.setMinimumWidth(250)
        grid.addWidget(self.helper_combo, 0, 1)

        grid.addWidget(QLabel("Target 模型:"), 1, 0)
        self.target_combo = QComboBox()
        self.target_combo.setMinimumWidth(250)
        grid.addWidget(self.target_combo, 1, 1)

        grid.addWidget(QLabel("评测并发数:"), 2, 0)
        self.workers_spin = QSpinBox()
        self.workers_spin.setRange(1, 32)
        self.workers_spin.setValue(4)
        self.workers_spin.setToolTip("评测会并发调用 Target 模型。若模型限流或不支持并发，请调回 1。")
        grid.addWidget(self.workers_spin, 2, 1)

        grid.addWidget(QLabel("长任务完成提醒:"), 3, 0)
        self.notify_check = QCheckBox("开启弹窗和提示音")
        self.notify_check.setChecked(True)
        self.notify_check.setToolTip("生成需求、生成 Prompt、生成评测数据、运行评测完成后提醒。")
        grid.addWidget(self.notify_check, 3, 1)

        layout.addLayout(grid)

        # Buttons
        btn_layout = QHBoxLayout()
        self.refresh_btn = QPushButton("刷新模型列表")
        self.refresh_btn.clicked.connect(self.refresh_models)
        btn_layout.addWidget(self.refresh_btn)

        self.confirm_btn = QPushButton("确认选择")
        self.confirm_btn.setObjectName("confirmBtn")
        self.confirm_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 24px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #1565C0; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.confirm_btn.clicked.connect(self._confirm_selection)
        btn_layout.addWidget(self.confirm_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # Status
        self.status_label = QLabel("")
        self.status_label.setStyleSheet("color: #666;")
        layout.addWidget(self.status_label)

        layout.addStretch()

    def refresh_models(self):
        self.status_label.setText("正在加载模型列表...")
        self.confirm_btn.setEnabled(False)
        self.refresh_btn.setEnabled(False)
        self.worker = ListModelsWorker()
        self.worker.finished.connect(self._on_models_loaded)
        self.worker.error.connect(self._on_models_error)
        self.worker.start()

    def _on_models_loaded(self, models):
        self.available_models = models
        self.helper_combo.clear()
        self.target_combo.clear()
        if not models:
            self.helper_combo.addItem("未找到可用模型，请配置 API Key")
            self.target_combo.addItem("未找到可用模型，请配置 API Key")
            self.status_label.setText("未找到可用模型。请在 code/models/api_keys.py 中配置 API Key。")
            self.status_label.setStyleSheet("color: red;")
            self.confirm_btn.setEnabled(False)
        else:
            for m in models:
                label = f"{m['prefix']} - {m['model_name']}"
                self.helper_combo.addItem(label, m)
                self.target_combo.addItem(label, m)

            if len(models) == 1:
                # Only one model available — auto-select and proceed to Step 2
                self.status_label.setText(f"自动选择唯一可用模型: {models[0]['model_name']}")
                self.status_label.setStyleSheet("color: #1565C0;")
                QApplication.processEvents()
                self._confirm_selection()
            else:
                self.status_label.setText(f"找到 {len(models)} 个可用模型。请选择 Helper 和 Target 模型。")
                self.status_label.setStyleSheet("color: green;")
                self.confirm_btn.setEnabled(True)
        self.refresh_btn.setEnabled(True)

    def _on_models_error(self, error_msg):
        self.status_label.setText(f"加载失败: {error_msg}")
        self.status_label.setStyleSheet("color: red;")
        self.refresh_btn.setEnabled(True)

    def _confirm_selection(self):
        if not self.available_models:
            QMessageBox.warning(self, "警告", "没有可用模型，请先配置 API Key。")
            return

        helper_data = self.helper_combo.currentData()
        target_data = self.target_combo.currentData()
        if not helper_data or not target_data:
            QMessageBox.warning(self, "警告", "请选择有效的模型。")
            return

        self.max_workers = self.workers_spin.value()

        self.status_label.setText("正在加载模型...")
        self.confirm_btn.setEnabled(False)
        self.refresh_btn.setEnabled(False)

        try:
            self.helper_llm = load_model(helper_data)
            self.target_llm = load_model(target_data)
            self.helper_info = helper_data
            self.target_info = target_data

            self.status_label.setText(
                f"✓ Helper: {helper_data['model_name']}  |  "
                f"Target: {target_data['model_name']}  |  并发: {self.max_workers}"
            )
            self.status_label.setStyleSheet("color: green; font-weight: bold;")

            # Emit signal or notify parent
            parent = self.parent()
            while parent and not hasattr(parent, 'on_step_complete'):
                parent = parent.parent()
            if parent and hasattr(parent, 'on_step_complete'):
                parent.on_step_complete(0, {
                    "helper_llm": self.helper_llm,
                    "target_llm": self.target_llm,
                    "helper_info": self.helper_info,
                    "target_info": self.target_info,
                    "max_workers": self.max_workers,
                    "notify_on_finish": self.notify_check.isChecked(),
                })

            QMessageBox.information(self, "完成", "模型选择完成，可以进入下一步。")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"模型加载失败: {e}")
            self.status_label.setText(f"模型加载失败: {e}")
            self.status_label.setStyleSheet("color: red;")

        self.confirm_btn.setEnabled(True)
        self.refresh_btn.setEnabled(True)


# ─── Step 2: Requirement Panel ──────────────────────────────────────────────


class RequirementPanel(QWidget):
    """Step 2: Input scenario name and requirements."""

    # 全局 Context 文档存储路径
    CONTEXT_DIR = os.path.join(PROJECT_ROOT, "data", "system_context")

    def __init__(self, parent=None):
        super().__init__(parent)
        self.helper_llm = None
        self.scenario_name = ""
        self.requirements = ""
        self.llm_requirements = ""
        self.selected_context_fields = set()
        self.context_field_specs = []
        self.context_selection_initialized = False
        self._setup_ui()

    def _get_context_files(self):
        """Return list of all global context file paths, sorted by name."""
        os.makedirs(self.CONTEXT_DIR, exist_ok=True)
        files = []
        for f in sorted(os.listdir(self.CONTEXT_DIR)):
            if not f.startswith("."):
                files.append(os.path.join(self.CONTEXT_DIR, f))
        return files

    def _context_exists(self):
        return len(self._get_context_files()) > 0

    def set_helper_llm(self, helper_llm):
        self.helper_llm = helper_llm

    # ─── Continue / resume methods ──────────────────────────────────────

    def _refresh_existing_scenarios(self):
        """Scan data/prompt/ for existing scenarios with saved versions."""
        prompt_dir = os.path.join(PROJECT_ROOT, "data", "prompt")
        self.continue_list.clear()
        self.continue_btn.setEnabled(False)
        self.continue_status.setText("")

        if not os.path.isdir(prompt_dir):
            self.continue_list.addItem("（暂无已有场景）")
            return

        found = False
        for scenario in sorted(os.listdir(prompt_dir)):
            scenario_path = os.path.join(prompt_dir, scenario)
            if not os.path.isdir(scenario_path) or scenario.startswith("."):
                continue
            versions = [f for f in os.listdir(scenario_path) if f.endswith(".py")]
            if not versions:
                continue

            # Determine status
            vm = VersionManager(scenario)
            v_list = vm.list_versions()
            latest_v = v_list[-1] if v_list else 0

            has_eval_data = False
            eval_path = os.path.join(PROJECT_ROOT, "data", "eval", f"{scenario}.json")
            if os.path.exists(eval_path):
                has_eval_data = True

            has_eval_script = False
            eval_script_path = os.path.join(PROJECT_ROOT, "code", "eval", f"eval_prompt_{scenario}.py")
            if os.path.exists(eval_script_path):
                has_eval_script = True

            has_result = False
            result_dir = os.path.join(PROJECT_ROOT, "data", "eval_result")
            if os.path.isdir(result_dir):
                for f in os.listdir(result_dir):
                    if f.startswith(scenario) and f.endswith(".json"):
                        has_result = True
                        break

            # Build status string
            status_parts = [f"v{latest_v}"]
            if has_eval_data:
                status_parts.append("有评测数据")
            if has_eval_script:
                status_parts.append("有评测脚本")
            if has_result:
                status_parts.append("有评测结果")

            label = f"{scenario}  |  {' | '.join(status_parts)}"
            item = QListWidgetItem(label)
            item.setData(Qt.UserRole, {
                "scenario": scenario,
                "latest_v": latest_v,
                "has_eval_data": has_eval_data,
                "has_eval_script": has_eval_script,
                "has_result": has_result,
            })
            self.continue_list.addItem(item)
            found = True

        if not found:
            self.continue_list.addItem("（暂无已有场景）")

    def _on_continue_scenario_selected(self, item):
        """When an existing scenario is selected, show status info."""
        data = item.data(Qt.UserRole)
        if not data:
            self.continue_btn.setEnabled(False)
            return

        s = data["scenario"]
        self.continue_btn.setEnabled(True)
        self.name_edit.setText(s)

        # Show detailed status
        lines = [f"场景: {s}"]
        vm = VersionManager(s)
        v_list = vm.list_versions()
        if v_list:
            lines.append(f"Prompt 版本: v{v_list[-1]}（共 {len(v_list)} 个版本）")
        if data["has_eval_data"]:
            lines.append("评测数据: ✅ 已保存")
        if data.get("has_eval_script"):
            lines.append("评测脚本: ✅ 已生成")
        if data["has_result"]:
            lines.append("评测结果: ✅ 已有历史结果")
        prompt_template = ""
        latest_v = v_list[-1] if v_list else data.get("latest_v")
        if latest_v:
            try:
                prompt_template = vm.load_version(latest_v)
            except Exception:
                prompt_template = ""
        synced, missing = self._sync_context_fields_from_prompt(prompt_template)
        if prompt_template:
            if synced:
                lines.append(f"已从当前 Prompt 自动勾选字段: {', '.join(synced)}")
            else:
                lines.append("当前 Prompt 未检测到可同步的 Context 占位符")
            if missing:
                lines.append(f"当前 Prompt 中有字段不在最新 Context 中: {', '.join(missing)}")
        self.continue_status.setPlainText("\n".join(lines))

    def _do_continue(self):
        """Load selected scenario and navigate to appropriate step."""
        item = self.continue_list.currentItem()
        if not item:
            QMessageBox.warning(self, "提示", "请先在列表中选择一个已有场景。")
            return
        data = item.data(Qt.UserRole)
        if not data:
            QMessageBox.warning(self, "提示", "场景数据无效。")
            return

        scenario = data["scenario"]
        vm = VersionManager(scenario)
        v_list = vm.list_versions()
        latest_v = v_list[-1] if v_list else 1

        # Load prompt template
        prompt_template = ""
        try:
            prompt_template = vm.load_version(latest_v)
        except Exception:
            pass

        # Load eval data
        eval_data = []
        eval_data_loaded = False
        try:
            eval_data = vm.load_eval_data()
            eval_data_loaded = True
        except Exception:
            pass

        # Load saved eval script if it exists
        saved_script = ""
        eval_script_path = os.path.join(vm.root_dir, "code", "eval", f"eval_prompt_{scenario}.py")
        if os.path.exists(eval_script_path):
            try:
                with open(eval_script_path, "r", encoding="utf-8") as f:
                    saved_script = f.read()
            except Exception:
                pass

        # Build context data
        ctx = {
            "scenario_name": scenario,
            "requirements": f"续接场景: {scenario}（从已有版本 v{latest_v} 继续）",
            "vm": vm,
            "continue_mode": True,
            "existing_version": latest_v,
            "version": latest_v,
            "prompt_template": prompt_template,
            "eval_data": eval_data,
            "all_context_fields": self._get_context_fields(),
            "context_fields": self._get_active_context_fields(),
            "selected_context_fields": self._get_active_context_fields(),
            "context_field_specs": self._get_active_context_field_specs(),
        }

        # ─── 续接：弹窗输入修改意见 → 自动优化 prompt → 转到 Step 3 ───
        parent = self.parent()
        while parent and not hasattr(parent, 'on_step_complete'):
            parent = parent.parent()
        if not parent:
            return

        selected_count = len(ctx.get("context_fields") or [])
        total_count = len(ctx.get("all_context_fields") or [])
        context_box = QMessageBox(self)
        context_box.setWindowTitle("确认 Context 字段")
        context_box.setIcon(QMessageBox.Question)
        context_box.setText(
            "继续此场景前，请先确认本次 Prompt 生成/调优允许使用的 Context 字段白名单。"
        )
        context_box.setInformativeText(
            f"当前已勾选 {selected_count}/{total_count} 个字段。\n\n"
            "Context 字段有没有要修改的地方？"
        )
        modify_btn = context_box.addButton("重新勾选字段", QMessageBox.YesRole)
        continue_btn = context_box.addButton("字段不改，继续", QMessageBox.NoRole)
        cancel_btn = context_box.addButton("取消", QMessageBox.RejectRole)
        context_box.setDefaultButton(continue_btn)
        context_box.exec_()
        clicked_btn = context_box.clickedButton()
        if clicked_btn == cancel_btn:
            return
        if clicked_btn == modify_btn:
            self.name_edit.setText(scenario)
            self.status_label.setText("请在 Context 字段列表中重新勾选本场景允许使用的字段；如有修改需求，填写「本次继续的修改需求」，再点击蓝色「继续此场景」。")
            self.status_label.setStyleSheet("color: #1565C0; font-weight: bold;")
            if hasattr(self, "context_field_list"):
                self.context_field_list.setFocus()
            if hasattr(self, "continue_revision_edit"):
                self.continue_revision_edit.setFocus()
            return
        has_full_assets = eval_data_loaded and data["has_result"]

        if has_full_assets:
            from code.pipeline.prompt_generator import load_system_templates
            inline_suggestions = ""
            if hasattr(self, "continue_revision_edit"):
                inline_suggestions = self.continue_revision_edit.toPlainText().strip()
            dialog = ContinuePromptDialog(
                load_system_templates(), self,
                initial_suggestions=inline_suggestions,
            )
            if dialog.exec_() == QDialog.Accepted and dialog.suggestions():
                if not self.helper_llm:
                    QMessageBox.warning(self, "提示",
                                       "Helper 模型未加载，无法自动优化。\n"
                                       "请先在 Step 1 中选择模型，再继续此场景。")
                else:
                    suggestions = dialog.suggestions()
                    template_ref = dialog.template_ref()
                    template_name = dialog.template_name()

                    parent.context.update(ctx)
                    if template_name:
                        parent.context["preferred_template_name"] = template_name
                    parent.prompt_panel.set_context(parent.context)
                    parent.prompt_panel.status_label.setText(
                        "正在后台根据修改意见生成候选 Prompt，请稍候..."
                    )
                    parent.prompt_panel.status_label.setStyleSheet(
                        "color: #F57C00; font-weight: bold;"
                    )
                    parent.prompt_panel.prompt_edit.setPlainText(
                        "正在后台生成候选 Prompt，界面可以继续响应..."
                    )
                    parent.go_to_step(3)

                    field_schema = {
                        "input_fields": ctx.get("context_field_specs") or [
                            {"name": name, "type": "string", "required": True}
                            for name in sorted(ctx["context_fields"])
                        ],
                        "output_fields": [],
                    }
                    self.continue_opt_worker = ContinuePromptOptimizeWorker(
                        self.helper_llm, scenario, prompt_template,
                        ctx["requirements"], suggestions,
                        template_ref=template_ref,
                        field_schema=field_schema,
                    )

                    def _on_continue_optimized(new_prompt):
                        ctx["prompt_template"] = new_prompt if new_prompt.strip() else prompt_template
                        ctx["candidate_reason"] = "已根据修改意见生成候选 Prompt，确认后才会保存为新版本。"
                        parent.context.update(ctx)
                        parent.prompt_panel.set_context(parent.context)
                        parent.prompt_panel.prompt_edit.setPlainText(ctx["prompt_template"])
                        parent.prompt_panel.current_prompt = ctx["prompt_template"]
                        parent.prompt_panel._update_placeholders()
                        parent.prompt_panel.status_label.setText(
                            f"✓ 已根据修改意见生成候选 Prompt。请审查，确认后才会保存为 v{latest_v + 1}。"
                        )
                        parent.prompt_panel.status_label.setStyleSheet("color: green; font-weight: bold;")
                        parent.prompt_panel.confirm_btn.setEnabled(True)
                        parent.prompt_panel.edit_btn.setEnabled(True)
                        parent.prompt_panel.regenerate_btn.setEnabled(True)
                        _notify_long_task_done(
                            self, "候选 Prompt 生成完成",
                            "候选 Prompt 已生成，请回到 Step 3 审查。"
                        )

                    def _on_continue_error(error_msg):
                        parent.prompt_panel.prompt_edit.setPlainText(prompt_template)
                        parent.prompt_panel.current_prompt = prompt_template
                        parent.prompt_panel._update_placeholders()
                        parent.prompt_panel.status_label.setText(
                            f"自动优化失败: {error_msg}。已加载原 Prompt，请手动编辑。"
                        )
                        parent.prompt_panel.status_label.setStyleSheet("color: red;")
                        parent.prompt_panel.confirm_btn.setEnabled(True)
                        parent.prompt_panel.edit_btn.setEnabled(True)
                        parent.prompt_panel.regenerate_btn.setEnabled(True)

                    self.continue_opt_worker.finished.connect(_on_continue_optimized)
                    self.continue_opt_worker.error.connect(_on_continue_error)
                    self.continue_opt_worker.progress.connect(parent.prompt_panel.status_label.setText)
                    self.continue_opt_worker.start()
                    return

        # go_to_step mapping: 1=model, 2=requirement, 3=prompt, 4=eval_data,
        #                     5=eval_script, 6=eval_run, 7=results, 8=optimize
        target_step = 3  # 统一：续接一律到 Step 3

        # Directly update context and jump
        parent.context.update(ctx)

        # Set context for all intermediate panels based on target step
        if target_step >= 3:
            parent.prompt_panel.set_context(parent.context)
        if target_step >= 4:
            parent.eval_data_panel.set_context(parent.context)
            if not eval_data_loaded and saved_script:
                parent.eval_data_panel.status_label.setText(
                    "⚠️ 评测数据文件不存在（旧版未保存），可重新生成。")
                parent.eval_data_panel.status_label.setStyleSheet("color: orange;")
        if target_step >= 5:
            parent.eval_script_panel.set_context(parent.context)
            # Auto-load saved eval script into editor
            if saved_script:
                parent.eval_script_panel.script_edit.setPlainText(saved_script)
                parent.eval_script_panel.script_edit.setReadOnly(False)
                parent.eval_script_panel.save_btn.setEnabled(True)
                parent.eval_script_panel.confirm_btn.setEnabled(True)
                parent.eval_script_panel.status_label.setText("✓ 已加载保存的评测脚本。")
                parent.eval_script_panel.status_label.setStyleSheet("color: green;")
        if target_step >= 6:
            parent.eval_run_panel.set_context(parent.context)
        if target_step >= 7:
            parent.results_panel.set_context(parent.context)

        parent.go_to_step(target_step)

    def _setup_ui(self):
        # Wrap in QScrollArea so vertical squeeze doesn't hide content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setStyleSheet("QScrollArea { border: none; }")
        scroll_widget = QWidget()
        layout = QVBoxLayout(scroll_widget)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 12, 20, 20)

        title = QLabel("Step 2: 需求输入")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        desc = QLabel("先在下框手动输入你的核心需求，再通过下方 PRD 文档导入由 LLM 提取补充需求。两种内容会自动合并。")
        desc.setWordWrap(True)
        desc.setStyleSheet("font-size: 14px; line-height: 1.5; padding: 2px 0 6px 0;")
        layout.addWidget(desc)

        # ─── Scenario name ───
        name_label = QLabel("场景名称（正式场景用业务名；临时测试请用 测试_具体用途）:")
        name_label.setStyleSheet("font-size: 14px; padding: 2px 0;")
        layout.addWidget(name_label)
        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText("例如: UserBasicInfoExtract / 测试_下一组交互式会话")
        self.name_edit.setStyleSheet("QLineEdit { font-size: 14px; padding: 6px 8px;}")
        layout.addWidget(self.name_edit)

        # ─── Requirements (manual input, shorter) ───
        req_label = QLabel("✍ 手动输入核心需求（对 Prompt 的要求和设想 — 必填）:")
        req_label.setStyleSheet("font-size: 14px; font-weight: bold; padding: 4px 0;")
        layout.addWidget(req_label)
        self.req_edit = QTextEdit()
        self.req_edit.setPlaceholderText(
            "请详细描述:\n"
            "  - 这个 prompt 要完成什么任务？\n"
            "  - 输入有哪些字段？各是什么类型？\n"
            "  - 输出是什么格式（JSON字段）？\n"
            "  - 有哪些规则或约束？\n"
            "  - 输出要体现什么风格和内容？"
        )
        self.req_edit.setMinimumHeight(80)
        self.req_edit.setMaximumHeight(120)
        self.req_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        _enable_width_wrap(self.req_edit)
        self.req_edit.setStyleSheet(
            "QTextEdit { font-size: 14px; line-height: 1.5; padding: 8px; }"
        )
        layout.addWidget(self.req_edit)

        # ─── Document import panel (always visible) ───
        self.docs_panel = QGroupBox("📂 导入 PRD 文档 — 自动提取需求（可选）")
        self.docs_panel.setStyleSheet("QGroupBox { font-size: 14px; }")
        docs_layout = QVBoxLayout(self.docs_panel)
        docs_layout.setContentsMargins(10, 16, 10, 10)

        docs_layout.addWidget(QLabel("从 data/raw_docs/ 选择场景文件夹 (LLM 自动提取结构化需求):"))

        # Horizontal split: folder list (left) + file display (right)
        folder_file_layout = QHBoxLayout()
        folder_file_layout.setSpacing(12)

        # Left: Folder list
        left_col = QVBoxLayout()
        left_col.addWidget(QLabel("场景文件夹:"))
        self.folder_list = QListWidget()
        self.folder_list.setMinimumHeight(100)
        self.folder_list.setMaximumHeight(200)
        self.folder_list.setStyleSheet(
            "QListWidget { font-size: 14px; }"
            "QListWidget::item { padding: 6px 8px; }"
        )
        self.folder_list.itemClicked.connect(self._on_folder_selected)
        left_col.addWidget(self.folder_list)
        folder_file_layout.addLayout(left_col, 1)

        # Right: File display
        right_col = QVBoxLayout()
        file_label = QLabel("📄 包含文件:")
        file_label.setStyleSheet("font-weight: bold; color: #555;")
        right_col.addWidget(file_label)
        self.files_display = QTextEdit()
        self.files_display.setReadOnly(True)
        self.files_display.setMinimumHeight(100)
        self.files_display.setMaximumHeight(200)
        self.files_display.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        _enable_width_wrap(self.files_display)
        self.files_display.setStyleSheet(
            "QTextEdit { font-size: 14px; padding: 8px; background: #f9f9f9;"
            " border: 1px solid #eee; border-radius: 3px; color: #333; }"
        )
        right_col.addWidget(self.files_display)
        folder_file_layout.addLayout(right_col, 1)

        docs_layout.addLayout(folder_file_layout)

        # Buttons
        doc_btn_layout = QHBoxLayout()
        self.upload_btn = QPushButton("📤 上传文档")
        self.upload_btn.setToolTip("从本地选择文件上传到 data/raw_docs/，上传时会提示输入文件夹名称")
        self.upload_btn.clicked.connect(self._upload_docs)
        doc_btn_layout.addWidget(self.upload_btn)

        self.refresh_docs_btn = QPushButton("刷新目录")
        self.refresh_docs_btn.clicked.connect(self._refresh_doc_folders)
        doc_btn_layout.addWidget(self.refresh_docs_btn)

        self.analyze_btn = QPushButton("LLM 自动分析需求")
        self.analyze_btn.setEnabled(False)
        self.analyze_btn.setStyleSheet(
            "QPushButton { background: #F57C00; color: white; padding: 8px 20px;"
            " border-radius: 4px; }"
            "QPushButton:hover { background: #E65100; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.analyze_btn.clicked.connect(self._llm_analyze_docs)
        doc_btn_layout.addWidget(self.analyze_btn)

        self.analyze_status = QLabel("")
        self.analyze_status.setStyleSheet("color: #666;")
        doc_btn_layout.addWidget(self.analyze_status)
        doc_btn_layout.addStretch()
        docs_layout.addLayout(doc_btn_layout)

        layout.addWidget(self.docs_panel)

        # ─── Global Context section ──────────────────────────────────────
        self.context_group = QGroupBox("全局字段说明 (Context 文档)")
        self.context_group.setToolTip("所有场景共享的字段定义文档，可上传多份（如一份带格式定义、一份字段说明），自动合并使用")
        self.context_group.setStyleSheet("QGroupBox { font-size: 14px; }")
        ctx_layout = QVBoxLayout(self.context_group)
        ctx_layout.setContentsMargins(10, 16, 10, 10)

        ctx_top = QHBoxLayout()
        self.context_status_label = QLabel("")
        self.context_status_label.setStyleSheet("font-size: 14px; padding: 4px 0;")
        ctx_top.addWidget(self.context_status_label)
        ctx_top.addStretch()

        self.upload_context_btn = QPushButton("📤 上传 Context（追加）")
        self.upload_context_btn.clicked.connect(self._upload_context)
        ctx_top.addWidget(self.upload_context_btn)

        self.delete_context_btn = QPushButton("🗑 删除")
        self.delete_context_btn.setStyleSheet("color: #C62828; font-size: 13px;")
        self.delete_context_btn.clicked.connect(self._delete_context)
        ctx_top.addWidget(self.delete_context_btn)

        ctx_layout.addLayout(ctx_top)

        field_toolbar = QHBoxLayout()
        field_toolbar.addWidget(QLabel("本场景使用字段（勾选后，后续生成和调优只能使用这些字段）:"))
        field_toolbar.addStretch()
        self.select_all_context_fields_btn = QPushButton("全选")
        self.select_all_context_fields_btn.clicked.connect(lambda: self._set_all_context_fields_checked(True))
        field_toolbar.addWidget(self.select_all_context_fields_btn)
        self.clear_context_fields_btn = QPushButton("清空")
        self.clear_context_fields_btn.clicked.connect(lambda: self._set_all_context_fields_checked(False))
        field_toolbar.addWidget(self.clear_context_fields_btn)
        ctx_layout.addLayout(field_toolbar)

        self.context_field_list = QTreeWidget()
        self.context_field_list.setMinimumHeight(220)
        self.context_field_list.setMaximumHeight(420)
        self.context_field_list.setColumnCount(4)
        self.context_field_list.setHeaderLabels(["字段 / 嵌套层级", "命名域", "类型", "上线占位符"])
        self.context_field_list.header().setSectionResizeMode(0, QHeaderView.Stretch)
        self.context_field_list.header().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.context_field_list.header().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.context_field_list.header().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.context_field_list.setSelectionMode(QAbstractItemView.NoSelection)
        self.context_field_list.setAlternatingRowColors(True)
        self.context_field_list.setUniformRowHeights(True)
        self.context_field_list.setStyleSheet(
            "QTreeWidget { font-size: 13px; }"
            "QTreeWidget::item { padding: 3px 6px; }"
        )
        self.context_field_list.itemChanged.connect(self._on_context_field_changed)
        ctx_layout.addWidget(self.context_field_list)

        self.context_field_hint = QLabel("")
        self.context_field_hint.setWordWrap(True)
        self.context_field_hint.setStyleSheet("color: #666; font-size: 12px;")
        ctx_layout.addWidget(self.context_field_hint)

        layout.addWidget(self.context_group)

        self._refresh_context_status()

        # ─── Existing scenarios（续接功能）───────────────────────────────
        self.continue_group = QGroupBox("已有场景（继续之前的工作）")
        self.continue_group.setStyleSheet("QGroupBox { font-size: 14px; }")
        cont_layout = QVBoxLayout(self.continue_group)
        cont_layout.setContentsMargins(10, 16, 10, 10)

        cont_top = QHBoxLayout()
        cont_top.addWidget(QLabel("以下场景已有保存的数据，点击选择后可继续："))
        self.refresh_continue_btn = QPushButton("刷新列表")
        self.refresh_continue_btn.clicked.connect(self._refresh_existing_scenarios)
        cont_top.addWidget(self.refresh_continue_btn)
        cont_layout.addLayout(cont_top)

        self.continue_list = QListWidget()
        self.continue_list.setMinimumHeight(60)
        self.continue_list.setMaximumHeight(150)
        self.continue_list.setStyleSheet(
            "QListWidget { font-size: 14px; }"
            "QListWidget::item { padding: 6px 8px; }"
        )
        self.continue_list.itemClicked.connect(self._on_continue_scenario_selected)
        cont_layout.addWidget(self.continue_list)

        self.continue_status = QTextEdit()
        self.continue_status.setReadOnly(True)
        self.continue_status.setMaximumHeight(80)
        self.continue_status.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        _enable_width_wrap(self.continue_status)
        self.continue_status.setStyleSheet(
            "QTextEdit { font-size: 14px; padding: 8px; background: #e3f2fd;"
            " border: 1px solid #bbdefb; border-radius: 4px; color: #1565C0; }"
        )
        cont_layout.addWidget(self.continue_status)

        self.continue_revision_label = QLabel("本次继续的修改需求（可选；重新勾选字段后也在这里写）:")
        self.continue_revision_label.setStyleSheet("font-weight: bold; color: #333;")
        cont_layout.addWidget(self.continue_revision_label)

        self.continue_revision_edit = QTextEdit()
        self.continue_revision_edit.setMinimumHeight(70)
        self.continue_revision_edit.setMaximumHeight(130)
        self.continue_revision_edit.setPlaceholderText(
            "例如：只使用 training_session.workout_vector 和 target_language；输出更短；不要生成未勾选字段。"
        )
        self.continue_revision_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        _enable_width_wrap(self.continue_revision_edit)
        self.continue_revision_edit.setStyleSheet(
            "QTextEdit { font-size: 14px; padding: 8px; background: #fffdf5;"
            " border: 1px solid #ffe082; border-radius: 3px; color: #333; }"
        )
        cont_layout.addWidget(self.continue_revision_edit)

        cont_btn_layout = QHBoxLayout()
        self.continue_btn = QPushButton("▶ 继续此场景")
        self.continue_btn.setEnabled(False)
        self.continue_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 24px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #1565C0; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.continue_btn.clicked.connect(self._do_continue)
        cont_btn_layout.addWidget(self.continue_btn)
        cont_btn_layout.addStretch()
        cont_layout.addLayout(cont_btn_layout)

        layout.addWidget(self.continue_group)
        self._refresh_existing_scenarios()

        # ─── LLM extracted requirements (editable, shown after analysis) ───
        self.llm_result_group = QGroupBox("LLM 综合最终需求（手动需求 + PRD 文档合并结果 — 确认时优先采用此版本）")
        self.llm_result_group.setStyleSheet("QGroupBox { font-size: 14px; }")
        llm_layout = QVBoxLayout(self.llm_result_group)
        llm_layout.setContentsMargins(10, 16, 10, 10)

        self.llm_result_edit = QTextEdit()
        self.llm_result_edit.setReadOnly(False)
        self.llm_result_edit.setMinimumHeight(250)
        self.llm_result_edit.setMaximumHeight(500)
        self.llm_result_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        _enable_width_wrap(self.llm_result_edit)
        self.llm_result_edit.setStyleSheet(
            "QTextEdit { font-size: 14px; padding: 8px; background: #fff8e1;"
            " border: 1px solid #ffe082; border-radius: 3px; color: #333; }"
        )
        self.llm_result_edit.setPlaceholderText("点击上方「LLM 自动分析需求」后，结果将显示在此处；你可以直接手动修改后再确认。")
        llm_layout.addWidget(self.llm_result_edit)

        self.clear_llm_btn = QPushButton("🗑 清除 LLM 结果")
        self.clear_llm_btn.setStyleSheet("color: #C62828; font-size: 13px;")
        self.clear_llm_btn.setEnabled(False)
        self.clear_llm_btn.clicked.connect(self._clear_llm_result)
        llm_layout.addWidget(self.clear_llm_btn)

        self.llm_result_group.setVisible(False)
        layout.addWidget(self.llm_result_group)

        # ─── Buttons ───
        btn_layout = QHBoxLayout()
        self.confirm_btn = QPushButton("✓ 完成输入 (确认需求)")
        self.confirm_btn.setStyleSheet(
            "QPushButton { background: #388E3C; color: white; padding: 8px 24px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #2E7D32; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.confirm_btn.clicked.connect(self._confirm)
        btn_layout.addWidget(self.confirm_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        self.status_label = QLabel("")
        self.status_label.setStyleSheet("font-size: 14px; padding: 4px 0;")
        layout.addWidget(self.status_label)

        layout.addStretch()

        # Finalize scroll area
        scroll.setWidget(scroll_widget)
        # Replace this panel's layout with one containing the scroll area
        outer_layout = QVBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.addWidget(scroll)

    def _confirm(self):
        name = normalize_scenario_name(self.name_edit.text())
        manual_text = self.req_edit.toPlainText().strip()

        if not name:
            QMessageBox.warning(self, "警告", "请输入场景名称。")
            return
        resolved_name = _resolve_scratch_scenario_name(self, name)
        if not resolved_name:
            return
        if resolved_name != name:
            name = resolved_name
            self.name_edit.setText(name)
        if not manual_text:
            current_continue_item = self.continue_list.currentItem() if hasattr(self, "continue_list") else None
            if current_continue_item and current_continue_item.data(Qt.UserRole):
                QMessageBox.warning(
                    self, "提示",
                    "你当前选中的是已有场景。请在「本次继续的修改需求」里填写修改点，然后点击蓝色「继续此场景」。\n\n"
                    "绿色「完成输入」用于新建场景。"
                )
                return
            QMessageBox.warning(self, "警告", "请输入需求描述。")
            return

        # Use LLM-merged result (bottom) if available, otherwise manual input (top)
        llm_text = self.llm_result_edit.toPlainText().strip()
        req = llm_text if llm_text else manual_text

        # Append global context doc to requirements if it exists
        ctx_content = self._read_context_content()
        if ctx_content:
            req += (
                "\n\n# 全局字段定义\n"
                + self._context_selection_note()
                + "\n以下字段定义文档仅供理解字段含义；当前场景只能使用上方已勾选字段：\n"
                + ctx_content
            )

        self.scenario_name = name
        self.requirements = req

        # Silently strip field references not found in context document
        ctx_fields = self._get_active_context_fields()
        if ctx_fields:
            refs = self._extract_field_refs_from_text(req)
            allowed_ref_names = set(ctx_fields) | {field.split(".")[-1] for field in ctx_fields}
            unknown = refs - allowed_ref_names
            for raw_f in unknown:
                req = re.sub(r'\{' + re.escape(raw_f) + r'\}', '', req)
                req = re.sub(r'\b' + re.escape(raw_f) + r'\b', '', req)
            req = re.sub(r'\n{3,}', '\n\n', req)
            req = req.strip()
            self.requirements = req
            if unknown and not req:
                QMessageBox.warning(self, "警告", "字段移除后需求内容为空，请补充手动输入的需求描述。")
                return

        self.status_label.setText(f"✓ 场景: {name} | 需求已记录")
        self.status_label.setStyleSheet("color: green; font-weight: bold;")

        # Check for existing versions
        vm = VersionManager(name)
        existing = vm.list_versions()
        if existing:
            reply = QMessageBox.question(
                self, "发现已有版本",
                f"场景 '{name}' 已有版本: v{', v'.join(map(str, existing))}\n\n"
                f"「是」= 加载最新版本继续优化  |  「否」= 从头开始",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                latest_v = existing[-1]
                try:
                    prompt_template = vm.load_version(latest_v)
                    eval_data = []
                    try:
                        eval_data = vm.load_eval_data()
                    except FileNotFoundError:
                        pass

                    parent = self.parent()
                    while parent and not hasattr(parent, 'on_step_complete'):
                        parent = parent.parent()
                    if parent:
                        parent.on_step_complete(1, {
                            "scenario_name": name,
                            "requirements": req,
                            "existing_version": latest_v,
                            "prompt_template": prompt_template,
                            "eval_data": eval_data,
                            "vm": vm,
                            "continue_mode": True,
                            "all_context_fields": self._get_context_fields(),
                            "context_fields": self._get_active_context_fields(),
                            "selected_context_fields": self._get_active_context_fields(),
                            "context_field_specs": self._get_active_context_field_specs(),
                        })
                    return
                except Exception as e:
                    QMessageBox.warning(self, "加载失败", f"加载已有版本失败: {e}")
            else:
                vm = VersionManager(name)
                parent = self.parent()
                while parent and not hasattr(parent, 'on_step_complete'):
                    parent = parent.parent()
                if parent:
                    parent.on_step_complete(1, {
                        "scenario_name": name,
                        "requirements": req,
                        "vm": vm,
                        "continue_mode": False,
                        "all_context_fields": self._get_context_fields(),
                        "context_fields": self._get_active_context_fields(),
                        "selected_context_fields": self._get_active_context_fields(),
                        "context_field_specs": self._get_active_context_field_specs(),
                    })
                return
        vm = VersionManager(name)
        parent = self.parent()
        while parent and not hasattr(parent, 'on_step_complete'):
            parent = parent.parent()
        if parent:
            parent.on_step_complete(1, {
                "scenario_name": name,
                "requirements": req,
                "vm": vm,
                "continue_mode": False,
                "all_context_fields": self._get_context_fields(),
                "context_fields": self._get_active_context_fields(),
                "selected_context_fields": self._get_active_context_fields(),
                "context_field_specs": self._get_active_context_field_specs(),
            })

    # ─── V2.1: Document import mode methods ───────────────────────────────

    def _get_raw_docs_dir(self):
        """Get the data/raw_docs/ directory path."""
        return os.path.join(PROJECT_ROOT, "data", "raw_docs")

    def _read_doc_file(self, file_path: str) -> str:
        """Read a document file and return its text content."""
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".docx":
                doc = docx.Document(file_path)
                return "\n".join(p.text for p in doc.paragraphs)
            elif ext == ".json":
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                return json.dumps(data, ensure_ascii=False, indent=2)
            elif ext in (".txt", ".md", ".py", ".yaml", ".yml", ".csv"):
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            else:
                return f"[不支持的文件格式: {ext}]"
        except Exception as e:
            return f"[读取失败: {e}]"

    def _refresh_doc_folders(self):
        """Scan data/raw_docs/ for scenario folders."""
        docs_dir = self._get_raw_docs_dir()
        os.makedirs(docs_dir, exist_ok=True)

        self.folder_list.clear()
        self.files_display.clear()
        self.analyze_btn.setEnabled(False)
        self.analyze_status.setText("")

        found = False
        for item in sorted(os.listdir(docs_dir)):
            item_path = os.path.join(docs_dir, item)
            if os.path.isdir(item_path) and not item.startswith("."):
                # Count readable files
                files = [f for f in os.listdir(item_path)
                         if os.path.isfile(os.path.join(item_path, f))
                         and not f.startswith(".")]
                file_exts = ", ".join(os.path.splitext(f)[1] for f in files[:6])
                label = f"{item}  ({len(files)} 个文件)"
                list_item = QListWidgetItem(label)
                list_item.setData(Qt.UserRole, item)
                self.folder_list.addItem(list_item)
                found = True

        if not found:
            no_folder_item = QListWidgetItem("(暂无场景文件夹，请在 data/raw_docs/ 下创建)")
            no_folder_item.setFlags(no_folder_item.flags() & ~Qt.ItemIsSelectable)
            self.folder_list.addItem(no_folder_item)

    def _upload_docs(self):
        """Upload files from local machine to data/raw_docs/{folder_name}/."""
        from PyQt5.QtWidgets import QInputDialog

        # 1. Select files
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择要上传的文档",
            "",
            "文档文件 (*.docx *.json *.txt *.md *.pdf *.csv *.yaml *.yml);;所有文件 (*)"
        )
        if not files:
            return

        # 2. Ask for folder name
        folder_name, ok = QInputDialog.getText(
            self, "新建场景文件夹",
            "请输入场景文件夹名称（英文，如 MyScenario）:",
            QLineEdit.Normal, ""
        )
        if not ok or not folder_name.strip():
            return
        folder_name = folder_name.strip()

        # 3. Create folder and copy files
        target_dir = os.path.join(self._get_raw_docs_dir(), folder_name)
        os.makedirs(target_dir, exist_ok=True)

        copied = []
        skipped = []
        for src_path in files:
            name = os.path.basename(src_path)
            dst_path = os.path.join(target_dir, name)

            # Avoid overwrite
            if os.path.exists(dst_path):
                base, ext = os.path.splitext(name)
                counter = 1
                while os.path.exists(dst_path):
                    dst_path = os.path.join(target_dir, f"{base}_{counter}{ext}")
                    counter += 1

            try:
                with open(src_path, "rb") as sf:
                    with open(dst_path, "wb") as df:
                        df.write(sf.read())
                copied.append(os.path.basename(dst_path))
            except Exception as e:
                skipped.append(f"{name} ({e})")

        # 4. Refresh folder list and select the new folder
        self._refresh_doc_folders()
        for i in range(self.folder_list.count()):
            item = self.folder_list.item(i)
            if item.data(Qt.UserRole) == folder_name:
                self.folder_list.setCurrentItem(item)
                self._on_folder_selected(item)
                break

        # 5. Show result
        msg = f"✅ 成功上传 {len(copied)} 个文件到 data/raw_docs/{folder_name}/"
        if skipped:
            msg += f"\n⚠️ 跳过 {len(skipped)} 个: {', '.join(skipped)}"
        QMessageBox.information(self, "上传完成", msg)

    def _on_folder_selected(self, item):
        """When a folder is selected, show its file list."""
        folder_name = item.data(Qt.UserRole)
        if not folder_name:
            return

        folder_path = os.path.join(self._get_raw_docs_dir(), folder_name)
        if not os.path.isdir(folder_path):
            return

        # Auto-fill scenario name
        self.name_edit.setText(folder_name)

        # List files
        files = []
        for f in sorted(os.listdir(folder_path)):
            fpath = os.path.join(folder_path, f)
            if os.path.isfile(fpath) and not f.startswith("."):
                size = os.path.getsize(fpath)
                files.append(f"  {f} ({size:,} bytes)")

        if files:
            self.files_display.setPlainText("\n".join(files))
            self.analyze_btn.setEnabled(True)
            self.analyze_status.setText("点击「LLM 自动分析需求」提取结构化需求")
            self.analyze_status.setStyleSheet("color: #F57C00;")
        else:
            self.files_display.setPlainText("（该文件夹为空或无支持的文档格式）")
            self.analyze_btn.setEnabled(False)

    def _llm_analyze_docs(self):
        """Call Helper LLM to analyze documents and extract requirements."""
        if not self.helper_llm:
            QMessageBox.warning(self, "警告",
                                "Helper 模型未加载。请先在 Step 1 中选择并确认模型。")
            return

        folder_name = self.name_edit.text().strip()
        if not folder_name:
            return

        folder_path = os.path.join(self._get_raw_docs_dir(), folder_name)
        if not os.path.isdir(folder_path):
            return

        # Read all documents
        self.analyze_btn.setEnabled(False)
        self.analyze_status.setText("正在读取文档...")
        self.analyze_status.setStyleSheet("color: #666;")
        self.llm_result_edit.setPlainText("正在读取并分析文档，请稍候...")
        self.llm_result_group.setVisible(True)
        QApplication.processEvents()

        docs_parts = []

        # Prepend global context document if it exists
        ctx_content = self._read_context_content()
        if ctx_content:
            docs_parts.insert(
                0,
                "===== [全局 Context] 字段定义文档 =====\n"
                + self._context_selection_note()
                + "\n以下字段定义文档仅供理解字段含义；当前场景只能使用上方已勾选字段：\n"
                + ctx_content
                + "\n"
            )

        # Read scenario-specific documents
        supported = (".docx", ".json", ".txt", ".md", ".py", ".yaml", ".yml", ".csv")
        for f in sorted(os.listdir(folder_path)):
            fpath = os.path.join(folder_path, f)
            if os.path.isfile(fpath) and os.path.splitext(f)[1].lower() in supported:
                content = self._read_doc_file(fpath)
                docs_parts.append(f"===== {f} =====\n{content}\n")

        if not docs_parts:
            self.llm_result_edit.setPlainText("（文件夹中无可支持的文档文件）")
            self.llm_result_group.setVisible(True)
            QMessageBox.warning(self, "警告", "文件夹中无可支持的文档文件。")
            self.analyze_btn.setEnabled(True)
            return

        all_docs = "\n".join(docs_parts)
        self.analyze_status.setText(f"已读取 {len(docs_parts)} 个文件，正在调用 LLM 分析...")
        self.analyze_status.setStyleSheet("color: #666;")

        # Read manual input to merge with PRD docs
        manual_input = self.req_edit.toPlainText().strip()

        # Start background worker
        self.doc_worker = RawDocsAnalyzeWorker(
            self.helper_llm, folder_name, all_docs, manual_input
        )
        self.doc_worker.finished.connect(self._on_docs_analyzed)
        self.doc_worker.error.connect(self._on_docs_analyze_error)
        self.doc_worker.progress.connect(self.analyze_status.setText)
        self.doc_worker.start()

    def _on_docs_analyzed(self, requirements):
        """Handle completed LLM analysis — show merged result in bottom editable area."""
        self.llm_requirements = requirements
        self.llm_result_edit.setPlainText(requirements)
        self.llm_result_edit.setReadOnly(False)
        self.llm_result_group.setVisible(True)
        self.clear_llm_btn.setEnabled(True)
        self.status_label.setText("✓ LLM 已合并手动需求与 PRD 文档；下方最终需求可手动修改，确认时将使用修改后的版本")
        self.status_label.setStyleSheet("color: green; font-weight: bold;")
        self.analyze_status.setText("✓ 分析完成")
        self.analyze_status.setStyleSheet("color: green;")
        self.analyze_btn.setEnabled(True)
        _notify_long_task_done(
            self, "需求分析完成",
            "LLM 综合最终需求已生成，可以在 Step 2 下方手动修改后确认。"
        )

    def _on_docs_analyze_error(self, error_msg):
        """Handle LLM analysis error."""
        self.llm_result_edit.setPlainText("")
        QMessageBox.critical(self, "分析失败", f"文档分析失败: {error_msg}")
        self.analyze_status.setText(f"分析失败: {error_msg}")
        self.analyze_status.setStyleSheet("color: red;")
        self.analyze_btn.setEnabled(True)

    def _clear_llm_result(self):
        """Clear the LLM analysis result."""
        self.llm_requirements = ""
        self.llm_result_edit.clear()
        self.llm_result_group.setVisible(False)
        self.clear_llm_btn.setEnabled(False)
        self.status_label.setText("已清除 LLM 提取的需求")

    # ─── Global Context methods ─────────────────────────────────────────

    def _read_context_content(self) -> str:
        """Read all global context files and concatenate their content."""
        ctx_files = self._get_context_files()
        if not ctx_files:
            return ""

        parts = []
        for ctx_path in ctx_files:
            fname = os.path.basename(ctx_path)
            try:
                ext = os.path.splitext(ctx_path)[1].lower()
                content = ""
                if ext in (".txt", ".md", ".py", ".yaml", ".yml", ".csv"):
                    with open(ctx_path, "r", encoding="utf-8") as f:
                        content = f.read()
                elif ext == ".json":
                    with open(ctx_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    content = json.dumps(data, ensure_ascii=False, indent=2)
                elif ext == ".docx":
                    doc = docx.Document(ctx_path)
                    lines = [p.text for p in doc.paragraphs]
                    for table in doc.tables:
                        for row in table.rows:
                            cells = " | ".join(cell.text.strip() for cell in row.cells)
                            if cells.strip():
                                lines.append(cells)
                    content = "\n".join(lines)
                else:
                    content = f"[不支持的 context 文件格式: {ext}]"
                if content.strip():
                    parts.append(f"===== {fname} =====\n{content}")
            except Exception as e:
                parts.append(f"===== {fname} =====\n[读取失败: {e}]")

        return "\n\n".join(parts)

    def _refresh_context_status(self):
        """Update the context status label."""
        ctx_files = self._get_context_files()
        if ctx_files:
            from datetime import datetime
            lines = [f"✅ 已上传 {len(ctx_files)} 个 Context 文档（全局共享）:"]
            for fp in ctx_files:
                fname = os.path.basename(fp)
                mtime = datetime.fromtimestamp(os.path.getmtime(fp)).strftime("%Y-%m-%d %H:%M")
                size = os.path.getsize(fp)
                lines.append(f"  · {fname}  ({size:,} bytes, {mtime})")
            self.context_status_label.setText("\n".join(lines))
            self.context_status_label.setStyleSheet("color: #2E7D32; font-size: 14px; padding: 4px 0;")
            self.delete_context_btn.setEnabled(True)
        else:
            self.context_status_label.setText("⚠️ 未上传 — 请上传字段定义文档，所有场景全局共享")
            self.context_status_label.setStyleSheet("color: #F57C00; font-size: 14px; padding: 4px 0;")
            self.delete_context_btn.setEnabled(False)
        self._refresh_context_field_checklist()

    def showEvent(self, event):
        super().showEvent(event)
        if hasattr(self, "context_field_list"):
            self._refresh_context_status()

    def _upload_context(self):
        """Upload additional context document (appends, does not replace)."""
        path, _ = QFileDialog.getOpenFileName(
            self, "上传 Context 文档（追加，不会覆盖已有文件）",
            "", "文档文件 (*.docx *.json *.txt *.md);;所有文件 (*)"
        )
        if not path:
            return

        os.makedirs(self.CONTEXT_DIR, exist_ok=True)

        # Copy new file (if name collides, append number suffix)
        base_name = os.path.basename(path)
        dst = os.path.join(self.CONTEXT_DIR, base_name)
        if os.path.exists(dst):
            base, ext = os.path.splitext(base_name)
            counter = 1
            while os.path.exists(dst):
                dst = os.path.join(self.CONTEXT_DIR, f"{base}_{counter}{ext}")
                counter += 1

        try:
            with open(path, "rb") as sf:
                with open(dst, "wb") as df:
                    df.write(sf.read())
            self._refresh_context_status()
            QMessageBox.information(self, "上传成功",
                                    f"Context 文档已追加保存:\n{dst}\n\n"
                                    "请在字段列表中勾选当前场景允许使用的字段。")
        except Exception as e:
            QMessageBox.critical(self, "上传失败", str(e))

    def _delete_context(self):
        """Delete all global context files."""
        ctx_files = self._get_context_files()
        if not ctx_files:
            return
        names = "\n".join(f"  · {os.path.basename(f)}" for f in ctx_files)
        reply = QMessageBox.question(
            self, "确认删除",
            f"确定要删除以下全局 Context 文档？\n{names}",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            for fp in ctx_files:
                os.remove(fp)
            self.selected_context_fields = set()
            self.context_selection_initialized = False
            self._refresh_context_status()
            QMessageBox.information(self, "已删除",
                                    f"已删除 {len(ctx_files)} 个 Context 文档。")

    def _refresh_context_field_checklist(self):
        """Render extracted Context fields as a checkable per-scenario whitelist."""
        if not hasattr(self, "context_field_list"):
            return
        specs = self._sort_context_field_specs(self._get_context_field_specs())
        all_fields = {spec["name"] for spec in specs}
        if self.context_selection_initialized:
            selected = set(self.selected_context_fields) & all_fields
        else:
            selected = set()
            self.context_selection_initialized = True
        self.selected_context_fields = selected
        self.context_field_specs = specs

        self.context_field_list.blockSignals(True)
        self.context_field_list.clear()
        node_map = {}
        selectable_names = {spec["name"] for spec in specs}
        group_paths = set()
        for spec in specs:
            parts = (spec.get("name") or "").split(".")
            for idx in range(1, len(parts)):
                group_paths.add(".".join(parts[:idx]))

        def ensure_node(path_parts, namespace):
            parent_item = None
            current_path = ""
            for idx, part in enumerate(path_parts):
                current_path = f"{current_path}.{part}" if current_path else part
                if current_path in node_map:
                    parent_item = node_map[current_path]
                    continue
                item = QTreeWidgetItem([part, namespace, "namespace" if idx == 0 else "object", f"{{{{.{current_path}}}}}"])
                item.setData(0, Qt.UserRole, {
                    "name": current_path,
                    "path": current_path,
                    "namespace": namespace,
                    "is_group": True,
                })
                item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                item.setCheckState(0, Qt.Unchecked)
                if parent_item is None:
                    self.context_field_list.addTopLevelItem(item)
                else:
                    parent_item.addChild(item)
                node_map[current_path] = item
                parent_item = item
            return parent_item

        for spec in specs:
            name = spec["name"]
            path = spec.get("path") or name
            type_name = spec.get("type") or "unknown"
            parts = name.split(".")
            namespace = spec.get("namespace") or (parts[0] if parts else name)
            if name in group_paths:
                group_item = ensure_node(parts, namespace)
                if group_item is not None:
                    group_item.setText(2, type_name)
                    group_item.setToolTip(0, f"Context path: {path}")
                    group_item.setToolTip(3, f"Prompt / deploy placeholder: {{{{.{name}}}}}")
                continue
            parent_item = ensure_node(parts[:-1], namespace) if len(parts) > 1 else None
            leaf_label = parts[-1] if parts else name
            item = QTreeWidgetItem([
                leaf_label,
                namespace,
                type_name,
                f"{{{{.{name}}}}}",
            ])
            leaf_spec = dict(spec)
            leaf_spec["namespace"] = namespace
            leaf_spec["is_group"] = False
            item.setData(0, Qt.UserRole, leaf_spec)
            item.setToolTip(0, f"Context path: {path}")
            item.setToolTip(3, f"Prompt / deploy placeholder: {{{{.{name}}}}}")
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            item.setCheckState(0, Qt.Checked if name in selected else Qt.Unchecked)
            if parent_item is None:
                self.context_field_list.addTopLevelItem(item)
            else:
                parent_item.addChild(item)
            node_map[name] = item

        self._refresh_context_tree_parent_states(selectable_names)
        self.context_field_list.expandAll()
        for col in range(4):
            self.context_field_list.resizeColumnToContents(col)
        self.context_field_list.blockSignals(False)
        self._update_context_field_hint()

    def _get_checked_context_fields_from_ui(self) -> set:
        if not hasattr(self, "context_field_list"):
            return set()
        checked = set()
        selectable_names = self._get_context_fields()
        for item in self._iter_context_tree_items():
            if item.checkState(0) == Qt.Checked:
                spec = item.data(0, Qt.UserRole) or {}
                name = spec.get("name") or item.text(0).split()[0]
                if name in selectable_names:
                    checked.add(name)
        return checked

    def _iter_context_tree_items(self):
        if not hasattr(self, "context_field_list"):
            return
        stack = [self.context_field_list.topLevelItem(i) for i in range(self.context_field_list.topLevelItemCount())]
        while stack:
            item = stack.pop(0)
            if item is None:
                continue
            yield item
            for i in range(item.childCount()):
                stack.append(item.child(i))

    def _set_context_tree_children_checked(self, item, state):
        for i in range(item.childCount()):
            child = item.child(i)
            child.setCheckState(0, state)
            self._set_context_tree_children_checked(child, state)

    def _refresh_context_tree_parent_states(self, selectable_names=None):
        selectable_names = selectable_names or self._get_context_fields()

        def update_item(item):
            if item.childCount() == 0:
                return item.checkState(0)
            child_states = [update_item(item.child(i)) for i in range(item.childCount())]
            if all(state == Qt.Checked for state in child_states):
                state = Qt.Checked
            elif all(state == Qt.Unchecked for state in child_states):
                state = Qt.Unchecked
            else:
                state = Qt.PartiallyChecked
            item.setCheckState(0, state)
            return state

        for i in range(self.context_field_list.topLevelItemCount()):
            update_item(self.context_field_list.topLevelItem(i))

    def _get_active_context_fields(self) -> set:
        all_fields = self._get_context_fields()
        if hasattr(self, "context_field_list") and self.context_field_list.topLevelItemCount() > 0:
            checked = self._get_checked_context_fields_from_ui()
            self.selected_context_fields = checked & all_fields
            self.context_selection_initialized = True
            return set(self.selected_context_fields)
        if self.selected_context_fields:
            return set(self.selected_context_fields) & all_fields
        return set(all_fields)

    def _get_active_context_field_specs(self) -> list:
        selected = self._get_active_context_fields()
        return [spec for spec in self._sort_context_field_specs(self._get_context_field_specs()) if spec.get("name") in selected]

    def _sort_context_field_specs(self, specs: list) -> list:
        priority = {
            "coach_persona": 0,
            "target_language": 1,
            "trainingsession": 2,
            "training_session": 2,
            "intragroup": 3,
            "intra_group": 3,
        }

        def root_key(name: str) -> str:
            return (name or "").split(".")[0].lower()

        def sort_key(spec):
            name = spec.get("name", "")
            root = root_key(name)
            root_rank = priority.get(name.lower(), priority.get(root, 50))
            depth = name.count(".")
            return (root_rank, root, depth, name.lower())

        return sorted(specs or [], key=sort_key)

    def _set_all_context_fields_checked(self, checked: bool):
        if not hasattr(self, "context_field_list"):
            return
        self.context_field_list.blockSignals(True)
        state = Qt.Checked if checked else Qt.Unchecked
        for item in self._iter_context_tree_items():
            item.setCheckState(0, state)
        self.context_field_list.blockSignals(False)
        self.selected_context_fields = self._get_checked_context_fields_from_ui()
        self.context_selection_initialized = True
        self._update_context_field_hint()

    def _sync_context_fields_from_prompt(self, prompt_template: str) -> tuple:
        """Check Context fields used by an existing prompt and uncheck unrelated fields."""
        if not hasattr(self, "context_field_list"):
            return [], []
        placeholders = set(extract_placeholders(prompt_template or ""))
        all_fields = self._get_context_fields()
        matched = sorted(placeholders & all_fields)
        missing = sorted(placeholders - all_fields)

        effective = set(matched)
        for selected in matched:
            prefix = selected + "."
            effective.update(field for field in all_fields if field.startswith(prefix))

        self.context_field_list.blockSignals(True)
        for item in self._iter_context_tree_items():
            spec = item.data(0, Qt.UserRole) or {}
            name = spec.get("name") or item.text(0).split()[0]
            item.setCheckState(0, Qt.Checked if name in effective else Qt.Unchecked)
        self._refresh_context_tree_parent_states(all_fields)
        self.context_field_list.blockSignals(False)

        self.selected_context_fields = self._get_checked_context_fields_from_ui() & all_fields
        self.context_selection_initialized = True
        self._update_context_field_hint()
        return matched, missing

    def _on_context_field_changed(self, item, _column=0):
        self.context_field_list.blockSignals(True)
        state = item.checkState(0)
        if state in (Qt.Checked, Qt.Unchecked):
            self._set_context_tree_children_checked(item, state)
        self._refresh_context_tree_parent_states()
        self.context_field_list.blockSignals(False)
        self.selected_context_fields = self._get_checked_context_fields_from_ui()
        self.context_selection_initialized = True
        self._update_context_field_hint()

    def _update_context_field_hint(self):
        all_count = len(self._get_context_fields())
        selected = self._get_checked_context_fields_from_ui()
        if all_count > 0:
            self.context_field_hint.setText(
                f"已选择 {len(selected)}/{all_count} 个可用字段。GLOBAL 是系统全局变量；其余第一层是 Context 命名域。"
                "展开后可按嵌套对象选择，勾选父节点会选中整支字段。"
            )
            return
        if all_count > 0:
            self.context_field_hint.setText(
                f"已选择 {len(selected)}/{all_count} 个可用字段。树的第一层是命名域，展开后可按嵌套对象选择；"
                "勾选父节点会选中整支字段，未勾选字段不会进入 Prompt 生成、导入校验或后续自动调优。"
            )
            return
        if all_count == 0:
            self.context_field_hint.setText("未识别到可选字段。")
            return
        self.context_field_hint.setText(
            f"已选择 {len(selected)}/{all_count} 个字段。未勾选字段不会进入 Prompt 生成、导入校验或后续自动调优。"
        )

    def _context_selection_note(self) -> str:
        selected_specs = self._get_active_context_field_specs()
        if not selected_specs:
            return "No Context fields are selected for this scenario. Prompt must not use any Context placeholders."
        lines = []
        for spec in selected_specs:
            name = spec.get("name", "")
            path = spec.get("path") or name
            type_name = spec.get("type") or "unknown"
            lines.append(f"- `{{{{.{name}}}}}` = Context path `{path}` ({type_name})")
        return (
            "Selected Context field whitelist. Prompt generation and optimization must use only these placeholders; "
            "unlisted fields must not be referenced, guessed, or restored:\n"
            + "\n".join(lines)
        )
        selected = sorted(self._get_active_context_fields())
        if not selected:
            return "当前场景未勾选任何 Context 字段：Prompt 不得使用任何 Context 占位符。"
        return (
            "当前场景已勾选允许使用的 Context 字段如下；Prompt 生成和调优必须严格限制在这些字段内，"
            "未列出的字段视为不可用，不得引用、不得作为占位符、不得恢复：\n"
            + ", ".join(f"`{name}`" for name in selected)
        )

    # ─── Context field verification ─────────────────────────────────────

    # Common generic JSON keys that appear in context documents but aren't
    # meaningful "business fields" that prompt templates should reference.
    _GENERIC_KEYS = {
        "name", "type", "types", "role", "status", "error", "count",
        "context", "value", "values", "key", "keys", "data", "info",
        "text", "label", "title", "desc", "description", "note", "notes",
        "flag", "tags", "meta", "id", "ids", "code", "score",
        "total", "avg", "min", "max", "sum", "timestamp",
    }
    _BUILTIN_GLOBAL_FIELDS = {
        "target_language",
        "coach_persona",
    }

    @staticmethod
    def _is_valid_field_name(word: str) -> bool:
        """Validate a potential snake_case field name (filter out fragments and generic keys)."""
        _fragments = {
            "ontext", "riteria", "eps", "ser_profile",
            "progression", "regression", "tatus", "fficer",
            "confirmed", "conunt",
        }
        if len(word) < 4:
            return False
        if word in _fragments:
            return False
        if word in RequirementPanel._GENERIC_KEYS:
            return False
        return True

    @staticmethod
    def _context_value_type(value) -> str:
        if value is None:
            return "null"
        if isinstance(value, bool):
            return "boolean"
        if isinstance(value, int):
            return "integer"
        if isinstance(value, float):
            return "number"
        if isinstance(value, str):
            return "string"
        if isinstance(value, list):
            return "array"
        if isinstance(value, dict):
            return "object"
        return type(value).__name__

    @staticmethod
    def _clean_json_like(text: str) -> str:
        text = RequirementPanel._strip_json_line_comments(text or "")
        text = re.sub(r"/\*.*?\*/", "", text, flags=re.S)
        text = re.sub(r",\s*([}\]])", r"\1", text)
        text = re.sub(
            r'([}\]"0-9eE])\s*\n\s*("?[A-Za-z_][A-Za-z0-9_]*"?\s*:)',
            r"\1,\n\2",
            text,
        )
        return text.strip()

    @staticmethod
    def _strip_json_line_comments(text: str) -> str:
        out = []
        i = 0
        in_string = False
        escape = False
        while i < len(text):
            ch = text[i]
            nxt = text[i + 1] if i + 1 < len(text) else ""
            if in_string:
                out.append(ch)
                if escape:
                    escape = False
                elif ch == "\\":
                    escape = True
                elif ch == '"':
                    in_string = False
                i += 1
                continue
            if ch == '"':
                in_string = True
                out.append(ch)
                i += 1
                continue
            if ch == "/" and nxt == "/":
                while i < len(text) and text[i] not in "\r\n":
                    i += 1
                continue
            out.append(ch)
            i += 1
        return "".join(out)

    @staticmethod
    def _extract_json_like_blocks(text: str) -> list:
        blocks = re.findall(r"```(?:json|JSON)?\s*(.*?)```", text or "", re.S)
        src = text or ""
        start = None
        depth = 0
        in_string = False
        escape = False
        for idx, ch in enumerate(src):
            if in_string:
                if escape:
                    escape = False
                elif ch == "\\":
                    escape = True
                elif ch == '"':
                    in_string = False
                continue
            if ch == '"':
                in_string = True
                continue
            if ch == "{":
                if depth == 0:
                    start = idx
                depth += 1
            elif ch == "}" and depth:
                depth -= 1
                if depth == 0 and start is not None:
                    block = src[start:idx + 1]
                    if re.search(r'"[A-Za-z_][A-Za-z0-9_]*"\s*:', block):
                        blocks.append(block)
                    start = None
        return blocks

    @staticmethod
    def _is_valid_context_path(path: str) -> bool:
        parts = (path or "").split(".")
        if not parts or any(not re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", p) for p in parts):
            return False
        if len(parts) == 1:
            lowered = parts[0].lower()
            return RequirementPanel._is_valid_field_name(lowered) or lowered in RequirementPanel._BUILTIN_GLOBAL_FIELDS
        return True

    def _add_context_spec(self, specs: dict, name: str, type_name: str,
                          path: str = None, source: str = "", required: bool = True):
        path = path or name
        if not self._is_valid_context_path(name):
            return
        specs[name] = {
            "name": name,
            "path": path,
            "type": type_name or "string",
            "required": required,
            "source_excerpt": source or f"Context path: {path}",
        }

    def _walk_context_object(self, value, prefix: str, specs: dict):
        if prefix:
            self._add_context_spec(
                specs, prefix, self._context_value_type(value),
                path=prefix, source=f"Context path: {prefix}"
            )
        if isinstance(value, dict):
            for key, child in value.items():
                if not isinstance(key, str) or not re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", key):
                    continue
                child_path = f"{prefix}.{key}" if prefix else key
                self._walk_context_object(child, child_path, specs)
        elif isinstance(value, list) and value and isinstance(value[0], dict) and prefix:
            spec = specs.get(prefix)
            if spec:
                child_keys = sorted({k for item in value[:5] if isinstance(item, dict) for k in item.keys() if isinstance(k, str)})
                if child_keys:
                    spec["source_excerpt"] += "; array item keys: " + ", ".join(child_keys[:12])
            for item in value[:5]:
                if not isinstance(item, dict):
                    continue
                for key, child in item.items():
                    if not isinstance(key, str) or not re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", key):
                        continue
                    child_path = f"{prefix}.{key}"
                    self._walk_context_object(child, child_path, specs)

    def _get_context_field_specs(self) -> list:
        ctx_content = self._read_context_content()
        specs = {}
        for name in sorted(self._BUILTIN_GLOBAL_FIELDS):
            self._add_context_spec(specs, name, "string", path=name, source="[GLOBAL] 系统全局变量，每个 Prompt 都可直接使用")
            specs[name]["namespace"] = "GLOBAL"
        if ctx_content:
            candidates = self._extract_json_like_blocks(ctx_content)
            for block in candidates:
                cleaned = self._clean_json_like(block)
                if not cleaned.startswith(("{", "[")):
                    continue
                try:
                    parsed = json.loads(cleaned)
                except Exception:
                    continue
                self._walk_context_object(parsed, "", specs)

            patterns = [
                r'"([A-Za-z_][A-Za-z0-9_]*)"\s*:',
                r'\|\s*([A-Za-z_][A-Za-z0-9_]*)\s*\|',
                r'`([A-Za-z_][A-Za-z0-9_]*)`',
                r'\{\{\s*([A-Za-z_][A-Za-z0-9_]*)\s*\}\}',
            ]
            for pattern in patterns:
                for m in re.finditer(pattern, ctx_content):
                    word = m.group(1)
                    existing_leaf_names = {name.split(".")[-1] for name in specs}
                    if self._is_valid_field_name(word.lower()) and word not in specs and word not in existing_leaf_names:
                        self._add_context_spec(specs, word, "string", path=word, source="[extracted from Context prose]")
        return sorted(specs.values(), key=lambda item: item.get("name", ""))

    def _get_context_fields(self) -> set:
        return {spec["name"] for spec in self._get_context_field_specs()}
        """
        提取 Context 文档中定义的所有字段名。
        返回小写字段名集合，如无 context 返回空集。
        使用多级回退模式以提高对不同文档格式的适应性。
        """
        ctx_content = self._read_context_content()
        if not ctx_content:
            return set(self._BUILTIN_GLOBAL_FIELDS)

        fields = set(self._BUILTIN_GLOBAL_FIELDS)

        # Pattern 1: snake_case + colon/paren
        # "investment_goal: ..." or "investment_goal": in JSON
        for m in re.finditer(r'([a-z][a-z0-9_]+[a-z0-9])\s*"?\s*[:：\（(]', ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        # Pattern 2: JSON quoted keys  "field_name":
        for m in re.finditer(r'"([a-z][a-z0-9_]+[a-z0-9])"\s*:', ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        # Pattern 3: markdown table cells | field_name |
        for m in re.finditer(r'\|\s*([a-z][a-z0-9_]+[a-z0-9])\s*\|', ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        # Pattern 4: inline backtick `field_name`
        for m in re.finditer(r'`([a-z][a-z0-9_]+[a-z0-9])`', ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        # Pattern 4b: prompt-style placeholders in context docs, e.g. {{target_language}}
        for m in re.finditer(r'\{\{\s*([a-z][a-z0-9_]+[a-z0-9])\s*\}\}', ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        # Pattern 5: bullet / numbered list items (line starts with - * + or digit.)
        # Captures field_name even without a colon after it
        for m in re.finditer(r'(?:^|\n)\s*[-*\d+.)]+[\s ]+([a-z][a-z0-9_]+[a-z0-9])',
                             ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        # Pattern 6: Chinese context — 字段/参数/属性/输入/输出 紧接 field_name
        for m in re.finditer(r'[字参属输][段数字性出入][\s：:，,]*([a-z][a-z0-9_]+[a-z0-9])',
                             ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        # Pattern 7: line that mixes Chinese characters and a snake_case word
        # (typical pattern in bilingual context documents)
        for m in re.finditer(r'[一-鿿].{0,10}?([a-z][a-z0-9_]{3,}[a-z0-9])',
                             ctx_content):
            word = m.group(1).lower()
            if self._is_valid_field_name(word):
                fields.add(word)

        return fields

    def _extract_field_refs_from_text(self, text: str) -> set:
        """
        从文本中提取疑似字段名的 snake_case 标识符。
        排除常见非字段英文单词和短名称。
        """
        skip_words = RequirementPanel._GENERIC_KEYS | {
            "true", "false", "none", "null", "string", "number", "boolean",
            "array", "object", "integer", "float", "dict", "list", "tuple",
            "format", "output", "input",
            "example", "field", "fields",
            "this", "that", "these", "those", "with", "from", "into",
            "json", "xml", "html", "csv", "yaml", "config", "configs",
        }
        refs = set()
        for m in re.finditer(r'\b[a-z][a-z0-9_]+[a-z0-9]\b', text, re.I):
            word = m.group().lower()
            # Require underscore and minimum length 6 to reduce noise
            if '_' not in word or len(word) < 6:
                continue
            if word in skip_words:
                continue
            refs.add(word)
        return refs



# ─── Step 3: Prompt Panel ───────────────────────────────────────────────────

class PromptPanel(QWidget):
    """Step 3-4: Generate and review prompt."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.helper_llm = None
        self.requirements = ""
        self.scenario_name = ""
        self.vm = None
        self.current_prompt = ""
        self.version = 0
        self.eval_data = []
        self.continue_mode = False
        self.context_fields = set()
        self.context_field_specs = []
        self.base_prompt = ""
        self.candidate_reason = ""
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        title = QLabel("Step 3: 生成 Prompt")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        # Info
        self.info_label = QLabel("")
        self.info_label.setWordWrap(True)
        self.info_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.info_label)

        # ─── V2.2: System template selector ───
        template_layout = QHBoxLayout()
        template_layout.addWidget(QLabel("参考模板:"))
        self.template_combo = QComboBox()
        self.template_combo.addItem("(无 - 自由生成)", "")
        self.template_combo.setMinimumWidth(250)
        self.template_combo.setToolTip("选择 system_info/ 中的预设模板来约束 Prompt 生成格式")
        template_layout.addWidget(self.template_combo)
        self.template_info_label = QLabel("")
        self.template_info_label.setStyleSheet("color: #999; font-size: 11px;")
        template_layout.addWidget(self.template_info_label)
        template_layout.addStretch()
        layout.addLayout(template_layout)

        # Buttons
        btn_layout = QHBoxLayout()
        self.generate_btn = QPushButton("自动生成 Prompt")
        self.generate_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 20px;"
            " border-radius: 4px; }"
            "QPushButton:hover { background: #1565C0; }"
        )
        self.generate_btn.clicked.connect(self._generate)
        btn_layout.addWidget(self.generate_btn)

        self.confirm_btn = QPushButton("✓ 确认 (C)")
        self.confirm_btn.setEnabled(False)
        self.confirm_btn.clicked.connect(self._confirm)
        btn_layout.addWidget(self.confirm_btn)

        self.edit_btn = QPushButton("✎ 编辑 (E)")
        self.edit_btn.setEnabled(False)
        self.edit_btn.clicked.connect(self._edit)
        btn_layout.addWidget(self.edit_btn)

        self.regenerate_btn = QPushButton("↻ 重新生成 (R)")
        self.regenerate_btn.setEnabled(False)
        self.regenerate_btn.clicked.connect(self._generate)
        btn_layout.addWidget(self.regenerate_btn)

        self.import_btn = QPushButton("导入 Prompt")
        self.import_btn.clicked.connect(self._import_prompt)
        btn_layout.addWidget(self.import_btn)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # Prompt display / editor
        self.prompt_edit = QPlainTextEdit()
        self.prompt_edit.setFont(QFont("Consolas", 10))
        self.prompt_edit.setMinimumHeight(200)
        self.prompt_edit.setReadOnly(True)
        self.prompt_edit.setPlaceholderText("点击「自动生成 Prompt」开始...")
        _enable_width_wrap(self.prompt_edit)
        layout.addWidget(self.prompt_edit)

        # Placeholders info
        self.placeholder_label = QLabel("")
        self.placeholder_label.setStyleSheet("color: #666;")
        self.placeholder_label.setWordWrap(True)
        self.placeholder_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.placeholder_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.placeholder_label)

        # Status
        self.status_label = QLabel("")
        layout.addWidget(self.status_label)

    def set_context(self, data: dict):
        self.helper_llm = data.get("helper_llm")
        self.requirements = data.get("requirements", "")
        self.scenario_name = data.get("scenario_name", "")
        self.vm = data.get("vm")
        self.eval_data = data.get("eval_data", [])
        self.continue_mode = data.get("continue_mode", False)
        self.context_fields = set(data.get("selected_context_fields") or data.get("context_fields", set()))
        self.context_field_specs = [
            spec for spec in (data.get("context_field_specs") or [])
            if spec.get("name") in self.context_fields
        ]
        self.candidate_reason = data.get("candidate_reason", "")
        self.preferred_template_name = data.get("preferred_template_name", "")
        self.info_label.setText(
            f"场景: {self.scenario_name}  |  需求已记录"
        )

        # Load available system templates (V2.2)
        self._refresh_templates()
        self._select_preferred_template()

        # If continuing from existing version
        if data.get("continue_mode") and data.get("prompt_template"):
            self.current_prompt = data["prompt_template"]
            self.version = data.get("existing_version", 0)
            try:
                self.base_prompt = self.vm.load_version(self.version) if self.vm and self.version else self.current_prompt
            except Exception:
                self.base_prompt = self.current_prompt
            self.prompt_edit.setPlainText(self.current_prompt)
            self.prompt_edit.setReadOnly(True)
            self._update_placeholders()

            # Detect if output schema was auto-injected
            if self.candidate_reason:
                self.status_label.setText(self.candidate_reason)
            elif "<!-- AUTO_OUTPUT_SCHEMA_START -->" in self.current_prompt:
                self.status_label.setText(
                    f"已加载 v{self.version} 的 Prompt（含自动注入的输出字段定义）。")
            else:
                self.status_label.setText(f"已加载 v{self.version} 的 Prompt。")
            self.status_label.setStyleSheet("color: green;")
            self.confirm_btn.setEnabled(True)
            self.edit_btn.setEnabled(True)
            self.regenerate_btn.setEnabled(True)
            self.generate_btn.setEnabled(False)
        else:
            self.current_prompt = ""
            self.base_prompt = ""
            self.version = 0
            self.prompt_edit.clear()
            self.prompt_edit.setReadOnly(True)
            self.placeholder_label.setText("")
            next_v = (self.vm.get_current_version() + 1) if self.vm else 1
            self.status_label.setText(f"请点击「生成 Prompt」。确认保存时将写入 v{next_v}。")
            self.status_label.setStyleSheet("color: #1565C0;")
            self.confirm_btn.setEnabled(False)
            self.edit_btn.setEnabled(False)
            self.regenerate_btn.setEnabled(False)
            self.generate_btn.setEnabled(True)
            self.import_btn.setEnabled(True)

    def _refresh_templates(self):
        """Load available system preset templates (V2.2)."""
        from code.pipeline.prompt_generator import load_system_templates
        templates = load_system_templates()
        self._template_map = {}  # name -> content
        self.template_combo.clear()
        self.template_combo.addItem("(无 - 自由生成)", "")
        for t in templates:
            display = f"{t['name']}  ({len(t['content'])} chars)"
            self.template_combo.addItem(display, t['name'])
            self._template_map[t['name']] = t['content']
        if templates:
            self.template_combo.setCurrentIndex(1)
        self.template_info_label.setText(
            f"已加载 {len(templates)} 个系统模板" if templates else "(无系统模板)"
        )

    def _select_preferred_template(self):
        if not self.preferred_template_name:
            return
        for i in range(self.template_combo.count()):
            if self.template_combo.itemData(i) == self.preferred_template_name:
                self.template_combo.setCurrentIndex(i)
                return

    def _field_schema(self):
        specs = [
            dict(spec) for spec in (self.context_field_specs or [])
            if spec.get("name") in set(self.context_fields or set())
        ]
        seen = {spec.get("name") for spec in specs}
        for name in sorted(set(self.context_fields or set()) - seen):
            specs.append({"name": name, "type": "string", "required": True})
        return {"input_fields": specs, "output_fields": []}

    def _generate(self):
        if not callable(self.helper_llm):
            QMessageBox.warning(self, "Helper 模型未加载", "请先回到 Step 1 选择并确认 Helper 模型。")
            return
        self.generate_btn.setEnabled(False)
        self.prompt_edit.setPlainText("正在生成 Prompt 模板...")
        self.prompt_edit.setReadOnly(True)
        self.placeholder_label.setText("正在根据已勾选 Context 字段白名单生成 Prompt...")
        self.placeholder_label.setStyleSheet("color: #666;")
        self.status_label.setText("正在调用 Helper 模型生成 Prompt...")
        self.status_label.setStyleSheet("color: #666;")

        # Get selected template content (V2.2)
        template_name = self.template_combo.currentData()
        template_ref = self._template_map.get(template_name) if template_name else None
        if template_ref:
            self.status_label.setText(f"正在使用模板「{template_name}」生成 Prompt...")

        field_schema = self._field_schema()
        self.worker = GeneratePromptWorker(
            self.helper_llm, self.requirements, template_ref,
            field_schema=field_schema,
        )
        self.worker.finished.connect(self._on_generated)
        self.worker.error.connect(self._on_generate_error)
        self.worker.start()

    def _import_prompt(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("导入 Prompt")
        dialog.resize(760, 560)
        layout = QVBoxLayout(dialog)

        target_scene = self.scenario_name or "测试_具体用途"
        desc = QLabel(
            f"粘贴已有 Prompt。当前将关联到场景「{target_scene}」。\n"
            "如果当前没有场景上下文，系统会要求填写具体用途，并保存到 data/prompt/测试_具体用途/。"
        )
        desc.setWordWrap(True)
        layout.addWidget(desc)

        editor = QPlainTextEdit()
        editor.setFont(QFont("Consolas", 10))
        editor.setPlaceholderText("在这里粘贴你的 Prompt...")
        _enable_width_wrap(editor)
        layout.addWidget(editor, 1)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        if dialog.exec_() != QDialog.Accepted:
            return
        prompt = editor.toPlainText().strip()
        if not prompt:
            QMessageBox.warning(self, "提示", "导入的 Prompt 为空。")
            return

        parent = self.parent()
        while parent and not hasattr(parent, 'context'):
            parent = parent.parent()
        if not self.context_fields and parent and hasattr(parent, "req_panel"):
            try:
                self.context_fields = parent.req_panel._get_active_context_fields()
                self.context_field_specs = parent.req_panel._get_active_context_field_specs()
                parent.context["context_fields"] = self.context_fields
                parent.context["selected_context_fields"] = self.context_fields
                parent.context["context_field_specs"] = self.context_field_specs
            except Exception:
                pass

        if not self.scenario_name:
            resolved_name = _resolve_scratch_scenario_name(self, "测试")
            if not resolved_name:
                return
            self.scenario_name = resolved_name
            self.vm = VersionManager(self.scenario_name)
            self.requirements = self.requirements or "导入 Prompt 评测"
            self.info_label.setText(f"场景: {self.scenario_name}  |  导入 Prompt")
            if parent:
                parent.context.update({
                    "scenario_name": self.scenario_name,
                    "requirements": self.requirements,
                    "vm": self.vm,
                    "continue_mode": False,
                    "version": 0,
                })

        prompt = normalize_placeholders_to_double_braces(prompt, set(self.context_fields or set()))
        self.current_prompt = prompt
        self.base_prompt = ""
        self.version = 0
        self.prompt_edit.setPlainText(prompt)
        self.prompt_edit.setReadOnly(True)
        self._update_placeholders()
        next_v = (self.vm.get_current_version() + 1) if self.vm else 1
        self.status_label.setText(f"✓ Prompt 已导入到场景「{self.scenario_name}」，确认后将保存为 v{next_v}。")
        self.status_label.setStyleSheet("color: green;")
        self.confirm_btn.setEnabled(True)
        self.edit_btn.setEnabled(True)
        self.regenerate_btn.setEnabled(True)
        self.generate_btn.setEnabled(True)

    def _on_generated(self, prompt):
        self.current_prompt = prompt
        self.base_prompt = ""
        self.version = 0
        self.prompt_edit.setPlainText(prompt)
        self.prompt_edit.setReadOnly(True)
        self._update_placeholders()
        self.status_label.setText("✓ Prompt 已生成，请审查 (C=确认 / E=编辑 / R=重新生成)")
        self.status_label.setStyleSheet("color: green;")
        self.confirm_btn.setEnabled(True)
        self.edit_btn.setEnabled(True)
        self.regenerate_btn.setEnabled(True)
        _notify_long_task_done(
            self, "Prompt 生成完成",
            "Prompt 已生成，请回到 Step 3 审查并确认。"
        )

    def _on_generate_error(self, error_msg):
        self.status_label.setText(f"生成失败: {error_msg}")
        self.status_label.setStyleSheet("color: red;")
        self.generate_btn.setEnabled(True)
        self.prompt_edit.setPlainText("")

    def _update_placeholders(self):
        placeholders = extract_placeholders(self.current_prompt)
        selected_fields = sorted(set(self.context_fields or set()))
        if placeholders:
            unknown = self._unknown_placeholders()
            if unknown:
                self.placeholder_label.setText(
                    "Prompt 使用了未允许的 Context 占位符: " + ", ".join(unknown)
                    + "\n如果你只勾选了某个对象下面的子字段，Prompt 必须改用这些精确子字段；只有勾选整个对象时才允许使用父级对象占位符。"
                )
                self.placeholder_label.setStyleSheet("color: #C62828; font-weight: bold;")
            else:
                placeholder_set = set(placeholders)
                unused = [name for name in selected_fields if name not in placeholder_set]
                text = (
                    f"已勾选 Context 字段({len(selected_fields)}): "
                    f"{', '.join(selected_fields) if selected_fields else '无'}\n"
                    f"Prompt 实际使用输入占位符({len(placeholders)}): {', '.join(placeholders)}"
                )
                if unused:
                    text += f"\n已勾选但当前 Prompt 未使用({len(unused)}): {', '.join(unused)}"
                self.placeholder_label.setText(text)
                self.placeholder_label.setStyleSheet("color: #1565C0;")
        else:
            self.placeholder_label.setText(
                f"已勾选 Context 字段({len(selected_fields)}): "
                f"{', '.join(selected_fields) if selected_fields else '无'}\n"
                "Prompt 未实际使用任何输入占位符：当前 Prompt 不会读取评测样本 input 字段。"
            )
            self.placeholder_label.setStyleSheet("color: #F57C00;")
    def _unknown_placeholders(self):
        placeholders = set(extract_placeholders(self.current_prompt))
        return sorted(placeholders - set(self.context_fields or set()))

    def _confirm(self):
        if not self.current_prompt.strip():
            QMessageBox.warning(self, "警告", "Prompt 内容为空。")
            return
        if not self.vm:
            self.scenario_name = self.scenario_name or "测试"
            resolved_name = _resolve_scratch_scenario_name(self, self.scenario_name)
            if not resolved_name:
                return
            self.scenario_name = resolved_name
            self.vm = VersionManager(self.scenario_name)
            self.requirements = self.requirements or "导入 Prompt 评测"
            self.info_label.setText(f"场景: {self.scenario_name}  |  导入 Prompt")
        normalized = normalize_placeholders_to_double_braces(
            self.current_prompt, set(self.context_fields or set())
        )
        if normalized != self.current_prompt:
            self.current_prompt = normalized
            self.prompt_edit.setPlainText(normalized)
            self.status_label.setText("已将字段占位符统一为上线 Go template 格式 {{.field_name}}。")
            self.status_label.setStyleSheet("color: #1565C0;")
        unknown = self._unknown_placeholders()
        if unknown:
            QMessageBox.critical(
                self, "占位符不允许使用",
                "Prompt 含有未在 Step 2 勾选的 Context 占位符，不能确认保存。\n\n"
                f"不允许的占位符: {', '.join(unknown)}\n\n"
                "请删除这些占位符，或回到 Step 2 勾选对应字段。"
            )
            self.status_label.setText("Prompt 含有未勾选 Context 字段，已阻止确认。")
            self.status_label.setStyleSheet("color: red; font-weight: bold;")
            return

        changed_from_base = self._prompt_changed_from_base()
        if self.version != 0 and not changed_from_base:
            self.status_label.setText(f"✓ Prompt 未发生变化，继续使用当前版本 v{self.version}")
            parent = self.parent()
            while parent and not hasattr(parent, 'on_step_complete'):
                parent = parent.parent()
            if parent:
                parent.on_step_complete(2, {
                    "scenario_name": self.scenario_name,
                    "requirements": self.requirements,
                    "prompt_template": self.current_prompt,
                    "version": self.version,
                    "vm": self.vm,
                    "context_fields": set(self.context_fields or set()),
                    "selected_context_fields": set(self.context_fields or set()),
                    "context_field_specs": list(self.context_field_specs or []),
                })
            return

        # V3: If eval_data is available, inject output schema to constrain model output
        schema_injected = False
        if self.eval_data and self.vm:
            from code.pipeline.prompt_generator import augment_prompt_with_output_schema
            augmented = augment_prompt_with_output_schema(self.current_prompt, self.eval_data)
            if augmented != self.current_prompt:
                self.current_prompt = augmented
                self.prompt_edit.setPlainText(augmented)
                schema_injected = True
                self.status_label.setText("已注入输出字段定义表，将保存为新版本…")
                self.status_label.setStyleSheet("color: #F57C00;")

        if self.version == 0 or changed_from_base:
            self.version = self.vm.save_version(self.current_prompt)
            v_path = os.path.join(self.vm.root_dir, "data", "prompt", self.scenario_name, f"v{self.version}.py")
            msg = f"✓ Prompt 已保存为 v{self.version}"
            if schema_injected:
                msg += "（含自动注入的输出字段定义）"
            msg += f"\n  保存路径: {v_path}"
            self.status_label.setText(msg)
        else:
            self.status_label.setText(f"✓ Prompt 未发生变化，继续使用当前版本 v{self.version}")

        parent = self.parent()
        while parent and not hasattr(parent, 'on_step_complete'):
            parent = parent.parent()
        if parent:
            parent.on_step_complete(2, {
                "scenario_name": self.scenario_name,
                "requirements": self.requirements,
                "prompt_template": self.current_prompt,
                "version": self.version,
                "vm": self.vm,
                "context_fields": set(self.context_fields or set()),
                "selected_context_fields": set(self.context_fields or set()),
                "context_field_specs": list(self.context_field_specs or []),
            })

    def _edit(self):
        self.prompt_edit.setReadOnly(False)
        self.prompt_edit.setFocus()
        self.prompt_edit.selectAll()
        self.status_label.setText("请编辑 Prompt，编辑完成后点击「确认 (C)」保存。")
        self.status_label.setStyleSheet("color: #F57C00;")
        self.confirm_btn.setText("✓ 保存编辑")
        self.confirm_btn.clicked.disconnect()
        self.confirm_btn.clicked.connect(self._save_edit)

    def _save_edit(self):
        edited = self.prompt_edit.toPlainText().strip()
        if edited:
            self.current_prompt = edited
            self._update_placeholders()
        self.prompt_edit.setReadOnly(True)
        self.status_label.setText("✓ 编辑已保存")
        self.status_label.setStyleSheet("color: green;")
        self.confirm_btn.setText("✓ 确认 (C)")
        self.confirm_btn.clicked.disconnect()
        self.confirm_btn.clicked.connect(self._confirm)

    def _prompt_changed_from_base(self):
        if not self.base_prompt:
            return bool(self.current_prompt.strip())
        return self.current_prompt.strip() != self.base_prompt.strip()


# ─── Step 4: Eval Data Panel ────────────────────────────────────────────────

class EvalDataPanel(QWidget):
    """Step 5: Generate evaluation data."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.helper_llm = None
        self.prompt_template = ""
        self.requirements = ""
        self.eval_data = []
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        title = QLabel("Step 4: 生成评测数据")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        desc = QLabel(
            "生成用于评测的测试数据集。没有种子数据时默认使用「全自动生成」。"
            "全自动模式由 Helper 模型同时构造 input 和期望 output（ground truth），"
            "运行评测时再用 Target 模型输出与 ground truth 对比。"
        )
        desc.setWordWrap(True)
        layout.addWidget(desc)

        # Mode selection
        mode_layout = QHBoxLayout()
        self.auto_radio = QCheckBox("全自动生成")
        self.seed_radio = QCheckBox("种子扩充")
        self.auto_radio.setChecked(True)
        mode_layout.addWidget(self.auto_radio)
        mode_layout.addWidget(self.seed_radio)
        mode_layout.addStretch()
        layout.addLayout(mode_layout)

        # Seed data input (shown when seed mode)
        self.seed_label = QLabel("种子数据 (JSON数组格式):")
        layout.addWidget(self.seed_label)
        self.seed_edit = QTextEdit()
        self.seed_edit.setPlaceholderText(
            '例如:\n[\n'
            '  {\n'
            '    "input": {"field1": "value1", "field2": "value2"},\n'
            '    "output": {"result": "expected_output"}\n'
            '  }\n'
            ']'
        )
        self.seed_edit.setMaximumHeight(150)
        _enable_width_wrap(self.seed_edit)
        layout.addWidget(self.seed_edit)

        # Buttons
        btn_layout = QHBoxLayout()
        self.generate_btn = QPushButton("生成评测数据")
        self.generate_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 20px;"
            " border-radius: 4px; }"
            "QPushButton:hover { background: #1565C0; }"
        )
        self.generate_btn.clicked.connect(self._generate)
        btn_layout.addWidget(self.generate_btn)

        self.confirm_btn = QPushButton("✓ 确认数据")
        self.confirm_btn.setEnabled(False)
        self.confirm_btn.clicked.connect(self._confirm)
        btn_layout.addWidget(self.confirm_btn)

        self.edit_btn = QPushButton("✎ 编辑")
        self.edit_btn.setEnabled(False)
        self.edit_btn.clicked.connect(self._enable_edit)
        btn_layout.addWidget(self.edit_btn)

        self.regenerate_btn = QPushButton("↻ 重新生成")
        self.regenerate_btn.setEnabled(False)
        self.regenerate_btn.clicked.connect(self._generate)
        btn_layout.addWidget(self.regenerate_btn)

        self.revise_prompt_btn = QPushButton("修改 Prompt")
        self.revise_prompt_btn.setEnabled(False)
        self.revise_prompt_btn.clicked.connect(self._revise_prompt)
        btn_layout.addWidget(self.revise_prompt_btn)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # Data display
        self.data_edit = QPlainTextEdit()
        self.data_edit.setFont(QFont("Consolas", 10))
        self.data_edit.setReadOnly(True)
        self.data_edit.setMinimumHeight(180)
        self.data_edit.setPlaceholderText("生成的评测数据将显示在这里...")
        _enable_width_wrap(self.data_edit)
        layout.addWidget(self.data_edit)

        self.status_label = QLabel("")
        layout.addWidget(self.status_label)

        # Connect radio buttons
        self.auto_radio.toggled.connect(self._toggle_mode)
        self.seed_radio.toggled.connect(self._toggle_mode)
        self._toggle_mode()

    def set_context(self, data: dict):
        self.helper_llm = data.get("helper_llm")
        self.prompt_template = data.get("prompt_template", "")
        self.requirements = data.get("requirements", "")

        existing_data = data.get("eval_data", [])
        if existing_data:
            self.eval_data = existing_data
            self.data_edit.setPlainText(json.dumps(existing_data, ensure_ascii=False, indent=2))
            self.data_edit.setReadOnly(True)
            self.confirm_btn.setEnabled(True)
            self.edit_btn.setEnabled(True)
            self.regenerate_btn.setEnabled(True)
            self.revise_prompt_btn.setEnabled(bool(self.prompt_template.strip()))
            self.status_label.setText(f"已加载 {len(existing_data)} 个评测样本。")
            self.status_label.setStyleSheet("color: green;")
        else:
            self.eval_data = []
            self.data_edit.clear()
            self.data_edit.setReadOnly(True)
            self.confirm_btn.setEnabled(False)
            self.edit_btn.setEnabled(False)
            self.regenerate_btn.setEnabled(False)
            self.revise_prompt_btn.setEnabled(bool(self.prompt_template.strip()))
            self.generate_btn.setEnabled(True)
            self.status_label.setText("")

    def _toggle_mode(self):
        is_seed = self.seed_radio.isChecked()
        self.seed_label.setVisible(is_seed)
        self.seed_edit.setVisible(is_seed)

    def _generate(self):
        if not callable(self.helper_llm):
            QMessageBox.warning(self, "Helper 模型未加载", "请先回到 Step 1 选择并确认 Helper 模型。")
            return
        self.generate_btn.setEnabled(False)
        self.data_edit.setPlainText("正在生成...")
        self.status_label.setText("正在调用 Helper 模型生成评测数据...")
        self.status_label.setStyleSheet("color: #666;")

        if self.auto_radio.isChecked():
            self.worker = GenerateEvalDataAutoWorker(
                self.helper_llm, self.prompt_template, self.requirements
            )
        else:
            seeds = self.seed_edit.toPlainText().strip()
            if not seeds:
                QMessageBox.warning(self, "警告", "请输入种子数据。")
                self.generate_btn.setEnabled(True)
                self.status_label.setText("")
                return
            try:
                json.loads(seeds)
            except json.JSONDecodeError as e:
                QMessageBox.warning(self, "JSON 格式错误", f"种子数据 JSON 解析失败: {e}")
                self.generate_btn.setEnabled(True)
                return
            self.worker = GenerateEvalDataSeedWorker(
                self.helper_llm, self.prompt_template, self.requirements, seeds
            )

        self.worker.finished.connect(self._on_generated)
        self.worker.error.connect(self._on_generate_error)
        self.worker.progress.connect(self.status_label.setText)
        self.worker.start()

    def _on_generated(self, data):
        self.eval_data = data
        formatted = json.dumps(data, ensure_ascii=False, indent=2)
        self.data_edit.setPlainText(formatted)
        self.data_edit.setReadOnly(True)
        self.status_label.setText(f"✓ 生成了 {len(data)} 个评测样本。请审查。")
        self.status_label.setStyleSheet("color: green;")
        self.confirm_btn.setEnabled(True)
        self.edit_btn.setEnabled(True)
        self.regenerate_btn.setEnabled(True)
        self.revise_prompt_btn.setEnabled(True)
        self.generate_btn.setEnabled(True)
        _notify_long_task_done(
            self, "评测数据生成完成",
            f"已生成 {len(data)} 个评测样本，请回到 Step 4 审查。",
            force=True,
        )

    def _on_generate_error(self, error_msg):
        self.status_label.setText(f"生成失败: {error_msg}")
        self.status_label.setStyleSheet("color: red;")
        self.generate_btn.setEnabled(True)
        self.data_edit.setPlainText("")

    def _revise_prompt(self):
        if not self.prompt_template.strip():
            QMessageBox.warning(self, "提示", "当前没有可修改的 Prompt，请先完成 Step 3。")
            return

        dialog = PromptRevisionDialog(self)
        if dialog.exec_() != QDialog.Accepted:
            return

        revision = dialog.revision_text()
        if not revision:
            QMessageBox.information(self, "提示", "没有输入修改方向，已取消。")
            return

        parent = self.parent()
        while parent and not hasattr(parent, 'context'):
            parent = parent.parent()
        if not parent:
            QMessageBox.warning(self, "提示", "无法找到主窗口上下文。")
            return
        if not parent.context.get("helper_llm"):
            QMessageBox.warning(self, "提示", "Helper 模型未加载，无法重新生成 Prompt。")
            return

        base_requirements = parent.context.get("requirements", self.requirements)
        parent.context.update({
            "prompt_template": self.prompt_template,
            "version": 0,
            "continue_mode": True,
            "eval_data": [],
            "eval_script": "",
            "results": None,
            "candidate_reason": "正在根据 Step 4 的修改方向生成候选 Prompt，确认后才会保存为新版本。",
        })

        parent.prompt_panel.set_context(parent.context)
        parent.go_to_step(3)
        parent.prompt_panel.status_label.setText("已收到修改方向，正在重新生成 Prompt...")
        parent.prompt_panel.status_label.setStyleSheet("color: #1565C0;")

        field_schema = {
            "input_fields": parent.context.get("context_field_specs") or [
                {"name": name, "type": "string", "required": True}
                for name in sorted(parent.context.get("context_fields", set()))
            ],
            "output_fields": [],
        }
        self.revise_worker = ContinuePromptOptimizeWorker(
            parent.context.get("helper_llm"),
            parent.context.get("scenario_name", ""),
            self.prompt_template,
            base_requirements,
            revision,
            field_schema=field_schema,
        )

        def _on_revised(new_prompt):
            parent.context.update({
                "prompt_template": new_prompt,
                "version": 0,
                "continue_mode": True,
                "candidate_reason": "已根据 Step 4 的修改方向生成候选 Prompt，确认后才会保存为新版本。",
            })
            parent.prompt_panel.set_context(parent.context)
            parent.go_to_step(3)

        def _on_revise_error(error_msg):
            parent.prompt_panel.status_label.setText(f"重新生成失败: {error_msg}")
            parent.prompt_panel.status_label.setStyleSheet("color: red;")

        self.revise_worker.finished.connect(_on_revised)
        self.revise_worker.error.connect(_on_revise_error)
        self.revise_worker.progress.connect(parent.prompt_panel.status_label.setText)
        self.revise_worker.start()

    def _enable_edit(self):
        self.data_edit.setReadOnly(False)
        self.data_edit.setFocus()
        self.status_label.setText("请编辑 JSON 数据，编辑后点击「确认数据」。")
        self.status_label.setStyleSheet("color: #F57C00;")
        self.confirm_btn.setText("✓ 保存编辑")

    def _confirm(self):
        if not self.data_edit.isReadOnly():
            # Save edit
            try:
                raw = self.data_edit.toPlainText().strip()
                parsed = json.loads(raw)
                if isinstance(parsed, list):
                    self.eval_data = parsed
                    formatted = json.dumps(parsed, ensure_ascii=False, indent=2)
                    self.data_edit.setPlainText(formatted)
                else:
                    QMessageBox.warning(self, "格式错误", "数据必须是 JSON 数组格式。")
                    return
            except json.JSONDecodeError as e:
                QMessageBox.warning(self, "JSON 错误", f"JSON 解析失败: {e}")
                return
            self.data_edit.setReadOnly(True)
            self.confirm_btn.setText("✓ 确认数据")

        if not self.eval_data:
            QMessageBox.warning(self, "警告", "评测数据为空。")
            return

        # Show save path info
        vm = None
        parent = self.parent()
        while parent and not hasattr(parent, 'context'):
            parent = parent.parent()
        if parent and hasattr(parent, 'context'):
            vm = parent.context.get("vm")

        save_msg = f"✓ 评测数据已确认 ({len(self.eval_data)} 个样本)"
        if vm and hasattr(vm, 'save_eval_data'):
            vm.save_eval_data(self.eval_data)
            eval_path = os.path.join(vm.root_dir, "data", "eval", f"{vm.scenario_name}.json")
            save_msg += f"\n  已保存到: {eval_path}"
        self.status_label.setText(save_msg)
        self.status_label.setStyleSheet("color: green;")

        parent = self.parent()
        while parent and not hasattr(parent, 'on_step_complete'):
            parent = parent.parent()
        if parent:
            parent.on_step_complete(3, {
                "eval_data": self.eval_data,
            })


# ─── Step 5: Eval Script Panel ──────────────────────────────────────────────

class EvalScriptPanel(QWidget):
    """Step 6: Generate evaluation script."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.scenario_name = ""
        self.prompt_template = ""
        self.eval_data = []
        self.target_info = None
        self.version = 1
        self.max_workers = 8
        self.vm = None
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        title = QLabel("Step 5: 生成评测脚本")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        desc = QLabel("系统将自动生成可独立运行的评测脚本。脚本会区分硬结构字段和自然语言文本字段，并在需要时自动使用 Helper Judge 做语义复判。")
        desc.setWordWrap(True)
        layout.addWidget(desc)

        btn_layout = QHBoxLayout()
        self.generate_btn = QPushButton("生成评测脚本")
        self.generate_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 20px;"
            " border-radius: 4px; }"
            "QPushButton:hover { background: #1565C0; }"
        )
        self.generate_btn.clicked.connect(self._generate)
        btn_layout.addWidget(self.generate_btn)

        self.save_btn = QPushButton("保存脚本")
        self.save_btn.setEnabled(False)
        self.save_btn.clicked.connect(self._save)
        btn_layout.addWidget(self.save_btn)

        self.confirm_btn = QPushButton("✓ 确认并继续")
        self.confirm_btn.setEnabled(False)
        self.confirm_btn.clicked.connect(self._confirm)
        btn_layout.addWidget(self.confirm_btn)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        self.script_edit = QPlainTextEdit()
        self.script_edit.setFont(QFont("Consolas", 10))
        self.script_edit.setReadOnly(True)
        self.script_edit.setMinimumHeight(250)
        self.script_edit.setPlaceholderText("评测脚本将显示在这里...")
        _enable_width_wrap(self.script_edit)
        layout.addWidget(self.script_edit)

        self.status_label = QLabel("")
        layout.addWidget(self.status_label)

    def set_context(self, data: dict):
        self.scenario_name = data.get("scenario_name", "")
        self.prompt_template = data.get("prompt_template", "")
        self.eval_data = data.get("eval_data", [])
        self.target_info = data.get("target_info")
        self.helper_info = data.get("helper_info")
        self.version = data.get("version", 1)
        self.max_workers = data.get("max_workers", 8)
        self.vm = data.get("vm")

    def _generate(self):
        if not self.eval_data:
            QMessageBox.warning(self, "警告", "暂无评测数据，请先到 Step 4 (评测数据) 重新生成评测数据。\n\n若已有已加载的评测脚本，可直接编辑后点击「确认并继续」跳过此步骤。")
            return
        if not self.target_info:
            QMessageBox.warning(self, "警告", "Target 模型信息缺失。")
            return

        self.generate_btn.setEnabled(False)
        self.script_edit.setPlainText("正在生成评测脚本...")
        self.status_label.setText("正在生成...")

        try:
            script = generate_eval_script(
                self.scenario_name, self.prompt_template, self.eval_data,
                self.target_info, self.version, self.max_workers,
                helper_info=self.helper_info
            )
            self.script_edit.setPlainText(script)
            self.script_edit.setReadOnly(False)
            self.status_label.setText("✓ 评测脚本已生成，可编辑修改后保存。")
            self.status_label.setStyleSheet("color: green;")
            self.save_btn.setEnabled(True)
            self.confirm_btn.setEnabled(True)
        except Exception as e:
            self.status_label.setText(f"生成失败: {e}")
            self.status_label.setStyleSheet("color: red;")

        self.generate_btn.setEnabled(True)

    def _save(self):
        script = self.script_edit.toPlainText()
        if self.vm:
            try:
                self.vm.save_eval_script(script)
                self.status_label.setText("✓ 评测脚本已保存。")
                self.status_label.setStyleSheet("color: green;")
            except Exception as e:
                QMessageBox.warning(self, "保存失败", str(e))
        else:
            path = QFileDialog.getSaveFileName(self, "保存评测脚本",
                                               f"eval_prompt_{self.scenario_name}.py",
                                               "Python Files (*.py)")
            if path[0]:
                with open(path[0], "w", encoding="utf-8") as f:
                    f.write(script)
                self.status_label.setText(f"✓ 已保存到: {path[0]}")

    def _confirm(self):
        # Auto-save script before confirming
        script = self.script_edit.toPlainText().strip()
        if script and self.vm:
            try:
                self.vm.save_eval_script(script)
                self.status_label.setText("✓ 评测脚本已保存。")
            except Exception as e:
                self.status_label.setText(f"保存失败: {e}")

        parent = self.parent()
        while parent and not hasattr(parent, 'on_step_complete'):
            parent = parent.parent()
        if parent:
            parent.on_step_complete(4, {})


# ─── Step 6: Eval Run Panel ─────────────────────────────────────────────────

class EvalRunPanel(QWidget):
    """Step 7: Run evaluation."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.target_llm = None
        self.prompt_template = ""
        self.eval_data = []
        self.scenario_name = ""
        self.version = 1
        self.max_workers = 8
        self.vm = None
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        title = QLabel("Step 6: 运行评测")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        # Info
        self.info_label = QLabel("")
        self.info_label.setWordWrap(True)
        layout.addWidget(self.info_label)

        # Progress
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        options_layout = QHBoxLayout()
        self.use_llm_judge_check = QCheckBox("启用全量语义 Judge（更严格，但会显著变慢）")
        self.use_llm_judge_check.setChecked(False)
        self.use_llm_judge_check.setToolTip(
            "关闭时系统仍会自动复判自然语言文本差异；开启时每个样本都会额外调用一次 Helper 模型。"
        )
        options_layout.addWidget(self.use_llm_judge_check)
        self.production_latency_check = QCheckBox("生产模式测速（移除自动评测 Schema）")
        self.production_latency_check.setChecked(False)
        self.production_latency_check.setToolTip(
            "仅用于估算真实应用延迟：调用 Target 模型前移除 AUTO_OUTPUT_SCHEMA 评测辅助区块。"
        )
        options_layout.addWidget(self.production_latency_check)
        options_layout.addStretch()
        layout.addLayout(options_layout)

        # Buttons
        btn_layout = QHBoxLayout()
        self.run_btn = QPushButton("▶ 开始评测")
        self.run_btn.setStyleSheet(
            "QPushButton { background: #388E3C; color: white; padding: 8px 24px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #2E7D32; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.run_btn.clicked.connect(self._run_eval)
        btn_layout.addWidget(self.run_btn)

        self.cancel_btn = QPushButton("取消")
        self.cancel_btn.setEnabled(False)
        self.cancel_btn.clicked.connect(self._cancel)
        btn_layout.addWidget(self.cancel_btn)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # Log output
        self.log_output = LogTextEdit()
        self.log_output.setMinimumHeight(200)
        layout.addWidget(self.log_output)

        self.status_label = QLabel("")
        layout.addWidget(self.status_label)

        # Keyboard shortcuts
        QShortcut(Qt.CTRL + Qt.Key_Return, self).activated.connect(self._run_eval)
        QShortcut(Qt.Key_Escape, self).activated.connect(self._cancel)

    def set_context(self, data: dict):
        self.target_llm = data.get("target_llm")
        self.helper_llm = data.get("helper_llm")
        self.prompt_template = data.get("prompt_template", "")
        self.eval_data = data.get("eval_data", [])
        self.scenario_name = data.get("scenario_name", "")
        self.version = data.get("version", 1)
        self.max_workers = data.get("max_workers", 8)
        self.vm = data.get("vm")

        self.info_label.setText(
            f"场景: {self.scenario_name}  |  v{self.version}  |  "
            f"样本数: {len(self.eval_data)}  |  并发: {self.max_workers}"
        )
        self.log_output.clear()
        self.log_output.log(f"评测准备就绪。点击「开始评测」运行。")

    def _run_eval(self):
        if not self.target_llm:
            QMessageBox.warning(self, "警告", "Target 模型未加载。")
            return
        if not self.eval_data:
            QMessageBox.warning(self, "警告", "评测数据为空。")
            return

        self.run_btn.setEnabled(False)
        self.cancel_btn.setEnabled(True)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.progress_bar.setMaximum(len(self.eval_data))
        self.log_output.clear()
        self.log_output.log("评测开始...")

        prompt_template = self.prompt_template
        if self.production_latency_check.isChecked():
            cleaned = strip_auto_output_schema(prompt_template)
            if cleaned != prompt_template:
                self.log_output.log(
                    f"生产模式测速：已移除自动评测 Schema，Prompt 长度 {len(prompt_template)} -> {len(cleaned)} 字符。"
                )
                prompt_template = cleaned

        result_path = None
        if self.vm:
            result_path = self.vm.get_result_path(self.version)
            self.log_output.log(f"结果将保存到: {result_path}")

        self.worker = RunEvalWorker(
            self.target_llm, prompt_template, self.eval_data,
            self.scenario_name, self.version, result_path, self.max_workers,
            helper_llm=self.helper_llm,
            use_llm_judge=self.use_llm_judge_check.isChecked(),
        )
        self.worker.progress.connect(self.log_output.log)
        self.worker.sample_done.connect(self._on_sample_done)
        self.worker.finished.connect(self._on_eval_done)
        self.worker.error.connect(self._on_eval_error)
        self.worker.start()

    def _on_sample_done(self, done, total):
        self.progress_bar.setValue(done)

    def _cancel(self):
        if hasattr(self, 'worker') and self.worker.isRunning():
            self.worker.cancelled = True
            self.log_output.log("正在取消评测（等待当前样本完成...）")
            self.cancel_btn.setEnabled(False)

    def _on_eval_done(self, results):
        self.run_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.progress_bar.setVisible(False)

        was_cancelled = getattr(self.worker, 'cancelled', False) if hasattr(self, 'worker') else False

        if was_cancelled:
            self.log_output.log("评测已取消。")
            self.status_label.setText("评测已取消")
            self.status_label.setStyleSheet("color: orange;")
            return

        self.log_output.log(f"评测完成!")
        self.log_output.log(f"总样本数: {results['total']}")
        self.log_output.log(f"正确数: {results['correct']}")
        self.log_output.log(f"准确率: {results['accuracy']:.2%}")
        if results.get("exact_match_accuracy") is not None:
            self.log_output.log(f"逐字准确率: {results['exact_match_accuracy']:.2%}")
        if results.get("structural_accuracy") is not None:
            self.log_output.log(f"结构准确率: {results['structural_accuracy']:.2%}")
        if results.get("semantic_eligible"):
            self.log_output.log(
                f"语义复判: {results.get('semantic_correct', 0)}/{results.get('semantic_eligible', 0)}"
            )
        self.log_output.log(f"评测样本平均耗时(含最多3次Target重试): {results['avg_time']:.2f}秒")
        if results.get("avg_target_attempt_time") is not None:
            self.log_output.log(
                f"单次Target完整响应平均耗时: {results['avg_target_attempt_time']:.2f}秒 "
                f"(Target调用 {results.get('target_attempts', 0)} 次)"
            )
        self.log_output.log(f"评测样本最小耗时: {results['min_time']:.2f}秒")
        self.log_output.log(f"评测样本最大耗时: {results['max_time']:.2f}秒")
        self.log_output.log("说明: 这里统计的是评测脚本等待完整响应并解析JSON的耗时，不等同于线上流式首字/首包延迟。")

        # Show error details
        error_count = 0
        for i, r in enumerate(results['results']):
            if not r.get('is_correct'):
                error_count += 1
                reason = r.get('error_reason', '未知错误')
                self.log_output.log(f"  样本{i}错误: {reason}")
                # Show raw response if available
                for k in sorted(r.keys()):
                    if k.startswith('response_') and not k.startswith('response_json_') and r[k]:
                        resp_text = str(r[k])[:200]
                        self.log_output.log(f"    {k}: {resp_text}")
                        break

        if error_count > 0:
            self.log_output.log(f"共 {error_count} 个样本失败，请查看上方日志了解具体原因")

        self.status_label.setText(
            f"准确率: {results['accuracy']:.2%}  |  正确: {results['correct']}/{results['total']}"
        )
        self.status_label.setStyleSheet(
            "color: green; font-weight: bold; font-size: 14px;"
        )

        parent = self.parent()
        while parent and not hasattr(parent, 'on_step_complete'):
            parent = parent.parent()
        if parent:
            parent.on_step_complete(5, {"results": results})
        _notify_long_task_done(
            self, "评测完成",
            f"评测已完成，准确率 {results['accuracy']:.2%}，正确 {results['correct']}/{results['total']}。"
        )

    def _on_eval_error(self, error_msg):
        self.run_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.progress_bar.setVisible(False)
        self.log_output.log_error(f"评测失败: {error_msg}")
        self.status_label.setText("评测失败")
        self.status_label.setStyleSheet("color: red;")


# ─── Step 7: Results Panel ──────────────────────────────────────────────────

class ResultsPanel(QWidget):
    """Step 8: Display evaluation results and error analysis."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.helper_llm = None
        self.prompt_template = ""
        self.results = None
        self.eval_data = []
        self.scenario_name = ""
        self.error_analysis = None
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        title = QLabel("Step 7: 结果分析与错误诊断")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        # Summary cards
        summary_layout = QHBoxLayout()
        self.accuracy_label = QLabel("准确率: --")
        self.accuracy_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self.accuracy_label.setStyleSheet(
            "background: #e3f2fd; padding: 12px; border-radius: 6px;"
            " font-size: 16px; font-weight: bold;"
        )
        summary_layout.addWidget(self.accuracy_label)

        self.time_label = QLabel("平均耗时: --")
        self.time_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self.time_label.setStyleSheet(
            "background: #fff3e0; padding: 12px; border-radius: 6px;"
            " font-size: 16px; font-weight: bold;"
        )
        summary_layout.addWidget(self.time_label)

        self.count_label = QLabel("正确/总数: --")
        self.count_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self.count_label.setStyleSheet(
            "background: #e8f5e9; padding: 12px; border-radius: 6px;"
            " font-size: 16px; font-weight: bold;"
        )
        summary_layout.addWidget(self.count_label)
        layout.addLayout(summary_layout)

        # Error pattern stats
        self.error_pattern_label = QLabel("")
        self.error_pattern_label.setWordWrap(True)
        self.error_pattern_label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        layout.addWidget(self.error_pattern_label)

        # Results table
        layout.addWidget(QLabel("样本详情:"))
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["样本", "状态", "耗时(秒)", "错误详情"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.itemDoubleClicked.connect(self._show_sample_detail)
        layout.addWidget(self.table)

        # Analyze button
        btn_layout = QHBoxLayout()
        self.analyze_btn = QPushButton("分析错误并生成优化建议")
        self.analyze_btn.setEnabled(False)
        self.analyze_btn.setStyleSheet(
            "QPushButton { background: #F57C00; color: white; padding: 8px 20px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #E65100; }"
        )
        self.analyze_btn.clicked.connect(self._analyze_errors)
        btn_layout.addWidget(self.analyze_btn)

        self.next_btn = QPushButton("▶ 去迭代优化")
        self.next_btn.setEnabled(False)
        self.next_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 24px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #1565C0; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.next_btn.clicked.connect(self._go_to_optimize)
        btn_layout.addWidget(self.next_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        self.status_label = QLabel("")
        layout.addWidget(self.status_label)

        # Keyboard shortcut
        QShortcut(Qt.CTRL + Qt.Key_Return, self).activated.connect(self._analyze_errors)

    def set_context(self, data: dict):
        self.helper_llm = data.get("helper_llm")
        self.prompt_template = data.get("prompt_template", "")
        self.results = data.get("results")
        self.eval_data = data.get("eval_data", [])
        self.scenario_name = data.get("scenario_name", "")
        self.error_analysis = None

        if self.results:
            self._display_results()

    def _display_results(self):
        r = self.results
        acc = r['accuracy']
        exact = r.get("exact_match_accuracy")
        if exact is not None:
            self.accuracy_label.setText(f"准确率: {acc:.2%} | 逐字: {exact:.2%}")
        else:
            self.accuracy_label.setText(f"准确率: {acc:.2%}")
        self.accuracy_label.setStyleSheet(
            f"background: {'#e8f5e9' if acc >= 0.8 else '#fff3e0' if acc >= 0.5 else '#ffebee'};"
            f" padding: 12px; border-radius: 6px; font-size: 16px; font-weight: bold;"
        )
        avg_attempt = r.get("avg_target_attempt_time")
        if avg_attempt is not None:
            self.time_label.setText(
                f"评测样本均耗时: {r['avg_time']:.2f}s | 单次Target均耗时: {avg_attempt:.2f}s"
            )
        else:
            self.time_label.setText(f"评测样本均耗时: {r['avg_time']:.2f}s")
        self.count_label.setText(f"正确/总数: {r['correct']}/{r['total']}")

        # Populate table
        self.table.setRowCount(len(r['results']))
        for i, sample_result in enumerate(r['results']):
            item_idx = QTableWidgetItem(str(i))
            self.table.setItem(i, 0, item_idx)

            is_correct = sample_result.get("is_correct", False)
            status_item = QTableWidgetItem("✓ 正确" if is_correct else "✗ 错误")
            status_item.setForeground(QColor("#2E7D32") if is_correct else QColor("#C62828"))
            self.table.setItem(i, 1, status_item)

            tc = sample_result.get("time_cost", 0)
            self.table.setItem(i, 2, QTableWidgetItem(f"{tc:.2f}"))

            error = sample_result.get("error_reason", "")
            self.table.setItem(i, 3, QTableWidgetItem(error if error else ""))

        self.analyze_btn.setEnabled(True)

    def _show_sample_detail(self, index):
        row = index.row()
        if not self.results:
            return
        r = self.results['results'][row]
        detail = json.dumps(r, ensure_ascii=False, indent=2)
        QMessageBox.information(self, f"样本 {row} 详情", detail)

    def _analyze_errors(self):
        if not self.results:
            return

        self.status_label.setText("正在分析错误...")
        self.status_label.setStyleSheet("color: #666;")
        self.analyze_btn.setEnabled(False)

        self.error_worker = AnalyzeErrorsWorker(self.results, self.eval_data)
        self.error_worker.finished.connect(self._on_analyzed)
        self.error_worker.error.connect(lambda e: self.status_label.setText(f"分析失败: {e}"))
        self.error_worker.start()

    def _on_analyzed(self, analysis):
        # Show error patterns
        patterns = analysis['error_patterns']
        pattern_text = "错误模式统计:  "
        for err_type, count in patterns.items():
            if count > 0:
                labels = {
                    "hard_structure_mismatch": "硬结构不匹配",
                    "semantic_quality_failure": "语义质量未达标",
                    "semantic_judge_unavailable": "缺少语义Judge",
                    "exact_text_mismatch": "逐字文案不一致",
                    "json_parse_error": "JSON解析错误",
                    "no_response": "无响应",
                    "unknown": "未知错误",
                }
                pattern_text += f"{labels.get(err_type, err_type)}: {count}个  "
        self.error_pattern_label.setText(pattern_text if analysis['errors'] else "所有样本全部正确!")
        self.error_pattern_label.setStyleSheet(
            "color: #c62828;" if analysis['errors'] else "color: #2E7D32;"
        )

        self.status_label.setText(
            f"错误分析完成: {len(analysis['errors'])} 个错误样本"
            if analysis['errors'] else "所有样本全部正确!"
        )
        self.status_label.setStyleSheet(
            "color: orange;" if analysis['errors'] else "color: green;"
        )

        # Store for next step
        self.error_analysis = analysis

        # Save error analysis to file
        if self.scenario_name:
            try:
                vm_root = self.parent()
                while vm_root and not hasattr(vm_root, 'context'):
                    vm_root = vm_root.parent()
                project_root = None
                if vm_root and hasattr(vm_root, 'context'):
                    project_root = vm_root.context.get("project_root")
                if not project_root:
                    project_root = os.path.normpath(
                        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                                     "..", "Virtual-Coach-main")
                    )
                eval_result_dir = os.path.join(project_root, "data", "eval_result")
                os.makedirs(eval_result_dir, exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                analysis_path = os.path.join(
                    eval_result_dir, f"{self.scenario_name}_analysis_{timestamp}.json"
                )
                with open(analysis_path, "w", encoding="utf-8") as f:
                    json.dump(analysis, f, ensure_ascii=False, indent=2)
                self.status_label.setText(
                    f"错误分析完成，已保存至 {analysis_path}"
                )
            except Exception as e:
                self.status_label.setText(f"错误分析完成，但保存文件失败: {e}")

        parent = self.parent()
        while parent and not hasattr(parent, 'on_step_complete'):
            parent = parent.parent()
        if parent and analysis['errors']:
            parent.on_step_complete(6, {
                "error_analysis": analysis,
                "results": self.results,
            })
        self.analyze_btn.setEnabled(True)
        self.next_btn.setEnabled(True)

    def _go_to_optimize(self):
        """Navigate to Step 8 (Iterative Optimization)."""
        parent = self.parent()
        while parent and not hasattr(parent, 'go_to_step'):
            parent = parent.parent()
        if parent:
            parent.go_to_step(8, {
                "error_analysis": self.error_analysis,
                "results": self.results,
            })


# ─── Step 8: Optimize Panel ─────────────────────────────────────────────────

class OptimizePanel(QWidget):
    """Step 9: Optimization suggestions and iteration."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.helper_llm = None
        self.target_llm = None
        self.prompt_template = ""
        self.eval_data = []
        self.results = None
        self.error_analysis = None
        self.scenario_name = ""
        self.version = 1
        self.target_info = None
        self.max_workers = 8
        self.vm = None
        self.requirements = ""
        self.context_fields = set()
        self.context_field_specs = []
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        title = QLabel("Step 8: 迭代优化")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        # Top action buttons
        btn_layout = QHBoxLayout()
        self.suggest_btn = QPushButton("获取优化建议")
        self.suggest_btn.setStyleSheet(
            "QPushButton { background: #F57C00; color: white; padding: 8px 20px;"
            " border-radius: 4px; }"
        )
        self.suggest_btn.clicked.connect(self._get_suggestions)
        btn_layout.addWidget(self.suggest_btn)

        self.auto_optimize_btn = QPushButton("LLM 自动优化")
        self.auto_optimize_btn.setEnabled(False)
        self.auto_optimize_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 20px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #1565C0; }"
        )
        self.auto_optimize_btn.setToolTip("基于 LLM 建议和手动建议，自动生成优化后的新版本 Prompt")
        self.auto_optimize_btn.clicked.connect(self._auto_optimize)
        btn_layout.addWidget(self.auto_optimize_btn)

        self.edit_prompt_btn = QPushButton("修改 Prompt")
        self.edit_prompt_btn.setEnabled(False)
        self.edit_prompt_btn.clicked.connect(self._edit_prompt)
        btn_layout.addWidget(self.edit_prompt_btn)

        self.save_version_btn = QPushButton("保存为新版本")
        self.save_version_btn.setEnabled(False)
        self.save_version_btn.clicked.connect(self._save_version)
        btn_layout.addWidget(self.save_version_btn)

        self.regen_data_btn = QPushButton("重新生成评测数据")
        self.regen_data_btn.setEnabled(False)
        self.regen_data_btn.clicked.connect(self._regen_data)
        btn_layout.addWidget(self.regen_data_btn)

        self.rerun_btn = QPushButton("重新评测")
        self.rerun_btn.setEnabled(False)
        self.rerun_btn.clicked.connect(self._rerun)
        btn_layout.addWidget(self.rerun_btn)

        self.deploy_btn = QPushButton("部署上线")
        self.deploy_btn.setEnabled(False)
        self.deploy_btn.setStyleSheet(
            "QPushButton { background: #388E3C; color: white; padding: 8px 20px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #2E7D32; }"
        )
        self.deploy_btn.clicked.connect(self._go_deploy)
        btn_layout.addWidget(self.deploy_btn)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # Splitter: suggestions + prompt editor
        splitter = QSplitter(Qt.Vertical)

        # LLM suggestions display (editable)
        sugg_widget = QWidget()
        sugg_layout = QVBoxLayout(sugg_widget)
        sugg_layout.setContentsMargins(0, 0, 0, 0)
        sugg_layout.addWidget(QLabel("LLM 优化建议 (可编辑):"))
        self.suggestions_display = QTextEdit()
        self.suggestions_display.setReadOnly(False)
        _enable_width_wrap(self.suggestions_display)
        self.suggestions_display.setFont(QFont("Consolas", 10))
        self.suggestions_display.setPlaceholderText("点击「获取优化建议」生成...")
        sugg_layout.addWidget(self.suggestions_display)
        splitter.addWidget(sugg_widget)

        # Manual suggestions (editable)
        manual_widget = QWidget()
        manual_layout = QVBoxLayout(manual_widget)
        manual_layout.setContentsMargins(0, 0, 0, 0)
        manual_layout.addWidget(QLabel("手动输入优化建议 (可编辑):"))
        self.manual_suggestions = QTextEdit()
        _enable_width_wrap(self.manual_suggestions)
        self.manual_suggestions.setFont(QFont("Consolas", 10))
        self.manual_suggestions.setPlaceholderText("在此手动输入您的优化建议，例如：\n1. 调整输出格式要求\n2. 增加对XX场景的处理说明\n3. 修正字段命名规范...")
        manual_layout.addWidget(self.manual_suggestions)
        splitter.addWidget(manual_widget)

        # Prompt editor
        editor_widget = QWidget()
        editor_layout = QVBoxLayout(editor_widget)
        editor_layout.setContentsMargins(0, 0, 0, 0)
        self.editor_title = QLabel("当前 Prompt (双击编辑):")
        editor_layout.addWidget(self.editor_title)
        self.prompt_editor = QPlainTextEdit()
        self.prompt_editor.setFont(QFont("Consolas", 10))
        self.prompt_editor.setMinimumHeight(150)
        self.prompt_editor.setReadOnly(True)
        _enable_width_wrap(self.prompt_editor)
        editor_layout.addWidget(self.prompt_editor)
        splitter.addWidget(editor_widget)

        layout.addWidget(splitter)

        # Version history
        version_widget = QWidget()
        version_layout = QHBoxLayout(version_widget)
        version_layout.setContentsMargins(0, 0, 0, 0)
        version_layout.addWidget(QLabel("版本历史:"))
        self.version_list = QListWidget()
        self.version_list.setMaximumHeight(80)
        self.version_list.itemClicked.connect(self._load_version)
        version_layout.addWidget(self.version_list)
        layout.addWidget(version_widget)

        self.status_label = QLabel("")
        layout.addWidget(self.status_label)

    def set_context(self, data: dict):
        self.helper_llm = data.get("helper_llm")
        self.target_llm = data.get("target_llm")
        self.vm = data.get("vm") or self.vm
        self.version = data.get("version", self.version)
        self.prompt_template = data.get("prompt_template") or self.prompt_template
        if not self.prompt_template and self.vm:
            try:
                load_version = self.version or self.vm.get_current_version()
                if load_version:
                    self.prompt_template = self.vm.load_version(load_version)
                    self.version = load_version
            except Exception:
                pass
        self.eval_data = data.get("eval_data", [])
        self.results = data.get("results")
        self.error_analysis = data.get("error_analysis")
        self.scenario_name = data.get("scenario_name", "")
        self.target_info = data.get("target_info")
        self.helper_info = data.get("helper_info")
        self.max_workers = data.get("max_workers", 8)
        self.requirements = data.get("requirements", "")
        self.context_fields = set(data.get("selected_context_fields") or data.get("context_fields", set()))
        self.context_field_specs = [
            spec for spec in (data.get("context_field_specs") or [])
            if spec.get("name") in self.context_fields
        ]

        self.prompt_editor.setPlainText(self.prompt_template)
        # Load manual suggestions: prefer persistent file, fall back to context
        saved_suggestions = self.vm.load_manual_suggestions() if self.vm else ""
        if saved_suggestions:
            self.manual_suggestions.setPlainText(saved_suggestions)
        elif "manual_suggestions" in data:
            self.manual_suggestions.setPlainText(data["manual_suggestions"])

        # Load version history
        if self.vm:
            versions = self.vm.list_versions()
            self.version_list.clear()
            for v in versions:
                item = QListWidgetItem(f"v{v}")
                item.setData(Qt.UserRole, v)
                if v == self.version:
                    item.setSelected(True)
                self.version_list.addItem(item)

        # Enable core buttons when context is ready
        has_errors = (self.error_analysis and self.error_analysis.get('errors'))
        has_manual = bool(self.manual_suggestions.toPlainText().strip())
        self.suggest_btn.setEnabled(self.error_analysis is not None)
        has_prompt = bool(self.prompt_template.strip())
        self.edit_prompt_btn.setEnabled(has_prompt)
        self.deploy_btn.setEnabled(has_prompt)
        if has_prompt:
            self.status_label.setText(f"已加载 v{self.version} 的 Prompt，可直接部署上线。")
            self.status_label.setStyleSheet("color: #1565C0;")
        elif not has_prompt:
            self.status_label.setText("未加载到 Prompt，请先回到 Step 3 生成/导入 Prompt，或在 Step 9 上传文件部署。")
            self.status_label.setStyleSheet("color: #C62828;")
        # 有错误或有手动建议时，都允许 LLM 自动优化
        if has_errors or has_manual:
            self.auto_optimize_btn.setEnabled(True)
            self.regen_data_btn.setEnabled(True)
            self.rerun_btn.setEnabled(True)
            self.save_version_btn.setEnabled(True)

        # 手动建议框内容变化时，联动启用/禁用自动优化按钮
        try:
            self.manual_suggestions.textChanged.disconnect()
        except Exception:
            pass
        self.manual_suggestions.textChanged.connect(self._on_manual_text_changed)

    def _on_manual_text_changed(self):
        """当手动建议框有内容时，启用 LLM 自动优化按钮。"""
        has_text = bool(self.manual_suggestions.toPlainText().strip())
        self.auto_optimize_btn.setEnabled(has_text)
        if has_text:
            self.regen_data_btn.setEnabled(True)
            self.rerun_btn.setEnabled(True)
            self.save_version_btn.setEnabled(True)

    def _get_suggestions(self):
        if not callable(self.helper_llm):
            QMessageBox.warning(self, "Helper 模型未加载", "请先回到 Step 1 选择并确认 Helper 模型。")
            return
        if not self.error_analysis:
            self.suggestions_display.setPlainText("尚未进行错误分析，请先在 Step 6 运行评测。")
            return

        if not self.error_analysis.get('errors'):
            self.suggestions_display.setPlainText(
                "✓ 所有样本全部通过，LLM 无自动优化建议。\n\n"
                "如有优化想法，请在下方「手动输入优化建议」框中输入，"
                "然后点击「LLM 自动优化」按钮。"
            )
            # 即使全对，也启用编辑和操作按钮
            self.edit_prompt_btn.setEnabled(True)
            self.save_version_btn.setEnabled(True)
            self.regen_data_btn.setEnabled(True)
            self.rerun_btn.setEnabled(True)
            self.deploy_btn.setEnabled(True)
            # 如果有手动建议，启用自动优化
            if self.manual_suggestions.toPlainText().strip():
                self.auto_optimize_btn.setEnabled(True)
            return

        self.suggest_btn.setEnabled(False)
        self.suggestions_display.setPlainText("正在生成优化建议...")
        self.status_label.setText("正在调用 Helper 模型分析...")
        self.status_label.setStyleSheet("color: #666;")

        self.worker = SuggestImprovementsWorker(
            self.helper_llm, self.prompt_template, self.error_analysis
        )
        self.worker.finished.connect(self._on_suggestions)
        self.worker.error.connect(lambda e: self.status_label.setText(f"建议生成失败: {e}"))
        self.worker.start()

    def _on_suggestions(self, suggestions):
        self.suggestions_display.setPlainText(suggestions)
        self.status_label.setText("✓ 优化建议已生成")
        self.status_label.setStyleSheet("color: green;")
        self.suggest_btn.setEnabled(True)
        self.auto_optimize_btn.setEnabled(True)
        self.edit_prompt_btn.setEnabled(True)
        self.save_version_btn.setEnabled(True)
        self.regen_data_btn.setEnabled(True)
        self.rerun_btn.setEnabled(True)
        self.deploy_btn.setEnabled(True)

    # ─── LLM auto-optimize: regenerate prompt from suggestions ──────────

    def _auto_optimize(self):
        """Use helper LLM to auto-regenerate prompt based on suggestions."""
        if not callable(self.helper_llm):
            QMessageBox.warning(self, "Helper 模型未加载", "请先回到 Step 1 选择并确认 Helper 模型。")
            return
        llm_suggestions = self.suggestions_display.toPlainText().strip()
        manual_suggestions = self.manual_suggestions.toPlainText().strip()

        if not llm_suggestions and not manual_suggestions:
            QMessageBox.warning(self, "提示", "请先点击「获取优化建议」生成 LLM 建议，或在手动建议区输入您的优化想法。")
            return

        combined_suggestions = ""
        if manual_suggestions:
            combined_suggestions += "# 【优先采用】手动优化建议\n" + manual_suggestions + "\n\n"
        if llm_suggestions:
            combined_suggestions += "# 【仅供参考】LLM 优化建议\n" + llm_suggestions

        error_details = ""
        if self.error_analysis:
            error_details = (
                f"总样本: {self.error_analysis['total']}, "
                f"准确率: {self.error_analysis.get('accuracy', 0):.2%}, "
                f"错误数: {len(self.error_analysis.get('errors', []))}"
            )
            patterns = self.error_analysis.get('error_patterns', {})
            pattern_desc = "; ".join(f"{k}={v}" for k, v in patterns.items() if v)
            if pattern_desc:
                error_details += "\n错误模式: " + pattern_desc

        self.auto_optimize_btn.setEnabled(False)
        self.auto_optimize_btn.setText("优化中...")
        self.status_label.setText("正在调用 Helper 模型生成优化后的 Prompt...")
        self.status_label.setStyleSheet("color: #666;")

        self._run_auto_optimize(combined_suggestions, error_details)

    def _field_schema(self):
        specs = [
            dict(spec) for spec in (self.context_field_specs or [])
            if spec.get("name") in set(self.context_fields or set())
        ]
        seen = {spec.get("name") for spec in specs}
        for name in sorted(set(self.context_fields or set()) - seen):
            specs.append({"name": name, "type": "string", "required": True})
        return {"input_fields": specs, "output_fields": []}

    def _unknown_placeholders_in(self, prompt: str):
        placeholders = set(extract_placeholders(prompt or ""))
        return sorted(placeholders - set(self.context_fields or set()))

    def _run_auto_optimize(self, suggestions_text: str, error_details: str):
        """Background worker for LLM auto-optimization."""
        from .workers import AutoOptimizeWorker
        field_schema = self._field_schema()
        self._opt_worker = AutoOptimizeWorker(
            self.helper_llm, self.prompt_template,
            error_details, suggestions_text,
            field_schema=field_schema,
        )
        self._opt_worker.finished.connect(self._on_auto_optimized)
        self._opt_worker.error.connect(self._on_auto_optimize_error)
        self._opt_worker.start()

    def _on_auto_optimized(self, new_prompt: str):
        """Handle the auto-optimized prompt result."""
        self.auto_optimize_btn.setEnabled(True)
        self.auto_optimize_btn.setText("LLM 自动优化")

        if not new_prompt.strip():
            self.status_label.setText("优化失败：生成的 Prompt 为空")
            self.status_label.setStyleSheet("color: red;")
            return

        # Update editor and auto-save as new version
        illegal = self._unknown_placeholders_in(new_prompt)
        if illegal:
            self.status_label.setText("优化结果含未勾选 Context 字段，已拒绝保存。")
            self.status_label.setStyleSheet("color: red;")
            QMessageBox.critical(
                self, "占位符不允许使用",
                "LLM 优化后的 Prompt 含有未勾选的 Context 字段，已拒绝保存。\n\n"
                f"不允许的占位符: {', '.join(illegal)}"
            )
            return

        self.prompt_editor.setPlainText(new_prompt)
        self.prompt_editor.setReadOnly(True)
        self.editor_title.setText("当前 Prompt (LLM 优化后):")
        self.editor_title.setStyleSheet("color: #1976D2; font-weight: bold;")
        self.prompt_template = new_prompt

        if self.vm:
            self.version = self.vm.save_version(new_prompt)
            self.vm.save_manual_suggestions(self.manual_suggestions.toPlainText())
            # Also regenerate eval script
            try:
                from code.pipeline.eval_code_generator import generate_eval_script
                script = generate_eval_script(
                    self.scenario_name, self.prompt_template, self.eval_data,
                    self.target_info, self.version, self.max_workers,
                    helper_info=self.helper_info
                )
                self.vm.save_eval_script(script)
            except Exception as e:
                pass

            # Refresh version list
            versions = self.vm.list_versions()
            self.version_list.clear()
            for v in versions:
                item = QListWidgetItem(f"v{v}")
                item.setData(Qt.UserRole, v)
                if v == self.version:
                    item.setSelected(True)
                self.version_list.addItem(item)

        self.status_label.setText(f"✓ LLM 优化完成，已保存为 v{self.version}")
        self.status_label.setStyleSheet("color: green; font-weight: bold;")
        self.rerun_btn.setEnabled(True)
        _notify_long_task_done(
            self, "LLM 自动优化完成",
            f"Prompt 已优化完成，并保存为 v{self.version}。",
            force=True,
        )

    def _on_auto_optimize_error(self, error_msg: str):
        self.auto_optimize_btn.setEnabled(True)
        self.auto_optimize_btn.setText("LLM 自动优化")
        self.status_label.setText(f"优化失败: {error_msg}")
        self.status_label.setStyleSheet("color: red;")

    def _edit_prompt(self):
        self.prompt_editor.setReadOnly(False)
        self.prompt_editor.setFocus()
        self.editor_title.setText("当前 Prompt (编辑中 - 修改后点击「保存为新版本」):")
        self.editor_title.setStyleSheet("color: #F57C00; font-weight: bold;")
        self.save_version_btn.setText("保存为新版本")
        self.status_label.setText("请修改 Prompt，修改完成后点击「保存为新版本」。")

    def _save_version(self):
        new_prompt = self.prompt_editor.toPlainText().strip()
        if not new_prompt:
            QMessageBox.warning(self, "警告", "Prompt 内容为空。")
            return

        if new_prompt == self.prompt_template:
            QMessageBox.information(self, "提示", "Prompt 未修改。")
            return

        illegal = self._unknown_placeholders_in(new_prompt)
        if illegal:
            QMessageBox.critical(
                self, "占位符不允许使用",
                "Prompt 含有未勾选的 Context 字段，不能保存。\n\n"
                f"不允许的占位符: {', '.join(illegal)}"
            )
            self.status_label.setText("Prompt 含有未勾选字段，已阻止保存。")
            self.status_label.setStyleSheet("color: red; font-weight: bold;")
            return

        self.prompt_template = new_prompt
        self.prompt_editor.setReadOnly(True)
        self.editor_title.setText("当前 Prompt:")
        self.editor_title.setStyleSheet("")

        if self.vm:
            self.version = self.vm.save_version(new_prompt)
            self.vm.save_manual_suggestions(self.manual_suggestions.toPlainText())
            self.status_label.setText(f"✓ 已保存为 v{self.version}（含手动建议）")
            self.status_label.setStyleSheet("color: green;")

            # Refresh version list
            versions = self.vm.list_versions()
            self.version_list.clear()
            for v in versions:
                item = QListWidgetItem(f"v{v}")
                item.setData(Qt.UserRole, v)
                if v == self.version:
                    item.setSelected(True)
                self.version_list.addItem(item)

            # Also regenerate eval script
            try:
                script = generate_eval_script(
                    self.scenario_name, self.prompt_template, self.eval_data,
                    self.target_info, self.version, self.max_workers,
                    helper_info=self.helper_info
                )
                self.vm.save_eval_script(script)
            except Exception as e:
                self.status_label.setText(f"已保存为 v{self.version}，但评测脚本生成失败: {e}")

        self.rerun_btn.setEnabled(True)
        self.save_version_btn.setText("保存为新版本")

    def _load_version(self, item):
        v = item.data(Qt.UserRole)
        if self.vm and v:
            try:
                prompt = self.vm.load_version(v)
                self.prompt_editor.setPlainText(prompt)
                self.prompt_editor.setReadOnly(True)
                self.version = v
                self.prompt_template = prompt
                self.status_label.setText(f"已加载 v{v} 的 Prompt")
                self.status_label.setStyleSheet("color: #1565C0;")
            except Exception as e:
                QMessageBox.warning(self, "加载失败", str(e))

    def _regen_data(self):
        """Regenerate eval data with current optimized prompt."""
        parent = self.parent()
        while parent and not hasattr(parent, 'go_to_step'):
            parent = parent.parent()
        if parent:
            # Read latest from editor (user may have edited without saving)
            editor_prompt = self.prompt_editor.toPlainText().strip()
            if editor_prompt and editor_prompt != self.prompt_template:
                illegal = self._unknown_placeholders_in(editor_prompt)
                if illegal:
                    QMessageBox.critical(
                        self, "占位符不允许使用",
                        "Prompt 含有未勾选的 Context 字段，不能用于重新生成评测数据。\n\n"
                        f"不允许的占位符: {', '.join(illegal)}"
                    )
                    return
                # Auto-save if editor content differs from last saved version
                self.prompt_template = editor_prompt
                if self.vm:
                    self.version = self.vm.save_version(editor_prompt)
                    self.status_label.setText(f"已自动保存为 v{self.version}")
                    self.status_label.setStyleSheet("color: #F57C00;")
            elif self.vm:
                self.vm.save_manual_suggestions(self.manual_suggestions.toPlainText())

            # Carry the optimized prompt to eval data step
            parent.go_to_step(4, {
                "prompt_template": self.prompt_template,
                "version": self.version,
                "eval_data": self.eval_data,
                "requirements": self.requirements,
            })
            # Ensure eval data panel receives the new prompt
            parent.eval_data_panel.set_context(parent.context)
            self.status_label.setText(f"已跳转到评测数据步骤，使用 v{self.version} 的 Prompt")
            self.status_label.setStyleSheet("color: #1565C0;")

    def _rerun(self):
        """Re-run evaluation with current prompt (jump to Step 6)."""
        parent = self.parent()
        while parent and not hasattr(parent, 'go_to_step'):
            parent = parent.parent()
        if parent:
            # Read latest from editor (user may have edited without saving)
            editor_prompt = self.prompt_editor.toPlainText().strip()
            if editor_prompt and editor_prompt != self.prompt_template:
                illegal = self._unknown_placeholders_in(editor_prompt)
                if illegal:
                    QMessageBox.critical(
                        self, "占位符不允许使用",
                        "Prompt 含有未勾选的 Context 字段，不能用于重新评测。\n\n"
                        f"不允许的占位符: {', '.join(illegal)}"
                    )
                    return
                self.prompt_template = editor_prompt
                if self.vm:
                    self.version = self.vm.save_version(editor_prompt)
                    # Also regenerate eval script
                    try:
                        from code.pipeline.eval_code_generator import generate_eval_script
                        script = generate_eval_script(
                            self.scenario_name, self.prompt_template, self.eval_data,
                            self.target_info, self.version, self.max_workers,
                            helper_info=self.helper_info
                        )
                        self.vm.save_eval_script(script)
                    except Exception:
                        pass
            self.status_label.setText(f"跳转到评测步骤，使用 v{self.version} 的 Prompt...")
            self.status_label.setStyleSheet("color: #1565C0;")
            parent.go_to_step(6, {
                "prompt_template": self.prompt_template,
                "version": self.version,
                "eval_data": self.eval_data,
                "context_fields": set(self.context_fields or set()),
                "selected_context_fields": set(self.context_fields or set()),
                "context_field_specs": list(self.context_field_specs or []),
            })

    def _go_deploy(self):
        """Navigate to Step 9 (Deploy) with current prompt."""
        parent = self.parent()
        while parent and not hasattr(parent, 'go_to_step'):
            parent = parent.parent()
        if parent:
            editor_prompt = self.prompt_editor.toPlainText().strip()
            if editor_prompt:
                self.prompt_template = editor_prompt
            elif not self.prompt_template and self.vm:
                try:
                    load_version = self.version or self.vm.get_current_version()
                    if load_version:
                        self.prompt_template = self.vm.load_version(load_version)
                        self.version = load_version
                        self.prompt_editor.setPlainText(self.prompt_template)
                except Exception:
                    pass
            if not self.prompt_template.strip():
                QMessageBox.warning(
                    self, "无法部署",
                    "当前没有可部署的 Prompt。\n\n"
                    "请先回到 Step 3 生成/导入 Prompt，或进入 Step 9 上传 Prompt 文件。"
                )
                return
            parent.go_to_step(9, {
                "prompt_template": self.prompt_template,
                "version": self.version,
                "scenario_name": self.scenario_name,
                "requirements": self.requirements,
                "eval_data": self.eval_data,
                "vm": self.vm,
            })


# ─── Step 9: Deploy Panel ───────────────────────────────────────────────────

class DeployPanel(QWidget):
    """Step 9: Convert, validate & deploy prompt to LLM Gateway."""

    STORAGE_KEY = "streambridge_config"

    def __init__(self, parent=None):
        super().__init__(parent)
        self.prompt_template = ""
        self.go_template = ""
        self.scenario_name = ""
        self.version = 1
        self.requirements = ""
        self.eval_data = []
        self.helper_llm = None
        self.deployed_template_id = None
        self.vm = None
        self._one_click_deploying = False
        self.uploaded_prompt_name = ""
        self._setup_ui()
        self._load_config()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        # ── Title ──
        title = QLabel("Step 9: 上线部署 (LLM Gateway)")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #1976D2;")
        layout.addWidget(title)

        # ── Server config ──
        cfg_group = QGroupBox("LLM Gateway 连接配置")
        cfg_grid = QGridLayout(cfg_group)
        cfg_grid.addWidget(QLabel("网关地址:"), 0, 0)
        self.base_url_edit = QLineEdit(os.getenv("GATEWAY_BASE_URL") or os.getenv("STREAMBRIDGE_BASE_URL", "http://127.0.0.1:8080"))
        cfg_grid.addWidget(self.base_url_edit, 0, 1)
        cfg_grid.addWidget(QLabel("API Key:"), 1, 0)
        self.api_key_edit = QLineEdit()
        self.api_key_edit.setEchoMode(QLineEdit.Password)
        cfg_grid.addWidget(self.api_key_edit, 1, 1)
        self.show_key_cb = QCheckBox("显示")
        self.show_key_cb.toggled.connect(
            lambda checked: self.api_key_edit.setEchoMode(
                QLineEdit.Normal if checked else QLineEdit.Password)
        )
        cfg_grid.addWidget(self.show_key_cb, 1, 2)
        cfg_grid.addWidget(QLabel("命名空间:"), 2, 0)
        self.namespace_edit = QLineEdit("default")
        cfg_grid.addWidget(self.namespace_edit, 2, 1)
        layout.addWidget(cfg_group)

        # ── Version selection ──
        ver_widget = QWidget()
        ver_layout = QHBoxLayout(ver_widget)
        ver_layout.setContentsMargins(0, 0, 0, 0)
        ver_layout.addWidget(QLabel("选择 Prompt 版本:"))
        self.version_combo = QComboBox()
        self.version_combo.setMinimumWidth(200)
        self.version_combo.currentIndexChanged.connect(self._on_version_selected)
        ver_layout.addWidget(self.version_combo)
        self.upload_prompt_btn = QPushButton("上传 Prompt 文件")
        self.upload_prompt_btn.clicked.connect(self._upload_prompt_file)
        ver_layout.addWidget(self.upload_prompt_btn)
        self.deploy_btn = QPushButton("部署上线")
        self.deploy_btn.setEnabled(False)
        self.deploy_btn.setStyleSheet(
            "QPushButton { background: #388E3C; color: white; padding: 8px 24px;"
            " border-radius: 4px; font-size: 14px; }"
            "QPushButton:hover { background: #2E7D32; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.deploy_btn.clicked.connect(self._deploy_online)
        ver_layout.addWidget(self.deploy_btn)
        ver_layout.addStretch()
        layout.addWidget(ver_widget)

        # ── Template editors ──
        splitter = QSplitter(Qt.Vertical)

        # Python template (read-only reference)
        py_widget = QWidget()
        py_layout = QVBoxLayout(py_widget)
        py_layout.setContentsMargins(0, 0, 0, 0)
        py_layout.addWidget(QLabel("原始 Prompt（双括号字段语法）:"))
        self.py_display = QPlainTextEdit()
        self.py_display.setReadOnly(True)
        self.py_display.setFont(QFont("Consolas", 10))
        self.py_display.setMinimumHeight(170)
        _enable_width_wrap(self.py_display)
        py_layout.addWidget(self.py_display)
        splitter.addWidget(py_widget)

        # Go template (editable after background conversion)
        go_widget = QWidget()
        go_layout = QVBoxLayout(go_widget)
        go_layout.setContentsMargins(0, 0, 0, 0)
        self.go_title = QLabel("后台转换后的 Go 模板 (可编辑，用于排查或微调):")
        go_layout.addWidget(self.go_title)
        self.go_editor = QPlainTextEdit()
        self.go_editor.setFont(QFont("Consolas", 10))
        self.go_editor.setMinimumHeight(170)
        self.go_editor.setPlaceholderText("点击「部署上线」后，系统会在后台自动转换并显示 Go 语法模板...")
        _enable_width_wrap(self.go_editor)
        go_layout.addWidget(self.go_editor)
        splitter.addWidget(go_widget)

        splitter.setSizes([220, 220])
        layout.addWidget(splitter)

        # ── Placeholders / expected vars ──
        ph_widget = QWidget()
        ph_layout = QHBoxLayout(ph_widget)
        ph_layout.setContentsMargins(0, 0, 0, 0)
        ph_layout.addWidget(QLabel("模板变量:"))
        self.var_list = QListWidget()
        self.var_list.setMinimumHeight(90)
        self.var_list.setMaximumHeight(130)
        ph_layout.addWidget(self.var_list)
        self.add_expected_var_btn = QPushButton("+")
        self.add_expected_var_btn.setFixedWidth(30)
        self.add_expected_var_btn.setToolTip("将选中变量加入验证清单")
        ph_layout.addWidget(self.add_expected_var_btn)
        layout.addWidget(ph_widget)

        # ── Action buttons ──
        btn_layout = QHBoxLayout()
        self.convert_btn = QPushButton("转换模板")
        self.convert_btn.setStyleSheet(
            "QPushButton { background: #1976D2; color: white; padding: 8px 20px;"
            " border-radius: 4px; font-size: 13px; }"
        )
        self.convert_btn.clicked.connect(self._convert_template)
        self.convert_btn.setVisible(False)
        btn_layout.addWidget(self.convert_btn)

        self.validate_btn = QPushButton("校验模板")
        self.validate_btn.setEnabled(False)
        self.validate_btn.setStyleSheet(
            "QPushButton { background: #F57C00; color: white; padding: 8px 20px;"
            " border-radius: 4px; font-size: 13px; }"
            "QPushButton:disabled { background: #ccc; }"
        )
        self.validate_btn.clicked.connect(self._validate_template)
        self.validate_btn.setVisible(False)
        btn_layout.addWidget(self.validate_btn)

        self.back_btn = QPushButton("返回迭代优化")
        self.back_btn.clicked.connect(self._go_back)
        btn_layout.addWidget(self.back_btn)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # ── Status / Log ──
        self.status_label = QLabel("")
        self.status_label.setWordWrap(True)
        layout.addWidget(self.status_label)

        result_group = QGroupBox("上线结果")
        result_layout = QVBoxLayout(result_group)
        self.deploy_result_display = QTextEdit()
        self.deploy_result_display.setReadOnly(True)
        _enable_width_wrap(self.deploy_result_display)
        self.deploy_result_display.setMinimumHeight(125)
        self.deploy_result_display.setMaximumHeight(170)
        self.deploy_result_display.setPlaceholderText("部署成功后会在这里显示 template_id、模板名、命名空间和版本。")
        result_layout.addWidget(self.deploy_result_display)
        layout.addWidget(result_group)

        self.log_output = QTextEdit()
        self.log_output.setReadOnly(True)
        _enable_width_wrap(self.log_output)
        self.log_output.setFont(QFont("Consolas", 9))
        self.log_output.setStyleSheet("background: #1e1e1e; color: #d4d4d4;")
        self.log_output.setMinimumHeight(150)
        self.log_output.setMaximumHeight(210)
        self.log_output.setPlaceholderText("操作日志...")
        layout.addWidget(self.log_output)

        # ── Deploy history ──
        self.deploy_history_group = QGroupBox("部署历史")
        hist_layout = QVBoxLayout(self.deploy_history_group)
        self.deploy_history_list = QListWidget()
        self.deploy_history_list.setMinimumHeight(130)
        self.deploy_history_list.setMaximumHeight(180)
        self.deploy_history_list.setStyleSheet(
            "QListWidget { font-size: 12px; background: #f9f9f9; }"
            "QListWidget::item { padding: 3px 6px; }"
        )
        hist_layout.addWidget(self.deploy_history_list)
        layout.addWidget(self.deploy_history_group)

    # ── Context ────────────────────────────────────────────────────────

    def set_context(self, data: dict):
        self.helper_llm = data.get("helper_llm")
        self.prompt_template = strip_auto_output_schema(data.get("prompt_template", ""))
        self.scenario_name = data.get("scenario_name", "")
        self.version = data.get("version", 1)
        self.requirements = data.get("requirements", "")
        self.eval_data = data.get("eval_data", [])

        self.py_display.setPlainText(self.prompt_template)
        self.uploaded_prompt_name = ""

        # Extract placeholders
        from code.pipeline.prompt_generator import extract_placeholders
        phs = extract_placeholders(self.prompt_template)
        self.var_list.clear()
        for ph in phs:
            self.var_list.addItem(ph)

        # Reset deploy state
        self.go_template = ""
        self.go_editor.clear()
        self.deployed_template_id = None
        self.deploy_result_display.clear()
        self._set_buttons(False, False, False)
        self.deploy_btn.setEnabled(bool(self.prompt_template.strip()))
        self.deploy_btn.setText("部署上线")
        self.status_label.setText("已加载 Prompt。选择版本后点击「部署上线」，系统会后台完成转换、校验和部署。")

        # Load version manager and populate version list
        self.vm = data.get("vm")
        self._refresh_version_list()

        # Load deploy history
        self._refresh_deploy_history()

    def _set_buttons(self, converted, validated, deployed):
        self.validate_btn.setEnabled(converted)
        self.deploy_btn.setEnabled(validated)
        if deployed:
            self.deploy_btn.setText("✓ 已部署")
            self.deploy_btn.setStyleSheet(
                "QPushButton { background: #388E3C; color: white; padding: 8px 24px;"
                " border-radius: 4px; font-size: 14px; }"
            )

    def _refresh_version_list(self):
        """Populate the version combo box from VersionManager."""
        # 先加载部署历史，用于标注已部署版本
        deploy_records = self._load_deploy_history()
        deployed_versions = set(r["version"] for r in deploy_records)

        self.version_combo.blockSignals(True)
        self.version_combo.clear()
        self.version_combo.addItem("— 从 Step 8 传入 —", None)  # placeholder

        if self.vm:
            try:
                versions = self.vm.list_versions()
                for v in versions:
                    label = f"v{v}"
                    if v in deployed_versions:
                        label += "  ✓ 已部署"
                    self.version_combo.addItem(label, v)
                    if v == self.version:
                        self.version_combo.setCurrentIndex(self.version_combo.count() - 1)
            except Exception:
                pass
        self.version_combo.blockSignals(False)

    def _on_version_selected(self, idx):
        """Load the selected prompt version into the editor."""
        if idx < 0:
            return
        v = self.version_combo.itemData(idx)
        if v is None:
            # User selected the placeholder — do nothing
            return
        try:
            prompt = strip_auto_output_schema(self.vm.load_version(v))
            self.prompt_template = prompt
            self.version = v
            self.uploaded_prompt_name = ""
            self.py_display.setPlainText(prompt)

            # Reset deploy state
            self.go_template = ""
            self.go_editor.clear()
            self.deployed_template_id = None
            self._set_buttons(False, False, False)
            self.deploy_btn.setEnabled(bool(prompt.strip()))
            self.deploy_btn.setText("部署上线")
            self.deploy_result_display.clear()

            self.status_label.setText(f"已加载 v{v} 的 Prompt，点击「部署上线」开始后台转换、校验和部署。")
            self.status_label.setStyleSheet("color: #1565C0;")
            self._log(f"已加载 v{v} 的 Prompt ({len(prompt)} 字符)")
        except Exception as e:
            QMessageBox.warning(self, "加载失败", str(e))

    def _log(self, msg: str):
        ts = datetime.now().strftime("%H:%M:%S")
        self.log_output.append(f"[{ts}] {msg}")

    # ── Config persistence ─────────────────────────────────────────────

    def _save_config(self):
        """Save LLM Gateway connection config to JSON file."""
        config = {
            "base_url": self.base_url_edit.text().strip(),
            "api_key": self.api_key_edit.text().strip(),
            "namespace": self.namespace_edit.text().strip(),
        }
        try:
            config_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", ".deploy_config")
            os.makedirs(config_dir, exist_ok=True)
            with open(os.path.join(config_dir, f"{self.STORAGE_KEY}.json"), "w") as f:
                json.dump(config, f)
        except Exception:
            pass

    def _load_config(self):
        """Load LLM Gateway connection config from JSON file."""
        config_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "..", ".deploy_config",
            f"{self.STORAGE_KEY}.json"
        )
        try:
            with open(config_path) as f:
                config = json.load(f)
            self.base_url_edit.setText(config.get("base_url", self.base_url_edit.text()))
            self.api_key_edit.setText(config.get("api_key", "") or self._default_api_key())
            self.namespace_edit.setText(config.get("namespace", "default"))
        except (FileNotFoundError, json.JSONDecodeError):
            self.api_key_edit.setText(self._default_api_key())

    def _default_api_key(self):
        """Read the existing project API key so deploy can be used without retyping it."""
        candidates = [
            os.path.join(PROJECT_ROOT, "code", "workflow", "llm_key.py"),
            os.path.join(PROJECT_ROOT, "code", "models", "api_keys.py"),
        ]
        patterns = [
            r'api_key\s*=\s*["\']([^"\']+)["\']',
            r'OPENAI_API_KEY\s*=\s*["\']([^"\']+)["\']',
        ]
        for path in candidates:
            if not os.path.exists(path):
                continue
            try:
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
                for pattern in patterns:
                    match = re.search(pattern, text)
                    if match and match.group(1).strip():
                        return match.group(1).strip()
            except Exception:
                continue
        return ""

    def _upload_prompt_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "选择 Prompt 文件",
            os.path.join(PROJECT_ROOT, "data", "prompt"),
            "Prompt Files (*.py *.txt *.md);;All Files (*)",
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw = f.read()
        except UnicodeDecodeError:
            with open(path, "r", encoding="gbk", errors="replace") as f:
                raw = f.read()
        except Exception as exc:
            QMessageBox.warning(self, "读取失败", str(exc))
            return

        prompt = self._extract_prompt_from_file(raw)
        if not prompt.strip():
            QMessageBox.warning(self, "提示", "文件中没有可部署的 Prompt 内容。")
            return

        stem = os.path.splitext(os.path.basename(path))[0]
        self.uploaded_prompt_name = stem
        self.scenario_name = self.scenario_name or stem
        self.version = 0
        self.prompt_template = strip_auto_output_schema(prompt.strip())
        self.py_display.setPlainText(self.prompt_template)
        self.go_template = ""
        self.go_editor.clear()
        self.deployed_template_id = None
        self.deploy_result_display.clear()
        self.deploy_btn.setEnabled(True)
        self.deploy_btn.setText("部署上线")

        from code.pipeline.prompt_generator import extract_placeholders
        self.var_list.clear()
        for ph in extract_placeholders(self.prompt_template):
            self.var_list.addItem(ph)

        self.status_label.setText(f"已上传 Prompt 文件：{os.path.basename(path)}。点击「部署上线」即可部署。")
        self.status_label.setStyleSheet("color: #1565C0;")
        self._log(f"已上传 Prompt 文件: {path}")

    def _extract_prompt_from_file(self, raw: str) -> str:
        match = re.search(
            r'PROMPT_TEMPLATE\s*=\s*(?P<quote>"""|\'\'\')(?P<body>.*?)(?P=quote)',
            raw,
            flags=re.S,
        )
        if match:
            return match.group("body")
        match = re.search(
            r'PROMPT_TEMPLATE\s*=\s*(?P<quote>["\'])(?P<body>.*?)(?P=quote)',
            raw,
            flags=re.S,
        )
        if match:
            return match.group("body")
        return raw

    def _deploy_template_name(self):
        return normalize_scenario_name(
            self.uploaded_prompt_name or self.scenario_name or "uploaded_prompt"
        )

    def _deploy_description(self):
        source = "上传文件" if self.uploaded_prompt_name else f"Prompt v{self.version}"
        return f"{source} via Prompt 自动化调试与评测系统"

    # ── Step 1: Convert ────────────────────────────────────────────────

    def _convert_template(self):
        if not self.prompt_template.strip():
            QMessageBox.warning(self, "提示", "Prompt 模板为空。")
            return

        if self._one_click_deploying or not callable(self.helper_llm):
            self.convert_btn.setEnabled(False)
            self.convert_btn.setText("转换中...")
            self.status_label.setText("正在使用本地规则转换模板语法...")
            self.status_label.setStyleSheet("color: #666;")
            self._log("使用本地规则转换 {{field}} → {{.field}}。")
            try:
                go_template = convert_python_to_go_template(self.prompt_template)
                self._on_converted(go_template, self.prompt_template)
            except Exception as exc:
                self._on_convert_error(str(exc))
            return

        self.convert_btn.setEnabled(False)
        self.convert_btn.setText("转换中...")
        self.status_label.setText("正在通过 LLM 转换模板语法...")
        self.status_label.setStyleSheet("color: #666;")
        self._log("开始 Python → Go 模板语法转换...")

        self._conv_worker = ConvertTemplateWorker(self.helper_llm, self.prompt_template)
        self._conv_worker.finished.connect(self._on_converted)
        self._conv_worker.error.connect(self._on_convert_error)
        self._conv_worker.start()

    def _deploy_online(self):
        """One-click deploy: convert template, validate it, then deploy to LLM Gateway."""
        base_url = self.base_url_edit.text().strip()
        api_key = self.api_key_edit.text().strip()
        name = self._deploy_template_name()

        if not base_url or not api_key:
            QMessageBox.warning(self, "提示", "请先配置网关地址和 API Key。")
            return
        if not self.prompt_template.strip():
            QMessageBox.warning(self, "提示", "Prompt 模板为空，请先选择一个 Prompt 版本。")
            return

        self._save_config()
        self._one_click_deploying = True
        self.deploy_result_display.clear()
        self.go_editor.clear()
        self.go_template = ""
        self.deployed_template_id = None
        self.convert_btn.setEnabled(False)
        self.validate_btn.setEnabled(False)
        self.deploy_btn.setEnabled(False)
        self.deploy_btn.setText("部署中...")
        self.status_label.setText("正在后台转换模板，随后会自动校验并部署上线...")
        self.status_label.setStyleSheet("color: #666;")
        self._log(f"一键部署开始：v{self.version} / {name}")
        self._convert_template()

    def _on_converted(self, go_template, original):
        self.go_template = go_template
        self.go_editor.setPlainText(go_template)
        self._log("✓ 模板转换完成")
        self.status_label.setText("✓ 转换完成，请检查 Go 模板后点击「校验模板」")
        self.status_label.setStyleSheet("color: green;")
        self.convert_btn.setEnabled(True)
        self.convert_btn.setText("转换模板")
        self._set_buttons(converted=True, validated=False, deployed=False)

        # Update variable list for Go template
        from code.pipeline.prompt_generator import extract_placeholders
        # Show both Python and Go style vars
        go_vars = re.findall(r'\.(\w+)}', go_template)
        self.var_list.clear()
        for v in sorted(set(go_vars)):
            self.var_list.addItem(f"(.{v})")

        if self._one_click_deploying:
            self.status_label.setText("✓ 模板转换完成，正在自动校验...")
            self._validate_template(auto=True)

    def _on_convert_error(self, msg):
        self._log(f"ERROR: {msg}")
        self.status_label.setText(f"转换失败: {msg}")
        self.status_label.setStyleSheet("color: red;")
        self.convert_btn.setEnabled(True)
        self.convert_btn.setText("转换模板")
        self.deploy_btn.setEnabled(True)
        self.deploy_btn.setText("部署上线")
        self._one_click_deploying = False

    # ── Step 2: Validate ───────────────────────────────────────────────

    def _validate_template(self, auto=False):
        base_url = self.base_url_edit.text().strip()
        api_key = self.api_key_edit.text().strip()
        go_content = self.go_editor.toPlainText().strip()

        if not base_url:
            QMessageBox.warning(self, "提示", "请输入 LLM Gateway 网关地址。")
            return
        if not api_key:
            QMessageBox.warning(self, "提示", "请输入 API Key。")
            return
        if not go_content:
            QMessageBox.warning(self, "提示", "Go 模板内容为空。")
            return

        self.validate_btn.setEnabled(False)
        self.validate_btn.setText("校验中...")
        self.status_label.setText("正在连接 LLM Gateway 校验模板...")
        self.status_label.setStyleSheet("color: #666;")
        self._log("开始模板验证...")

        # Save config on first validate
        self._save_config()

        # Extract expected vars from Go template
        expected_vars = list(set(re.findall(r'\.(\w+)}', go_content)))

        self._val_worker = GatewayValidateWorker(
            base_url, api_key, go_content, expected_vars
        )
        self._val_worker.finished.connect(self._on_validated)
        self._val_worker.error.connect(self._on_validate_error)
        self._val_worker.start()

    def _on_validated(self, data):
        http_status = data.get("_http_status", 0)
        success = data.get("_success", False)

        # Pretty-print the API response
        resp_text = json.dumps(
            {k: v for k, v in data.items() if not k.startswith("_")},
            ensure_ascii=False, indent=2
        )

        if success and http_status < 300:
            self._log(f"✓ 验证通过 (HTTP {http_status})")
            self._log(f"响应: {resp_text}")
            self.status_label.setText("✓ 模板语法验证通过，可以部署")
            self.status_label.setStyleSheet("color: green; font-weight: bold;")
            self._set_buttons(converted=True, validated=True, deployed=False)
            if self._one_click_deploying or auto:
                self.status_label.setText("✓ 模板校验通过，正在自动部署上线...")
                self._deploy(auto=True)
        else:
            self._log(f"✗ 验证失败 (HTTP {http_status})")
            self._log(f"响应: {resp_text}")
            self.status_label.setText(f"验证失败 (HTTP {http_status})，请修改后重试")
            self.status_label.setStyleSheet("color: red;")
            self._set_buttons(converted=True, validated=False, deployed=False)
            self.deploy_btn.setEnabled(True)
            self.deploy_btn.setText("部署上线")
            self._one_click_deploying = False

        self.validate_btn.setEnabled(True)
        self.validate_btn.setText("校验模板")

    def _on_validate_error(self, msg):
        self._log(f"ERROR: {msg}")
        self.status_label.setText(f"验证异常: {msg}")
        self.status_label.setStyleSheet("color: red;")
        self.validate_btn.setEnabled(True)
        self.validate_btn.setText("校验模板")
        self.deploy_btn.setEnabled(True)
        self.deploy_btn.setText("部署上线")
        self._one_click_deploying = False

    # ── Step 3: Deploy ─────────────────────────────────────────────────

    def _deploy(self, auto=False):
        base_url = self.base_url_edit.text().strip()
        api_key = self.api_key_edit.text().strip()
        namespace = self.namespace_edit.text().strip() or "default"
        name = self._deploy_template_name()
        go_content = self.go_editor.toPlainText().strip()
        desc = self._deploy_description()

        if not base_url or not api_key:
            QMessageBox.warning(self, "提示", "请先配置网关地址和 API Key。")
            return
        if not go_content:
            QMessageBox.warning(self, "提示", "Go 模板内容为空。")
            return

        if not auto:
            reply = QMessageBox.question(
                self, "确认部署",
                f"即将部署模板到 LLM Gateway:\n\n"
                f"网关: {base_url}\n"
                f"命名空间: {namespace}\n"
                f"模板名称: {name}\n\n"
                f"确认继续？",
                QMessageBox.Yes | QMessageBox.No,
            )
            if reply != QMessageBox.Yes:
                return

        self.deploy_btn.setEnabled(False)
        self.deploy_btn.setText("部署中...")
        self.status_label.setText("正在部署到 LLM Gateway...")
        self.status_label.setStyleSheet("color: #666;")
        self._log(f"开始部署模板 '{name}' 到命名空间 '{namespace}'...")

        self._dep_worker = GatewayDeployWorker(
            base_url, api_key, namespace, name, go_content,
            description=desc, is_active=True,
        )
        self._dep_worker.finished.connect(self._on_deployed)
        self._dep_worker.error.connect(self._on_deploy_error)
        self._dep_worker.start()

    def _on_deployed(self, data):
        http_status = data.get("_http_status", 0)
        success = data.get("_success", False)
        resp_text = json.dumps(
            {k: v for k, v in data.items() if not k.startswith("_")},
            ensure_ascii=False, indent=2
        )

        if success and http_status < 300:
            # Extract template_id from response if available
            tid = self._extract_template_id(data)
            if tid:
                self.deployed_template_id = tid
            self._log(f"✓ 部署成功 (HTTP {http_status})")
            self._log(f"响应: {resp_text}")
            self.status_label.setText(f"✓ 部署成功！模板 ID: {tid or '未知'}")
            self.status_label.setStyleSheet("color: green; font-weight: bold;")
            self._set_buttons(converted=True, validated=True, deployed=True)
            result_text = (
                f"template_id: {tid or '未知'}\n"
                f"namespace: {self.namespace_edit.text().strip() or 'default'}\n"
                f"template_name: {self._deploy_template_name()}\n"
                f"prompt_version: {'上传文件' if self.uploaded_prompt_name else 'v' + str(self.version)}\n"
                f"base_url: {self.base_url_edit.text().strip()}"
            )
            self.deploy_result_display.setPlainText(result_text)
            # 保存部署记录
            self._save_deploy_record(tid)
            self._refresh_deploy_history()
            QMessageBox.information(
                self, "部署成功",
                f"模板已成功部署到 LLM Gateway！\n\n"
                f"命名空间: {self.namespace_edit.text().strip() or 'default'}\n"
                f"模板名: {self._deploy_template_name()}\n"
                f"版本: {'上传文件' if self.uploaded_prompt_name else 'v' + str(self.version)}\n"
                f"{'模板 ID: ' + tid if tid else ''}\n\n"
                f"部署记录已保存到 data/prompt/{self.scenario_name}/.deploy_log.json"
            )
        else:
            err_msg = data.get("message", data.get("error", str(data)))
            self._log(f"✗ 部署失败 (HTTP {http_status})")
            self._log(f"响应: {resp_text}")
            self.status_label.setText(f"部署失败: {err_msg}")
            self.status_label.setStyleSheet("color: red;")
            self.deploy_btn.setEnabled(True)
            self.deploy_btn.setText("部署上线")
        self._one_click_deploying = False

    def _on_deploy_error(self, msg):
        self._log(f"ERROR: {msg}")
        self.status_label.setText(f"部署异常: {msg}")
        self.status_label.setStyleSheet("color: red;")
        self.deploy_btn.setEnabled(True)
        self.deploy_btn.setText("部署上线")
        self._one_click_deploying = False

    def _extract_template_id(self, data):
        """Best-effort extraction for different LLM Gateway response shapes."""
        if not isinstance(data, dict):
            return ""
        for key in ("template_id", "templateId", "id"):
            value = data.get(key)
            if value:
                return str(value)
        nested = data.get("data")
        if isinstance(nested, dict):
            for key in ("template_id", "templateId", "id"):
                value = nested.get(key)
                if value:
                    return str(value)
            for value in nested.values():
                if isinstance(value, dict):
                    found = self._extract_template_id(value)
                    if found:
                        return found
        return ""

    # ── Deploy history tracking ──────────────────────────────────────────

    def _deploy_log_path(self):
        if not self.vm:
            return None
        return os.path.join(self.vm.prompt_dir, ".deploy_log.json")

    def _save_deploy_record(self, template_id):
        log_path = self._deploy_log_path()
        if not log_path:
            return
        records = []
        if os.path.exists(log_path):
            try:
                with open(log_path, "r", encoding="utf-8") as f:
                    records = json.load(f)
            except Exception:
                records = []
        records.append({
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "version": self.version,
            "template_id": template_id,
            "template_name": self._deploy_template_name(),
            "namespace": self.namespace_edit.text().strip() or "default",
            "base_url": self.base_url_edit.text().strip(),
        })
        with open(log_path, "w", encoding="utf-8") as f:
            json.dump(records, f, ensure_ascii=False, indent=2)
        self._log(f"部署记录已保存: {log_path}")

    def _load_deploy_history(self):
        log_path = self._deploy_log_path()
        if not log_path or not os.path.exists(log_path):
            return []
        try:
            with open(log_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []

    def _refresh_deploy_history(self):
        records = self._load_deploy_history()
        self.deploy_history_list.clear()
        if not records:
            self.deploy_history_list.addItem("（暂无部署记录）")
            self.deploy_history_group.setTitle("部署历史（无记录）")
        else:
            self.deploy_history_group.setTitle(f"部署历史（共 {len(records)} 次）")
            for r in reversed(records):
                label = (f"v{r['version']} → {r.get('template_name', '?')}  "
                         f"[{r['timestamp']}]  ID: {r.get('template_id', '?')}")
                self.deploy_history_list.addItem(label)
            # 显示最近一次部署的版本
            last = records[-1]
            self.deploy_history_list.addItem("")
            self.deploy_history_list.addItem(
                f"▶ 最近部署: v{last['version']} ({last['timestamp']})"
            )

    def _go_back(self):
        """Navigate back to Step 8 (Optimize)."""
        parent = self.parent()
        while parent and not hasattr(parent, 'go_to_step'):
            parent = parent.parent()
        if parent:
            parent.go_to_step(8, {
                "prompt_template": self.prompt_template,
                "version": self.version,
            })


# ─── Welcome Panel ──────────────────────────────────────────────────────────

class WelcomePanel(QWidget):
    """Welcome / start panel."""
    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignCenter)

        title = QLabel("Prompt 自动化调试与评测系统")
        title.setStyleSheet("font-size: 28px; font-weight: bold; color: #1976D2;")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        subtitle = QLabel("Virtual Coach - Prompt Engineering Pipeline")
        subtitle.setStyleSheet("font-size: 16px; color: #666;")
        subtitle.setAlignment(Qt.AlignCenter)
        layout.addWidget(subtitle)

        layout.addSpacing(30)

        workflow = QLabel(
            "完整工作流程:\n\n"
            "① 模型选择  →  ② 需求输入  →  ③ 生成 Prompt\n"
            "④ 评测数据  →  ⑤ 评测脚本  →  ⑥ 运行评测\n"
            "⑦ 结果分析  →  ⑧ 迭代优化  →  ⑨ 上线部署\n\n"
            "点击左侧「1. 模型选择」开始"
        )
        workflow.setStyleSheet("font-size: 14px; color: #333; padding: 20px;")
        workflow.setAlignment(Qt.AlignCenter)
        layout.addWidget(workflow)

        layout.addSpacing(20)

        tips = QLabel(
            "使用提示:\n"
            "• 先配置 API Key (code/models/api_keys.py)\n"
            "• Step 2 支持「文档导入」模式：将场景文档放入 data/raw_docs/ 下即可\n"
            "• Step 3 支持「参考模板」：system_info/ 中的预设模板可约束 Prompt 格式\n"
            "• 版本管理支持断点续接"
        )
        tips.setStyleSheet("font-size: 12px; color: #999; padding: 10px;")
        tips.setAlignment(Qt.AlignLeft)
        layout.addWidget(tips)

        layout.addStretch()


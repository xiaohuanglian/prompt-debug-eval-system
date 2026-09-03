import json
import os
import re

from .json_utils import extract_last_complete_json

try:
    import docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


# 纯文本格式（直接 open 读取）
_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown",
    ".json", ".jsonl",
    ".csv", ".tsv",
    ".py", ".yaml", ".yml",
    ".xml", ".html",
    ".log", ".cfg", ".ini", ".conf",
    ".rst", ".tex",
}

# 所有支持的格式（含二进制格式）
SUPPORTED_EXTENSIONS = _TEXT_EXTENSIONS | {".docx"}


META_PROMPT_ANALYZE_DOCS = """你是一个专业的Prompt工程师。以下是用户提供的一组参考文档，包含了业务规则、数据说明、示例等信息。
请你仔细阅读所有文档内容，然后总结出一份完整的 prompt 需求描述。

# 参考文档

{documents}

{field_schema_block}# 你需要总结的内容

请从以上文档中提取并整理以下信息，形成一份结构化的需求描述：

1. **任务描述**：这个 prompt 要完成什么任务？用一段话清晰描述。
2. **输入字段**：输入有哪些字段？各是什么类型？逐一列出。
3. **输出格式**：输出是什么格式（JSON字段）？给出完整的字段定义。
4. **规则与约束**：有哪些业务规则或约束条件？尽可能完整地列出。
5. **示例**（如果文档中有）：给出输入输出示例。

# 输出要求
- 直接输出总结后的需求描述纯文本，不要用代码块包裹
- 内容要完整、准确，不要遗漏文档中的关键规则
- 如果文档中有矛盾之处，请指出并给出你的理解
- **严禁**自创字段名；输入/输出字段必须完全沿用文档原文或上方"已确认字段表"中给出的名字"""


# ---------- 字段表抽取（用于杜绝 LLM 编造字段） ----------
META_PROMPT_EXTRACT_FIELDS = """你是一个严格的字段抽取器。下面给出业务文档，请抽取该场景下 prompt 真正需要的"输入字段"与"输出字段"。

# 业务文档
{documents}

# 抽取规则（极其重要）
1. 只抽取**文档原文里有据可查**的字段。每个字段必须给出 `source_excerpt`，这必须是一段**逐字（verbatim）摘自文档**的片段（10-120 字符），可用于人工核对来源。
2. **严禁**自创字段、严禁基于"看起来合理"扩写字段。若文档完全没有提及，相关字段**直接省略**，由人工后续补充。
3. 字段名优先沿用文档中出现的英文/中文写法。若文档只给中文别名，用合适的 snake_case 英文映射并把中文别名放进 source_excerpt。
4. 类型限定在：`string` / `integer` / `number` / `boolean` / `array` / `object` / `null` 之一；若文档语焉不详，标 `string` 并在 source_excerpt 中注明"类型未明示"。
5. `required` 字段：文档明示必填→true；明示可空→false；未明示→true（保守）。

# 输出格式（必须为合法 JSON，且仅一个 JSON 对象，用 ```json``` 包裹）
```json
{{
  "input_fields": [
    {{"name": "field_name", "type": "string", "required": true, "source_excerpt": "<verbatim 原文片段>"}}
  ],
  "output_fields": [
    {{"name": "field_name", "type": "string", "required": true, "source_excerpt": "<verbatim 原文片段>"}}
  ]
}}
```
若某一类字段在文档中完全没有依据，对应数组留空 `[]`，不要编造。"""


_FIELD_TYPES = {"string", "integer", "number", "boolean", "array", "object", "null"}


def _normalize_excerpt(text: str) -> str:
    """把原文/excerpt 都做空白归一化，便于做 substring 匹配。"""
    return re.sub(r"\s+", "", text or "")


def _validate_field(field: dict, documents_normalized: str) -> tuple:
    """返回 (是否合法, 原因)。"""
    if not isinstance(field, dict):
        return False, "字段非字典"
    name = field.get("name")
    if not isinstance(name, str) or not name.strip():
        return False, "name 缺失或非字符串"
    excerpt = field.get("source_excerpt")
    if not isinstance(excerpt, str) or not excerpt.strip():
        return False, f"{name}: source_excerpt 缺失"
    if _normalize_excerpt(excerpt) not in documents_normalized:
        return False, f"{name}: source_excerpt 不在文档中（疑似编造）"
    return True, ""


def extract_field_schema(helper_llm: callable, documents: str) -> dict:
    """
    用 helper 模型严格抽取字段表，并按 source_excerpt 是否能在文档里找到做过滤。

    返回:
        {
            "input_fields": [{"name", "type", "required", "source_excerpt"}, ...],
            "output_fields": [...],
            "dropped": [{"name", "reason"}, ...],   # 被过滤掉的字段（疑似编造）
        }

    异常:
        RuntimeError: LLM 无响应或输出无法解析
    """
    prompt = META_PROMPT_EXTRACT_FIELDS.format(documents=documents)
    response = helper_llm(prompt)
    if response is None:
        raise RuntimeError("Helper 模型未返回有效响应。")

    parsed = extract_last_complete_json(response)
    if not isinstance(parsed, dict):
        raise RuntimeError(
            f"无法从模型输出中解析字段表 JSON。模型原始输出（截断 500 字符）:\n{response[:500]}"
        )

    documents_normalized = _normalize_excerpt(documents)

    cleaned_input, cleaned_output, dropped = [], [], []
    for raw in parsed.get("input_fields") or []:
        ok, reason = _validate_field(raw, documents_normalized)
        if ok:
            cleaned_input.append(_canonicalize_field(raw))
        else:
            dropped.append({"name": (raw or {}).get("name", "?"), "kind": "input", "reason": reason})
    for raw in parsed.get("output_fields") or []:
        ok, reason = _validate_field(raw, documents_normalized)
        if ok:
            cleaned_output.append(_canonicalize_field(raw))
        else:
            dropped.append({"name": (raw or {}).get("name", "?"), "kind": "output", "reason": reason})

    return {
        "input_fields": cleaned_input,
        "output_fields": cleaned_output,
        "dropped": dropped,
    }


def _canonicalize_field(field: dict) -> dict:
    """规范化字段：保证 type / required 落在白名单上。"""
    name = field["name"].strip()
    type_str = (field.get("type") or "string").strip().lower()
    if type_str not in _FIELD_TYPES:
        type_str = "string"
    required = field.get("required")
    if not isinstance(required, bool):
        required = True
    return {
        "name": name,
        "type": type_str,
        "required": required,
        "source_excerpt": field["source_excerpt"].strip(),
    }


def format_field_schema_for_review(schema: dict) -> str:
    """把字段表渲染为人类可读的 markdown，供 confirm_or_edit 显示。"""
    parts = []
    has_any_field = False
    for kind, label in (("input_fields", "输入字段"), ("output_fields", "输出字段")):
        fields = schema.get(kind) or []
        parts.append(f"## {label} ({len(fields)} 个)")
        parts.append("| name | type | required | source_excerpt |")
        parts.append("|---|---|---|---|")
        if not fields:
            parts.append("| _（无 —— 文档未明示，可在编辑时按此格式补充行）_ |  |  |  |")
        else:
            has_any_field = True
            for f in fields:
                excerpt = (f.get("source_excerpt") or "").replace("|", "\\|").replace("\n", " ")
                if len(excerpt) > 80:
                    excerpt = excerpt[:80] + "…"
                parts.append(
                    f"| `{f['name']}` | {f['type']} | "
                    f"{'是' if f['required'] else '否'} | {excerpt} |"
                )
        parts.append("")

    dropped = schema.get("dropped") or []
    if dropped:
        parts.append(f"## 已自动剔除的字段 ({len(dropped)} 个，source_excerpt 在文档中找不到，疑似编造)")
        for d in dropped:
            parts.append(f"- [{d['kind']}] `{d['name']}` — {d['reason']}")
        parts.append("")

    parts.append("---")
    parts.append("提示：选择 [E] 编辑时，请直接粘贴上方 markdown 表格（或 JSON）；新增字段必须填写 source_excerpt"
                 "（若来自文档外的人工补充，请用 `[人工补充] <说明>` 起头，便于追溯）。")
    if not has_any_field:
        parts.append("")
        parts.append("由于自动抽取结果为空，下面给出一份可直接复制改写的模板：")
        parts.append("")
        parts.append("```markdown")
        parts.append("## 输入字段")
        parts.append("| name | type | required | source_excerpt |")
        parts.append("|---|---|---|---|")
        parts.append("| `student_name` | string | 是 | [人工补充] 学员姓名 |")
        parts.append("")
        parts.append("## 输出字段")
        parts.append("| name | type | required | source_excerpt |")
        parts.append("|---|---|---|---|")
        parts.append("| `summary` | string | 是 | [人工补充] 训练日小结 |")
        parts.append("```")
        parts.append("")
        parts.append("或等价 JSON：")
        parts.append("```json")
        parts.append('{')
        parts.append('  "input_fields":  [{"name": "student_name", "type": "string", "required": true, "source_excerpt": "[人工补充] 学员姓名"}],')
        parts.append('  "output_fields": [{"name": "summary",      "type": "string", "required": true, "source_excerpt": "[人工补充] 训练日小结"}]')
        parts.append('}')
        parts.append("```")
    return "\n".join(parts)


def parse_field_schema_from_edit(text: str) -> dict:
    """
    解析用户在 confirm_or_edit 阶段提交的字段表。
    优先 JSON，回退 markdown 表格。

    返回与 extract_field_schema 同结构（不含 dropped）。
    """
    text = text.strip()
    if not text:
        raise ValueError("字段表为空")

    # 1) JSON
    parsed = extract_last_complete_json(text)
    if isinstance(parsed, dict) and ("input_fields" in parsed or "output_fields" in parsed):
        return _normalize_parsed_schema(parsed)

    # 2) markdown 表格
    return _parse_markdown_schema(text)


def _normalize_parsed_schema(parsed: dict) -> dict:
    out_input, out_output = [], []
    for raw in parsed.get("input_fields") or []:
        if isinstance(raw, dict) and raw.get("name"):
            out_input.append(_canonicalize_field({
                "name": raw["name"],
                "type": raw.get("type") or "string",
                "required": raw.get("required", True),
                "source_excerpt": raw.get("source_excerpt") or "[人工补充]",
            }))
    for raw in parsed.get("output_fields") or []:
        if isinstance(raw, dict) and raw.get("name"):
            out_output.append(_canonicalize_field({
                "name": raw["name"],
                "type": raw.get("type") or "string",
                "required": raw.get("required", True),
                "source_excerpt": raw.get("source_excerpt") or "[人工补充]",
            }))
    return {"input_fields": out_input, "output_fields": out_output, "dropped": []}


_SECTION_HEADER_RE = re.compile(r"^##\s+(输入字段|输出字段)", re.M)
_TABLE_ROW_RE = re.compile(r"^\|\s*`?([^|`]+?)`?\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*(.*?)\s*\|$", re.M)


def _parse_markdown_schema(text: str) -> dict:
    """简易 markdown 表格解析，按 ## 输入字段 / ## 输出字段 切分。"""
    sections = list(_SECTION_HEADER_RE.finditer(text))
    if not sections:
        raise ValueError("既不是合法 JSON，也找不到 `## 输入字段` / `## 输出字段` 段落")

    blocks = {}
    for i, m in enumerate(sections):
        kind = m.group(1)
        start = m.end()
        end = sections[i + 1].start() if i + 1 < len(sections) else len(text)
        blocks[kind] = text[start:end]

    def parse_block(block: str) -> list:
        rows = []
        for rm in _TABLE_ROW_RE.finditer(block or ""):
            name = rm.group(1).strip()
            type_str = rm.group(2).strip()
            required_str = rm.group(3).strip()
            excerpt = rm.group(4).strip()
            # 跳过表头行
            if name.lower() in {"name", "字段", "字段名"} or set(type_str) <= {"-"}:
                continue
            if not name:
                continue
            required = required_str in {"是", "true", "True", "Y", "y", "1"}
            rows.append(_canonicalize_field({
                "name": name,
                "type": type_str,
                "required": required,
                "source_excerpt": excerpt or "[人工补充]",
            }))
        return rows

    return {
        "input_fields": parse_block(blocks.get("输入字段", "")),
        "output_fields": parse_block(blocks.get("输出字段", "")),
        "dropped": [],
    }


def _render_field_schema_for_meta(schema: dict) -> str:
    """精简版字段表，作为「白名单」喂给 meta-prompt。"""
    lines = ["## 已确认字段表（必须严格沿用 name，不得新增/改名/拆分/合并）"]
    for kind, label in (("input_fields", "输入字段"), ("output_fields", "输出字段")):
        fields = schema.get(kind) or []
        lines.append(f"\n### {label}")
        if not fields:
            lines.append("（无）")
            continue
        for f in fields:
            lines.append(
                f"- `{f['name']}` ({f['type']}, {'必填' if f['required'] else '可空'})"
            )
    return "\n".join(lines) + "\n\n"


def read_raw_docs(scenario_name: str, raw_docs_dir: str = "data/raw_docs") -> str:
    """
    读取 data/raw_docs/{scenario_name}/ 下所有支持的文件，
    拼接为一个结构化的文档字符串。

    参数:
        scenario_name: 场景名称（对应子文件夹名）
        raw_docs_dir: 原始文档根目录

    返回:
        str: 拼接后的文档内容

    异常:
        FileNotFoundError: 目录不存在
        ValueError: 目录下没有可读取的文件
    """
    doc_dir = os.path.join(raw_docs_dir, scenario_name)

    if not os.path.isdir(doc_dir):
        raise FileNotFoundError(
            f"文档目录不存在: {doc_dir}\n"
            f"请创建该目录并放入相关文档（规则描述、数据说明、示例等）。"
        )

    parts = []
    file_count = 0

    # 按文件名排序，保证顺序稳定
    for filename in sorted(os.listdir(doc_dir)):
        filepath = os.path.join(doc_dir, filename)
        if not os.path.isfile(filepath):
            continue

        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            continue

        try:
            if ext == ".docx":
                if not _HAS_DOCX:
                    print(f"  跳过文件 {filename}: 需要安装 python-docx（pip install python-docx）")
                    continue
                doc = docx.Document(filepath)
                content = "\n".join(p.text for p in doc.paragraphs)
            else:
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
        except (UnicodeDecodeError, PermissionError, Exception) as e:
            print(f"  跳过文件 {filename}: {e}")
            continue

        if not content.strip():
            continue

        file_count += 1
        parts.append(f"## 文件: {filename}\n\n{content}")

    if file_count == 0:
        raise ValueError(
            f"目录 {doc_dir} 下没有找到可读取的文档文件。\n"
            f"支持的格式: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    print(f"  已读取 {file_count} 个文件。")
    return "\n\n---\n\n".join(parts)


def read_system_info(system_info_dir: str = "system_info") -> str:
    """
    读取 system_info/ 目录下所有文件，作为系统预设的 Prompt 参考模板和规范。

    返回:
        str: 拼接后的系统信息内容，如果目录不存在或为空则返回空字符串
    """
    if not os.path.isdir(system_info_dir):
        return ""

    parts = []

    for filename in sorted(os.listdir(system_info_dir)):
        filepath = os.path.join(system_info_dir, filename)
        if not os.path.isfile(filepath):
            continue

        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            continue

        try:
            if ext == ".docx":
                if not _HAS_DOCX:
                    continue
                doc = docx.Document(filepath)
                content = "\n".join(p.text for p in doc.paragraphs)
            else:
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
        except Exception:
            continue

        if not content.strip():
            continue

        parts.append(f"## 文件: {filename}\n\n{content}")

    if not parts:
        return ""

    return "\n\n---\n\n".join(parts)


def analyze_docs_to_requirements(helper_llm: callable, documents: str,
                                 field_schema: dict = None) -> str:
    """
    使用 helper 模型分析文档内容，自动提取需求描述。

    参数:
        helper_llm: helper 模型调用函数
        documents: 拼接后的文档内容
        field_schema: 已确认的字段表（来自 extract_field_schema + 人工确认）。
            非空时会作为"白名单"注入到 meta-prompt 中，避免需求总结再次漂走字段名。

    返回:
        str: LLM 总结的需求描述
    """
    if field_schema:
        field_block = _render_field_schema_for_meta(field_schema)
    else:
        field_block = ""

    prompt = META_PROMPT_ANALYZE_DOCS.format(
        documents=documents,
        field_schema_block=field_block,
    )
    response = helper_llm(prompt)

    if response is None:
        raise RuntimeError("Helper 模型未返回有效响应，请检查模型配置。")

    return response.strip()
